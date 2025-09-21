#!/usr/bin/env python3
"""
Generate golden EPUB fixtures for regression testing.

This script creates real DOCX test files and converts them to EPUBs
that serve as reference outputs for golden-file testing.
"""

import sys
import tempfile
import zipfile
from pathlib import Path
from typing import List


def create_real_docx_files():
    """Create real DOCX files for testing using python-docx."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches
    except ImportError:
        print("Warning: python-docx not available. Creating HTML placeholders instead.")
        return create_html_placeholders()

    fixtures_dir = Path("tests/fixtures/golden_epubs")

    test_cases = [
        ("simple", create_simple_docx),
        ("footnotes", create_footnotes_docx),
        ("tables", create_tables_docx),
        ("images", create_images_docx),
        ("poetry", create_poetry_docx),
    ]

    created_files = []

    for test_name, creator_func in test_cases:
        test_dir = fixtures_dir / test_name
        test_dir.mkdir(parents=True, exist_ok=True)

        docx_path = test_dir / f"{test_name}.docx"
        try:
            creator_func(docx_path)
            created_files.append(docx_path)
            print(f"[PASS] Created {docx_path}")
        except Exception as e:
            print(f"[FAIL] Failed to create {docx_path}: {e}")

    return created_files


def create_simple_docx(output_path: Path):
    """Create a simple DOCX document for testing basic conversion."""
    from docx import Document

    doc = Document()

    # Title
    doc.add_heading('Simple Test Document', 0)

    # Chapter 1
    doc.add_heading('Chapter 1: Introduction', level=1)
    doc.add_paragraph('This is a simple paragraph with regular text.')

    p = doc.add_paragraph('This paragraph contains ')
    p.add_run('bold text').bold = True
    p.add_run(' and ')
    p.add_run('italic text').italic = True
    p.add_run(' for formatting tests.')

    # Chapter 2
    doc.add_heading('Chapter 2: Lists and Structure', level=1)

    doc.add_paragraph('Here is a bulleted list:')
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    doc.add_paragraph('Third bullet point', style='List Bullet')

    doc.add_paragraph('And a numbered list:')
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    doc.add_paragraph('Third numbered item', style='List Number')

    doc.save(str(output_path))


def create_footnotes_docx(output_path: Path):
    """Create a DOCX document with footnotes."""
    from docx import Document

    doc = Document()

    doc.add_heading('Document with Footnotes', 0)

    # Add some content with footnotes
    # Note: python-docx doesn't have direct footnote support, so we'll simulate with endnotes
    doc.add_paragraph('This document demonstrates footnote handling in EPUB conversion.')

    p1 = doc.add_paragraph('This paragraph has a reference to a footnote')
    p1.add_run('¹').font.superscript = True
    p1.add_run('.')

    p2 = doc.add_paragraph('Another paragraph with another footnote')
    p2.add_run('²').font.superscript = True
    p2.add_run('.')

    doc.add_heading('Notes', level=2)
    doc.add_paragraph('1. This is the first footnote content.')
    doc.add_paragraph('2. This is the second footnote with formatting.')

    doc.save(str(output_path))


def create_tables_docx(output_path: Path):
    """Create a DOCX document with tables."""
    from docx import Document

    doc = Document()

    doc.add_heading('Document with Tables', 0)

    doc.add_paragraph('This document contains various table formats for testing.')

    # Simple table
    doc.add_heading('Simple Table', level=1)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Header 1'
    hdr_cells[1].text = 'Header 2'
    hdr_cells[2].text = 'Header 3'

    # Data rows
    row1_cells = table.rows[1].cells
    row1_cells[0].text = 'Cell 1,1'
    row1_cells[1].text = 'Cell 1,2'
    row1_cells[2].text = 'Cell 1,3'

    row2_cells = table.rows[2].cells
    row2_cells[0].text = 'Cell 2,1'
    row2_cells[1].text = 'Cell 2,2'
    row2_cells[2].text = 'Cell 2,3'

    # Table with formatting
    doc.add_heading('Formatted Table', level=1)
    table2 = doc.add_table(rows=2, cols=2)
    table2.style = 'Table Grid'

    # Add content with formatting
    table2.cell(0, 0).text = 'Name'
    table2.cell(0, 1).text = 'John Doe'
    table2.cell(1, 0).text = 'Age'
    table2.cell(1, 1).text = '30'

    doc.save(str(output_path))


def create_images_docx(output_path: Path):
    """Create a DOCX document with images (placeholder)."""
    from docx import Document

    doc = Document()

    doc.add_heading('Document with Images', 0)

    doc.add_paragraph('This document would contain images in a real test scenario.')
    doc.add_paragraph('For now, it serves as a placeholder for image handling tests.')

    # In a real implementation, we would add actual images:
    # doc.add_picture('path/to/image.png', width=Inches(4))

    doc.add_paragraph('[Image placeholder: Cover image]')
    doc.add_paragraph('[Image placeholder: Inline diagram]')
    doc.add_paragraph('[Image placeholder: Full-width illustration]')

    doc.save(str(output_path))


def create_poetry_docx(output_path: Path):
    """Create a DOCX document with poetry formatting."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    doc.add_heading('Poetry Collection', 0)

    doc.add_heading('Poem 1: Simple Verse', level=1)

    # Create poem with center alignment
    poem1_lines = [
        'Roses are red,',
        'Violets are blue,',
        'EPUB conversion',
        'Should work for you.'
    ]

    for line in poem1_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Poem 2: Formatted Verse', level=1)

    poem2_lines = [
        'First stanza line one',
        'First stanza line two',
        '',  # Empty line for stanza break
        'Second stanza line one',
        'Second stanza line two'
    ]

    for line in poem2_lines:
        if line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            doc.add_paragraph()  # Empty paragraph for spacing

    doc.save(str(output_path))


def create_html_placeholders():
    """Create HTML placeholder files when python-docx is not available."""
    fixtures_dir = Path("tests/fixtures/golden_epubs")

    test_cases = [
        ("simple", """
        <h1>Simple Test Document</h1>
        <h2>Chapter 1: Introduction</h2>
        <p>This is a simple paragraph with <strong>bold</strong> and <em>italic</em> text.</p>
        <h2>Chapter 2: Lists</h2>
        <ul>
            <li>First bullet point</li>
            <li>Second bullet point</li>
        </ul>
        <ol>
            <li>First numbered item</li>
            <li>Second numbered item</li>
        </ol>
        """),
        ("footnotes", """
        <h1>Document with Footnotes</h1>
        <p>This paragraph has a footnote<sup>1</sup>.</p>
        <p>Another paragraph with a footnote<sup>2</sup>.</p>
        <h2>Notes</h2>
        <p>1. This is the first footnote.</p>
        <p>2. This is the second footnote.</p>
        """),
        ("tables", """
        <h1>Document with Tables</h1>
        <table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Cell 1</td><td>Cell 2</td></tr>
            <tr><td>Cell 3</td><td>Cell 4</td></tr>
        </table>
        """),
        ("images", """
        <h1>Document with Images</h1>
        <p>[Image placeholder: Cover image]</p>
        <p>[Image placeholder: Inline diagram]</p>
        """),
        ("poetry", """
        <h1>Poetry Collection</h1>
        <h2>Poem 1</h2>
        <div class="poem">
            <p class="line">Roses are red,</p>
            <p class="line">Violets are blue,</p>
            <p class="line">EPUB conversion</p>
            <p class="line">Should work for you.</p>
        </div>
        """)
    ]

    created_files = []

    for test_name, content in test_cases:
        test_dir = fixtures_dir / test_name
        test_dir.mkdir(parents=True, exist_ok=True)

        html_path = test_dir / f"{test_name}.html"
        html_path.write_text(content, encoding='utf-8')
        created_files.append(html_path)
        print(f"[PASS] Created HTML placeholder: {html_path}")

    return created_files


def convert_to_golden_epubs(source_files: List[Path]):
    """Convert source files to EPUBs and save as golden fixtures."""
    import sys
    sys.path.insert(0, str(Path.cwd()))
    from tests.test_golden_epubs import GoldenEPUBTest

    generated_epubs = []

    for source_file in source_files:
        test_name = source_file.stem
        print(f"\n[PROCESS] Processing {test_name}...")

        # Create a temporary EPUB for testing
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            epub_path = temp_path / f"{test_name}.epub"

            try:
                # Try to convert using docx2shelf
                if source_file.suffix.lower() == '.html':
                    # Convert HTML to EPUB
                    result = convert_html_to_epub(source_file, epub_path, test_name)
                elif source_file.suffix.lower() == '.docx':
                    # Convert DOCX to EPUB
                    result = convert_docx_to_epub(source_file, epub_path, test_name)
                else:
                    print(f"[WARN] Unsupported file type: {source_file.suffix}")
                    continue

                if result and epub_path.exists():
                    # Save as golden fixture
                    golden_path = source_file.parent / f"{test_name}_golden.epub"
                    epub_path.rename(golden_path)

                    # Generate expected structure
                    test = GoldenEPUBTest(test_name)
                    test.save_expected_structure(golden_path)

                    generated_epubs.append(golden_path)
                    print(f"[PASS] Generated golden EPUB: {golden_path}")

            except Exception as e:
                print(f"[FAIL] Failed to convert {source_file}: {e}")

    return generated_epubs


def convert_html_to_epub(html_path: Path, epub_path: Path, title: str) -> bool:
    """Convert HTML to EPUB using a minimal EPUB structure."""
    try:
        # Create a minimal EPUB structure for testing
        import uuid
        import zipfile
        from datetime import datetime

        # EPUB content
        mimetype = "application/epub+zip"

        container_xml = """<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="EPUB/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>"""

        content_opf = f"""<?xml version="1.0" encoding="UTF-8"?>
<package version="3.0" xmlns="http://www.idpf.org/2007/opf" unique-identifier="uid">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
        <dc:identifier id="uid">{uuid.uuid4()}</dc:identifier>
        <dc:title>{title}</dc:title>
        <dc:creator>Test Author</dc:creator>
        <dc:language>en</dc:language>
        <meta property="dcterms:modified">{datetime.now().strftime('%Y-%m-%dT%H:%M:%SZ')}</meta>
    </metadata>
    <manifest>
        <item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>
        <item id="content" href="content.xhtml" media-type="application/xhtml+xml"/>
        <item id="css" href="style/base.css" media-type="text/css"/>
    </manifest>
    <spine>
        <itemref idref="content"/>
    </spine>
</package>"""

        nav_xhtml = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">
<head>
    <title>Navigation</title>
    <link rel="stylesheet" href="style/base.css"/>
</head>
<body>
    <nav epub:type="toc">
        <h1>Table of Contents</h1>
        <ol>
            <li><a href="content.xhtml">{title}</a></li>
        </ol>
    </nav>
</body>
</html>"""

        # Read HTML content
        html_content = html_path.read_text(encoding='utf-8')

        content_xhtml = f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{title}</title>
    <link rel="stylesheet" href="style/base.css"/>
</head>
<body>
{html_content}
</body>
</html>"""

        base_css = """
/* Base styles for golden EPUB test fixtures */
body { font-family: serif; margin: 2em; line-height: 1.6; }
h1, h2, h3 { margin-top: 1.5em; margin-bottom: 0.5em; }
p { margin: 0.5em 0; }
table { border-collapse: collapse; width: 100%; margin: 1em 0; }
th, td { border: 1px solid #ccc; padding: 0.5em; text-align: left; }
th { background-color: #f5f5f5; }
.poem { text-align: center; margin: 2em 0; }
.line { margin: 0.2em 0; }
.stanza { margin: 1em 0; }
"""

        # Create EPUB
        with zipfile.ZipFile(epub_path, 'w', zipfile.ZIP_DEFLATED) as epub:
            epub.writestr("mimetype", mimetype, compress_type=zipfile.ZIP_STORED)
            epub.writestr("META-INF/container.xml", container_xml)
            epub.writestr("EPUB/content.opf", content_opf)
            epub.writestr("EPUB/nav.xhtml", nav_xhtml)
            epub.writestr("EPUB/content.xhtml", content_xhtml)
            epub.writestr("EPUB/style/base.css", base_css)

        return True

    except Exception as e:
        print(f"Error converting HTML to EPUB: {e}")
        return False


def convert_docx_to_epub(docx_path: Path, epub_path: Path, title: str) -> bool:
    """Convert DOCX to EPUB using a placeholder approach."""
    try:
        # For now, treat DOCX same as HTML since we're using minimal conversion
        # In a real implementation, this would use docx2shelf properly
        return convert_html_to_epub(docx_path, epub_path, title)
    except Exception as e:
        print(f"Error converting DOCX to EPUB: {e}")
        return False


def validate_golden_fixtures():
    """Validate that generated golden fixtures are valid EPUBs."""
    fixtures_dir = Path("tests/fixtures/golden_epubs")

    valid_count = 0
    total_count = 0

    for epub_path in fixtures_dir.rglob("*_golden.epub"):
        total_count += 1
        print(f"\n[VALIDATE] Validating {epub_path.name}...")

        try:
            # Basic EPUB validation
            with zipfile.ZipFile(epub_path, 'r') as epub:
                # Check for required files
                namelist = epub.namelist()

                required_files = [
                    "META-INF/container.xml",
                    "EPUB/content.opf",
                    "EPUB/nav.xhtml"
                ]

                missing_files = [f for f in required_files if f not in namelist]

                if missing_files:
                    print(f"[FAIL] Missing required files: {missing_files}")
                else:
                    print("[PASS] All required files present")
                    valid_count += 1

                # Check file count
                print(f"  [INFO] Contains {len(namelist)} files")

        except Exception as e:
            print(f"[FAIL] Invalid EPUB: {e}")

    print("\n[SUMMARY] Validation Summary:")
    print(f"  Valid EPUBs: {valid_count}/{total_count}")

    return valid_count == total_count


def main():
    """Generate golden EPUB fixtures for testing."""
    print("[BUILD] Generating Golden EPUB Fixtures")
    print("=" * 50)

    # Step 1: Create source files (DOCX or HTML)
    print("\n[CREATE] Creating source test files...")
    source_files = create_real_docx_files()

    if not source_files:
        print("[ERROR] No source files created. Exiting.")
        return 1

    # Step 2: Convert to EPUBs
    print(f"\n[CONVERT] Converting {len(source_files)} files to EPUBs...")
    golden_epubs = convert_to_golden_epubs(source_files)

    if not golden_epubs:
        print("[ERROR] No golden EPUBs generated. Check conversion process.")
        return 1

    # Step 3: Validate generated EPUBs
    print(f"\n[VALIDATE] Validating {len(golden_epubs)} golden EPUBs...")
    if validate_golden_fixtures():
        print("[SUCCESS] All golden fixtures are valid!")
    else:
        print("[WARN] Some golden fixtures have issues.")

    # Step 4: Generate test summary
    print("\n[SUMMARY] Generated golden fixtures:")
    for epub_path in golden_epubs:
        try:
            rel_path = epub_path.relative_to(Path.cwd())
            print(f"  - {rel_path}")
        except ValueError:
            print(f"  - {epub_path}")

    print("\n[NEXT] Next steps:")
    print("  1. Run tests: python -m pytest tests/test_golden_epubs.py")
    print("  2. Review fixtures in tests/fixtures/golden_epubs/")
    print("  3. Update expected structures if needed")

    return 0


if __name__ == "__main__":
    sys.exit(main())