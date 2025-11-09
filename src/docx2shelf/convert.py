from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from .path_utils import get_safe_temp_path

IMG_NS = {
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
}


def split_html_by_heading(html: str, level: str) -> list[str]:
    """Split a single HTML string into chunks at <h1> or <h2> boundaries.

    Uses a simple regex; good enough for initial implementation and Pandoc output.
    """
    tag = level.lower()
    # Normalize newlines to avoid regex surprises
    s = html.replace("\r\n", "\n")
    # Split but keep the heading with the following content
    pattern = re.compile(rf"(<{tag}[^>]*>.*?</{tag}>)", re.IGNORECASE | re.DOTALL)
    parts = pattern.split(s)
    if len(parts) <= 1:
        return [s]
    chunks: list[str] = []
    current: list[str] = []
    for i, part in enumerate(parts):
        if i % 2 == 1:  # this is a heading
            if current:
                chunks.append("".join(current))
                current = []
            current.append(part)
        else:
            current.append(part)
    if current:
        chunks.append("".join(current))
    # Wrap chunks into section tags for cleanliness
    return [f"<section>{c}</section>" for c in chunks if c.strip()]


def split_html_by_heading_level(html: str, level: str) -> list[str]:
    """Split HTML by a specific heading level (h3, h4, h5, h6)."""
    if level not in ["h3", "h4", "h5", "h6"]:
        return split_html_by_heading(html, level)

    tag = level.lower()
    s = html.replace("\r\n", "\n")
    pattern = re.compile(rf"(<{tag}[^>]*>.*?</{tag}>)", re.IGNORECASE | re.DOTALL)
    parts = pattern.split(s)
    if len(parts) <= 1:
        return [s]
    chunks: list[str] = []
    current: list[str] = []
    for i, part in enumerate(parts):
        if i % 2 == 1:  # this is a heading
            if current:
                chunks.append("".join(current))
                current = []
            current.append(part)
        else:
            current.append(part)
    if current:
        chunks.append("".join(current))
    return [f"<section>{c}</section>" for c in chunks if c.strip()]


def split_html_mixed(html: str, mixed_pattern: str) -> list[str]:
    """Split HTML using mixed strategy based on pattern.

    Pattern examples:
    - 'h1,pagebreak' - Split at h1 OR pagebreak
    - 'h1:main,pagebreak:appendix' - Split at h1 for main content, pagebreak for appendix
    """
    if not mixed_pattern:
        return split_html_by_heading(html, "h1")

    # Parse the pattern
    strategies = []
    for part in mixed_pattern.split(","):
        part = part.strip()
        if ":" in part:
            strategy, section = part.split(":", 1)
            strategies.append((strategy.strip(), section.strip()))
        else:
            strategies.append((part, None))

    # For now, implement simple OR logic - split at any of the specified points
    html_sections = [html]

    for strategy, section_type in strategies:
        new_sections = []
        for section in html_sections:
            if strategy == "pagebreak":
                parts = split_html_by_pagebreak(section)
            elif strategy in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                if strategy in ["h1", "h2"]:
                    parts = split_html_by_heading(section, strategy)
                else:
                    parts = split_html_by_heading_level(section, strategy)
            else:
                parts = [section]
            new_sections.extend(parts)
        html_sections = new_sections

    return html_sections


def split_html_by_pagebreak(html: str) -> list[str]:
    """Split HTML at common pagebreak markers.

    Looks for <hr class="pagebreak">, elements with style containing
    page-break-(before|after): always, or explicit <!-- PAGEBREAK --> comments.
    """
    s = html.replace("\r\n", "\n")
    # Insert a sentinel at break points
    patterns = [
        r"<hr[^>]*class=\"[^\"]*pagebreak[^\"]*\"[^>]*/?>",
        r"<[^>]*style=\"[^\"]*page-break-(before|after)\s*:\s*always[^\"]*\"[^>]*>",
        r"<!--\s*PAGEBREAK\s*-->",
    ]
    sentinel = "\n<!--__SPLIT__-->\n"
    for pat in patterns:
        s = re.sub(pat, sentinel, s, flags=re.IGNORECASE)
    parts = s.split(sentinel)
    chunks = [p for p in parts if p.strip()]
    return [f"<section>{c}</section>" for c in chunks]


def convert_file_to_html(input_path: Path, context: dict | None = None) -> tuple[list[str], list[Path], str]:
    """Convert input file to HTML chunks and gather any extracted resources.

    Strategy:
    - For .md and .txt, use Pandoc.
    - For .docx, try Pandoc first, then fall back to python-docx.
    - Uses performance optimizations for large files.
    """
    from .plugins import plugin_manager, load_default_plugins
    from .performance import PerformanceMonitor, BuildCache, ParallelImageProcessor

    # Initialize context if not provided
    if context is None:
        context = {}

    # Initialize performance monitoring
    monitor = PerformanceMonitor()
    monitor.start_monitoring()

    # Check for build cache
    cache_dir = Path.home() / ".docx2shelf" / "cache"
    cache = BuildCache(cache_dir)

    # Initialize image processor
    image_processor = ParallelImageProcessor()
    image_processor.set_cache(cache)

    # Load plugins
    load_default_plugins()

    # Execute pre-convert hooks for DOCX files
    actual_input_path = input_path
    if input_path.suffix.lower() == ".docx":
        actual_input_path = plugin_manager.execute_pre_convert_hooks(input_path, context)

    suffix = actual_input_path.suffix.lower()

    if suffix in (".md", ".txt", ".html", ".htm"):
        # Check Pandoc availability with detailed error messages
        from .tools import get_pandoc_status

        status = get_pandoc_status()
        if not status["overall_available"]:
            error_msg = f"Pandoc is required to convert {suffix} files.\n"

            if not status["pandoc_binary"]["available"]:
                error_msg += f"Issue: {status['pandoc_binary']['message']}\n"
                error_msg += "Solution: Run 'docx2shelf tools install pandoc'\n"

            if not status["pypandoc_library"]["available"]:
                error_msg += f"Issue: {status['pypandoc_library']['message']}\n"
                error_msg += "Solution: Run 'pip install pypandoc'\n"

            raise RuntimeError(error_msg.strip())

        try:
            import pypandoc  # type: ignore

            if suffix in (".html", ".htm"):
                file_format = "html"
            else:
                file_format = "markdown" if suffix == ".md" else "plain"

            html = pypandoc.convert_file(
                str(actual_input_path), to="html", format=file_format, extra_args=["--wrap=none"]
            )
            # Apply post-convert hooks
            processed_html = plugin_manager.execute_post_convert_hooks(html, context)
            # For now, we don't split these files, return as a single chunk
            return [f"<section>{processed_html}</section>"], [], ""
        except Exception as e:
            error_msg = f"Pandoc conversion failed for {suffix} file: {e}\n"
            error_msg += "This might be due to:\n"
            error_msg += "- Unsupported content in the input file\n"
            error_msg += "- Pandoc version compatibility issues\n"
            error_msg += "- File encoding problems\n"
            error_msg += f"Run 'docx2shelf doctor' to check your Pandoc installation"
            raise RuntimeError(error_msg)

    elif suffix == ".docx":
        # Check cache first
        cache_key = cache.generate_cache_key(actual_input_path)
        cached_result = cache.get_cached_conversion(cache_key)

        if cached_result:
            monitor.add_phase_time("cache_hit", 0.0)
            chunks, resources, styles = cached_result
        else:
            # Use performance-optimized conversion
            with monitor.phase_timer("docx_conversion"):
                chunks, resources, styles = docx_to_html_optimized(
                    actual_input_path, cache, image_processor, monitor
                )

            # Cache the result
            cache.cache_conversion(cache_key, (chunks, resources, styles))

        # Apply post-convert hooks to each chunk with parallel processing
        with monitor.phase_timer("post_processing"):
            processed_chunks = []
            for chunk in chunks:
                processed_chunk = plugin_manager.execute_post_convert_hooks(chunk, context)
                processed_chunks.append(processed_chunk)

        # Finalize monitoring
        monitor.stop_monitoring()

        return processed_chunks, resources, styles

    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _process_tracked_changes(run_element) -> tuple[str, bool]:
    """Process tracked changes in a run element.

    Returns:
        tuple: (processed_text, should_include)
    """
    # Handle insertions (w:ins)
    insertions = run_element.findall('.//w:ins', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if insertions:
        # Extract text from insertions and mark as accepted
        text_parts = []
        for ins in insertions:
            for t_elem in ins.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                if t_elem.text:
                    text_parts.append(t_elem.text)
        return ' '.join(text_parts), True

    # Handle deletions (w:del) - skip by default but could be made configurable
    deletions = run_element.findall('.//w:del', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    if deletions:
        return "", False

    # Handle move operations (w:moveFrom, w:moveTo)
    move_from = run_element.findall('.//w:moveFrom', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    move_to = run_element.findall('.//w:moveTo', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

    if move_from:
        return "", False  # Skip move source
    if move_to:
        # Extract text from move destination
        text_parts = []
        for move in move_to:
            for t_elem in move.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                if t_elem.text:
                    text_parts.append(t_elem.text)
        return ' '.join(text_parts), True

    return "", True  # No tracked changes found


def _process_comments(run_element, document) -> str:
    """Process comment references in a run element.

    Returns:
        HTML string with comment markers
    """
    comment_refs = run_element.findall('.//w:commentRangeStart', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
    comment_refs.extend(run_element.findall('.//w:commentReference', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}))

    if not comment_refs:
        return ""

    comment_html = []
    for ref in comment_refs:
        comment_id = ref.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
        if comment_id:
            try:
                # Try to extract comment text if comments part exists
                if hasattr(document.part, 'comments_part') and document.part.comments_part:
                    comment_nodes = document.part.comments_part._element.xpath(
                        f'.//w:comment[@w:id="{comment_id}"]//w:t',
                        namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                    )
                    comment_text = ' '.join(node.text or '' for node in comment_nodes).strip()
                    if comment_text:
                        comment_html.append(f'<span class="comment" title="{comment_text}">💬</span>')
            except Exception:
                # Fallback to just marking comment presence
                comment_html.append(f'<span class="comment">💬</span>')

    return ''.join(comment_html)


def _process_table(table_element) -> str:
    """Process a table element into HTML.

    Returns:
        HTML table string
    """
    try:
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        # Create a temporary table object to work with
        table = Table(table_element, None)

        html_parts = ['<table>']

        for row in table.rows:
            html_parts.append('<tr>')
            for cell in row.cells:
                # Process cell content
                cell_content = []
                for paragraph in cell.paragraphs:
                    p_text = paragraph.text.strip()
                    if p_text:
                        cell_content.append(f'<p>{p_text}</p>')

                cell_html = ''.join(cell_content) if cell_content else '<p></p>'
                html_parts.append(f'<td>{cell_html}</td>')
            html_parts.append('</tr>')

        html_parts.append('</table>')
        return ''.join(html_parts)

    except Exception as e:
        # Fallback to simple text representation
        return f'<p><em>[Table content - processing error: {str(e)}]</em></p>'


def _rasterize_complex_element(element, element_type: str, tempdir: Path) -> str:
    """Rasterize complex elements that can't be cleanly translated to HTML.

    This is an optional fallback that creates an image representation
    of layout elements to preserve author intention.
    """
    try:
        from xml.etree import ElementTree as ET
        import hashlib

        # Extract basic properties from the element
        element_id = f"{element_type}_{hashlib.md5(str(element).encode()).hexdigest()[:8]}"

        # Try to extract text content if available
        text_content = ""
        if hasattr(element, 'text') and element.text:
            text_content = element.text.strip()
        elif hasattr(element, 'element'):
            # Try to extract text from Word XML
            try:
                for text_elem in element.element.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if text_elem.text:
                        text_content += text_elem.text
            except (AttributeError, IndexError):
                # Fallback if text extraction fails
                pass

        # Create semantic markup based on element type
        if element_type.lower() == 'textbox':
            css_class = "text-box"
            if text_content:
                return f'<aside class="{css_class}" data-id="{element_id}"><p>{text_content}</p></aside>'
            else:
                return f'<aside class="{css_class}" data-id="{element_id}"><p>[Text Box Content]</p></aside>'

        elif element_type.lower() in ('shape', 'drawing'):
            css_class = "drawing-element"
            if text_content:
                return f'<figure class="{css_class}" data-id="{element_id}"><figcaption>{text_content}</figcaption></figure>'
            else:
                return f'<figure class="{css_class}" data-id="{element_id}"><figcaption>[Drawing Element]</figcaption></figure>'

        else:
            # Generic complex element
            css_class = "complex-element"
            content = text_content if text_content else f"[{element_type}]"
            return f'<div class="{css_class}" data-type="{element_type}" data-id="{element_id}">{content}</div>'

    except Exception:
        return f'<div class="complex-element" data-type="{element_type}">[{element_type}]</div>'


def _process_text_box_or_shape(element, tempdir: Path, rasterize_fallback: bool = True) -> str:
    """Process text boxes and shapes into HTML.

    Returns:
        HTML representation of the content
    """
    try:
        # Extract text from text boxes and shapes
        text_elements = element.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        text_content = ' '.join(elem.text or '' for elem in text_elements).strip()

        # Check for complex layout properties that might require rasterization
        has_complex_layout = False
        try:
            # Check for positioning, rotation, 3D effects, etc.
            drawing_elements = element.findall('.//wp:anchor', namespaces=IMG_NS)
            drawing_elements.extend(element.findall('.//wp:inline', namespaces=IMG_NS))
            if drawing_elements:
                has_complex_layout = True
        except Exception:
            pass

        if text_content:
            if has_complex_layout and rasterize_fallback:
                # Use rasterization for complex layouts
                return _rasterize_complex_element(element, "text-box", tempdir)
            else:
                return f'<div class="text-box"><p>{text_content}</p></div>'
        else:
            if has_complex_layout and rasterize_fallback:
                return _rasterize_complex_element(element, "shape", tempdir)
            else:
                return '<div class="shape">[Shape or text box]</div>'
    except Exception:
        if rasterize_fallback:
            return _rasterize_complex_element(element, "unknown", tempdir)
        else:
            return '<div class="shape">[Shape or text box]</div>'


def _load_style_mapping(docx_path: Path) -> dict:
    """Load style mapping from default styles.json and optional user override.

    Checks for:
    1. Default styles.json in package
    2. User styles.json in same directory as DOCX file
    3. User styles.json in current working directory
    """
    # Start with default styles
    default_styles_path = Path(__file__).parent / "styles.json"
    styles_data = {}

    if default_styles_path.exists():
        try:
            with open(default_styles_path, 'r', encoding='utf-8') as f:
                styles_data = json.load(f)
        except Exception as e:
            print(f"Warning: Could not load default styles.json: {e}", file=sys.stderr)

    # Check for user override in same directory as DOCX
    user_styles_path = docx_path.parent / "styles.json"
    if user_styles_path.exists():
        try:
            with open(user_styles_path, 'r', encoding='utf-8') as f:
                user_styles = json.load(f)
                # Merge user styles with defaults
                for category in ["paragraph_styles", "run_styles", "character_styles", "css_classes"]:
                    if category in user_styles:
                        if category not in styles_data:
                            styles_data[category] = {}
                        styles_data[category].update(user_styles[category])
            print(f"Loaded user style overrides from {user_styles_path}", file=sys.stderr)
        except Exception as e:
            print(f"Warning: Could not load user styles.json: {e}", file=sys.stderr)

    # Check for user override in current working directory
    cwd_styles_path = Path.cwd() / "styles.json"
    if cwd_styles_path.exists() and cwd_styles_path != user_styles_path:
        try:
            with open(cwd_styles_path, 'r', encoding='utf-8') as f:
                cwd_styles = json.load(f)
                # Merge with existing styles
                for category in ["paragraph_styles", "run_styles", "character_styles", "css_classes"]:
                    if category in cwd_styles:
                        if category not in styles_data:
                            styles_data[category] = {}
                        styles_data[category].update(cwd_styles[category])
            print(f"Loaded additional style overrides from {cwd_styles_path}", file=sys.stderr)
        except Exception as e:
            print(f"Warning: Could not load working directory styles.json: {e}", file=sys.stderr)

    return styles_data


def extract_styles_css(styles_data: dict) -> str:
    """Extract CSS rules from styles.json css_classes section."""
    css_rules = []
    css_classes = styles_data.get("css_classes", {})

    for class_name, class_data in css_classes.items():
        if isinstance(class_data, dict) and "css" in class_data:
            css_rules.append(class_data["css"])

    return "\n".join(css_rules)


def _process_equation(element) -> str:
    """Process mathematical equations.

    Returns:
        HTML representation of the equation
    """
    try:
        # Try to extract equation text content
        text_elements = element.findall('.//m:t', namespaces={'m': 'http://schemas.openxmlformats.org/officeDocument/2006/math'})
        if not text_elements:
            text_elements = element.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})

        equation_text = ' '.join(elem.text or '' for elem in text_elements).strip()

        if equation_text:
            return f'<span class="equation">{equation_text}</span>'
        else:
            return '<span class="equation">[Equation]</span>'
    except Exception:
        return '<span class="equation">[Equation]</span>'


def docx_to_html_optimized(docx_path: Path, cache, image_processor, monitor) -> tuple[list[str], list[Path], str]:
    """Convert DOCX to HTML chunks with performance optimizations.

    Uses streaming reading, parallel image processing, and performance monitoring.
    """
    from .performance import StreamingDocxReader, MemoryOptimizer

    # Initialize memory optimizer
    optimizer = MemoryOptimizer()
    optimizer.start_monitoring()

    try:
        # Try Pandoc first (fastest for most documents)
        with monitor.phase_timer("pandoc_conversion"):
            # Check Pandoc availability
            from .tools import get_pandoc_status
            status = get_pandoc_status()

            if status["overall_available"]:
                try:
                    import pypandoc  # type: ignore
                    html = pypandoc.convert_file(str(docx_path), to="html", extra_args=["--wrap=none"])
                    chunks = split_html_by_heading(html, level="h1")

                    # Extract and process images in parallel
                    with monitor.phase_timer("image_processing"):
                        resources = image_processor.process_images(docx_path)

                    # Load styles
                    with monitor.phase_timer("style_loading"):
                        styles_data = _load_style_mapping(docx_path)
                        styles_css = extract_styles_css(styles_data)

                    return chunks, resources, styles_css
                except Exception as e:
                    monitor.add_warning(f"Pandoc conversion failed: {e}")
            else:
                # Pandoc not available, log the reason
                if not status["pandoc_binary"]["available"]:
                    monitor.add_warning(f"Pandoc binary unavailable: {status['pandoc_binary']['message']}")
                if not status["pypandoc_library"]["available"]:
                    monitor.add_warning(f"pypandoc library unavailable: {status['pypandoc_library']['message']}")
                monitor.add_warning("Using fallback DOCX conversion - install Pandoc for better results")

        # Fallback to streaming python-docx reader for large files
        with monitor.phase_timer("streaming_conversion"):
            return _docx_to_html_streaming(docx_path, image_processor, monitor, optimizer)

    finally:
        optimizer.cleanup()


def _docx_to_html_streaming(docx_path: Path, image_processor, monitor, optimizer) -> tuple[list[str], list[Path], str]:
    """Convert DOCX to HTML using streaming approach for large documents."""
    from .performance import StreamingDocxReader
    import xml.etree.ElementTree as ET

    # Use streaming reader to minimize memory usage
    with StreamingDocxReader(docx_path, chunk_size=512 * 1024) as reader:  # 512KB chunks
        # Process document XML in chunks
        chunks = []
        current_chunk = []
        resources = []

        # Extract embedded images in parallel
        with monitor.phase_timer("image_extraction"):
            try:
                image_data = reader.get_embedded_images()
                resources = image_processor.process_image_data(image_data)
            except Exception as e:
                monitor.add_warning(f"Image extraction failed: {e}")

        # Stream and parse document XML
        with monitor.phase_timer("xml_streaming"):
            full_xml = ""
            for xml_chunk in reader.stream_document_xml():
                full_xml += xml_chunk
                # Trigger garbage collection periodically for large documents
                if len(full_xml) > 5 * 1024 * 1024:  # 5MB
                    optimizer.trigger_gc()

        # Parse and convert to HTML
        with monitor.phase_timer("html_conversion"):
            try:
                # Use the existing python-docx fallback but with optimized parsing
                chunks, _, styles = _parse_docx_xml_optimized(full_xml, docx_path, monitor)
            except Exception as e:
                monitor.add_warning(f"XML parsing failed: {e}")
                # Ultra-minimal fallback
                chunks = ["<section><p>Document content could not be parsed</p></section>"]
                styles = ""

        return chunks, resources, styles


def _parse_docx_xml_optimized(xml_content: str, docx_path: Path, monitor) -> tuple[list[str], list[Path], str]:
    """Parse DOCX XML with memory and performance optimizations."""
    import xml.etree.ElementTree as ET
    from xml.etree.ElementTree import ParseError

    try:
        # Parse XML with optimized settings
        root = ET.fromstring(xml_content)

        # Use the existing docx_to_html logic but with streaming optimizations
        # For now, fall back to the existing implementation but monitor performance
        with monitor.phase_timer("legacy_conversion"):
            chunks, resources, styles = _fallback_docx_conversion(docx_path)

        return chunks, resources, styles

    except ParseError as e:
        monitor.add_warning(f"XML parsing error: {e}")
        return ["<section><p>Error parsing document XML</p></section>"], [], ""


def _fallback_docx_conversion(docx_path: Path) -> tuple[list[str], list[Path], str]:
    """Fallback conversion using existing logic."""
    # Use the existing docx_to_html function logic as fallback
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore

        doc = Document(docx_path)
        chunks = []
        current_chunk = []
        resources = []

        # Simple paragraph-by-paragraph conversion
        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                if current_chunk:
                    chunks.append(f"<section>{''.join(current_chunk)}</section>")
                    current_chunk = []
                level = 1 if 'Heading 1' in para.style.name else 2
                current_chunk.append(f"<h{level}>{para.text}</h{level}>")
            else:
                current_chunk.append(f"<p>{para.text}</p>")

        if current_chunk:
            chunks.append(f"<section>{''.join(current_chunk)}</section>")

        if not chunks:
            chunks = ["<section><p>No content found</p></section>"]

        # Load styles
        styles_data = _load_style_mapping(docx_path)
        styles_css = extract_styles_css(styles_data)

        return chunks, resources, styles_css

    except Exception as e:
        return ["<section><p>Error processing document</p></section>"], [], ""


def docx_to_html(docx_path: Path) -> tuple[list[str], list[Path], str]:
    """Convert DOCX to HTML chunks and gather any extracted resources.

    Strategy:
    1) Try Pandoc via pypandoc for rich conversion.
    2) Fallback to a lightweight python-docx paragraph/headings extraction.
    """
    # Load style mapping with user override support
    styles_data = _load_style_mapping(docx_path)
    paragraph_styles_map = styles_data.get("paragraph_styles", {})
    run_styles_map = styles_data.get("run_styles", {})
    character_styles_map = styles_data.get("character_styles", {})
    css_classes = styles_data.get("css_classes", {})
    # Check Pandoc availability first
    from .tools import get_pandoc_status
    status = get_pandoc_status()

    if status["overall_available"]:
        try:
            import pypandoc  # type: ignore

            html = pypandoc.convert_file(str(docx_path), to="html", extra_args=["--wrap=none"])
            # Split at h1 by default; caller can later decide via CLI how to split
            chunks = split_html_by_heading(html, level="h1")
            # Load styles for potential CSS injection even with Pandoc
            styles_data = _load_style_mapping(docx_path)
            styles_css = extract_styles_css(styles_data)
            return chunks, [], styles_css
        except Exception as e:
            print(f"Warning: Pandoc conversion failed ({e}), using fallback")
    else:
        # Log why Pandoc is not available
        print("Warning: Pandoc not available, using fallback DOCX conversion")
        if not status["pandoc_binary"]["available"]:
            print(f"  - {status['pandoc_binary']['message']}")
        if not status["pypandoc_library"]["available"]:
            print(f"  - {status['pypandoc_library']['message']}")
        print("  - Run 'docx2shelf doctor' for detailed diagnostics")

    # 2) python-docx fallback
    try:
        from docx import Document  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception as e:
        raise RuntimeError(
            "No converter found. Install pypandoc (Pandoc recommended) or python-docx."
        ) from e

    document = Document(str(docx_path))
    parts: list[str] = []

    # Temp dir for extracted images
    # Use safe temp directory with Unicode support
    tempdir = get_safe_temp_path("docx2shelf_pandoc")
    tempdir.mkdir(parents=True, exist_ok=True)
    images: dict[str, Path] = {}

    # Process document elements including tables
    document_elements = []
    for element in document.element.body:
        if element.tag.endswith('}p'):  # Paragraph
            document_elements.append(('paragraph', element))
        elif element.tag.endswith('}tbl'):  # Table
            document_elements.append(('table', element))
        elif element.tag.endswith('}sectPr'):  # Section properties
            continue
        else:
            # Handle other elements like text boxes, shapes, etc.
            document_elements.append(('other', element))

    def flush_section(buf: list[str], footnotes: list[str] | None = None):
        if not buf:
            return
        body = "".join(buf)
        if footnotes:
            notes_html = (
                '<hr/><section class="footnotes"><ol>' + "".join(footnotes) + "</ol></section>"
            )
            body += notes_html
        parts.append("<section>" + body + "</section>")
        buf.clear()

    buf: list[str] = []
    note_idx = 0
    current_notes: list[str] = []
    # List handling state
    current_list_type: str | None = None  # 'ul' | 'ol'
    list_items: list[str] = []

    def flush_list():
        nonlocal current_list_type, list_items
        if current_list_type and list_items:
            buf.append(f"<{current_list_type}>" + "".join(list_items) + f"</{current_list_type}>")
        current_list_type = None
        list_items = []

    pending_img: str | None = None

    def _get_run_html(run, run_styles_map, initial_txt):
        txt = initial_txt
        # Apply basic inline formatting
        if getattr(run, "bold", False):
            txt = f"<strong>{txt}</strong>"
        if getattr(run, "italic", False):
            txt = f"<em>{txt}</em>"

        # Enhanced run style mapping
        try:
            # Check for underline
            if hasattr(run, 'underline') and run.underline:
                txt = f"<u>{txt}</u>"

            # Check for strike-through
            if hasattr(run, 'font') and hasattr(run.font, 'strike') and run.font.strike:
                txt = f"<s>{txt}</s>"

            # Check for superscript/subscript
            if hasattr(run, 'font') and hasattr(run.font, 'superscript') and run.font.superscript:
                txt = f"<sup>{txt}</sup>"
            elif hasattr(run, 'font') and hasattr(run.font, 'subscript') and run.font.subscript:
                txt = f"<sub>{txt}</sub>"

            # Check for small caps
            if hasattr(run, 'font') and hasattr(run.font, 'small_caps') and run.font.small_caps:
                txt = f'<span class="small-caps">{txt}</span>'

            # Check for custom styles via run style
            if hasattr(run, 'style') and run.style and hasattr(run.style, 'name'):
                style_name = run.style.name.lower().replace(' ', '-')
                # Map common Word styles to semantic HTML
                if 'code' in style_name or 'monospace' in style_name:
                    txt = f"<code>{txt}</code>"
                elif 'emphasis' in style_name or 'stress' in style_name:
                    txt = f"<em>{txt}</em>"
                elif 'strong' in style_name or 'intense' in style_name:
                    txt = f"<strong>{txt}</strong>"
                elif style_name not in ('normal', 'default'):
                    # Apply as CSS class for custom styles
                    txt = f'<span class="style-{style_name}">{txt}</span>'

        except AttributeError:
            # Fallback if style attributes don't exist
            pass

        return txt

    for element_type, element in document_elements:
        if element_type == 'table':
            # Process table
            flush_list()
            table_html = _process_table(element)
            buf.append(table_html)
            continue
        elif element_type == 'other':
            # Process other elements (text boxes, shapes, equations)
            flush_list()
            if 'textbox' in element.tag.lower() or 'shape' in element.tag.lower():
                other_html = _process_text_box_or_shape(element, tempdir)
            elif 'math' in element.tag.lower() or 'equation' in element.tag.lower():
                other_html = _process_equation(element)
            else:
                # Skip unknown elements or add generic handling
                continue
            buf.append(other_html)
            continue
        elif element_type != 'paragraph':
            continue

        # Process paragraph
        try:
            from docx.text.paragraph import Paragraph
            p = Paragraph(element, document)
        except Exception:
            continue

        style = (p.style.name or "").lower()
        # Detect list paragraphs
        is_num = False
        is_list = False
        try:
            pPr = p._p.pPr  # type: ignore[attr-defined]
            is_list = pPr is not None and pPr.numPr is not None  # type: ignore[attr-defined]
        except Exception:
            is_list = False
        if "list" in style or "bullet" in style or "number" in style:
            is_list = True
            is_num = "number" in style

        # Build paragraph content with basic runs formatting and images, merging hyperlinks
        run_html: list[str] = []
        current_link_href: str | None = None
        current_link_buf: list[str] = []
        for run in p.runs:
            # Process tracked changes
            tracked_text, should_include = _process_tracked_changes(run.element)
            if not should_include:
                continue

            # Process comments
            comment_html = _process_comments(run.element, document)

            txt = run.text or ""
            if tracked_text:
                txt = tracked_text
            # Images in this run?
            try:
                blips = run.element.xpath(".//a:blip", namespaces=IMG_NS)
            except Exception:
                blips = []
            for blip in blips:
                rid = blip.get(qn("r:embed"))
                if rid and rid in document.part.related_parts:
                    part = document.part.related_parts[rid]
                    filename = Path(part.partname).name
                    # Try to get alt text via docPr
                    alt = ""
                    try:
                        docprs = run.element.xpath(".//wp:docPr", namespaces=IMG_NS)
                        if docprs:
                            alt = docprs[0].get("descr") or docprs[0].get("title") or ""
                    except Exception:
                        alt = ""
                    if filename not in images:
                        out = tempdir / filename
                        out.write_bytes(part.blob)
                        images[filename] = out
                    alt_attr = f' alt="{alt}"' if alt else ' alt=""'
                    run_html.append(f'<img src="images/{filename}"{alt_attr} />')
            # Footnote/endnote references
            try:
                fns = run.element.xpath(
                    ".//w:footnoteReference",
                    namespaces={
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    },
                )
            except Exception:
                fns = []
            try:
                ens = run.element.xpath(
                    ".//w:endnoteReference",
                    namespaces={
                        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                    },
                )
            except Exception:
                ens = []
            # Manual page breaks
            try:
                brs = run.element.xpath('.//w:br[@w:type="page"]', namespaces=IMG_NS)
            except Exception:
                brs = []
            if brs:
                # Insert a pagebreak sentinel to be split later
                run_html.append("<!-- PAGEBREAK -->")
            for ref in fns + ens:
                note_idx += 1
                ref_id = ref.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id")
                sup = f'<sup id="fnref{note_idx}"><a href="#fn{note_idx}">{note_idx}</a></sup>'
                run_html.append(sup)
                # Retrieve note text if possible
                note_text = None
                try:
                    if (
                        "footnote" in ref.tag
                        and hasattr(document.part, "footnotes_part")
                        and document.part.footnotes_part is not None
                    ):
                        # type: ignore[attr-defined]
                        nodes = document.part.footnotes_part._element.xpath(
                            f'.//w:footnote[@w:id="{ref_id}"]//w:t',
                            namespaces={
                                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                            },
                        )
                        note_text = "".join(n.text or "" for n in nodes).strip()
                    elif (
                        "endnote" in ref.tag
                        and hasattr(document.part, "endnotes_part")
                        and document.part.endnotes_part is not None
                    ):
                        nodes = document.part.endnotes_part._element.xpath(
                            f'.//w:endnote[@w:id="{ref_id}"]//w:t',
                            namespaces={
                                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                            },
                        )
                        note_text = "".join(n.text or "" for n in nodes).strip()
                except Exception:
                    note_text = None
                if not note_text:
                    note_text = "(note)"
                current_notes.append(
                    f'<li id="fn{note_idx}"><p>{note_text} <a href="#fnref{note_idx}">↩</a></p></li>'
                )

            if txt:
                # Apply basic inline formatting and mapped run styles
                formatted_text = _get_run_html(run, run_styles_map, txt)
                # Add comment markers if present
                if comment_html:
                    formatted_text += comment_html
                run_html.append(formatted_text)
                # Hyperlinks: wrap if run is inside a hyperlink element
                try:
                    parent = run._r.getparent()
                    if parent is not None and parent.tag.endswith("}hyperlink"):
                        from docx.oxml.ns import qn  # type: ignore

                        rid = parent.get(qn("r:id")) if parent is not None else None
                        anchor = (
                            parent.get(
                                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}anchor"
                            )
                            if parent is not None
                            else None
                        )
                        href = None
                        if rid and rid in p.part.rels:
                            href = p.part.rels[rid].target_ref  # type: ignore[attr-defined]
                        elif anchor:
                            href = f"#{anchor}"
                        if href:
                            if current_link_href == href:
                                current_link_buf.append(txt)
                                txt = ""
                            else:
                                # flush previous
                                if current_link_href is not None:
                                    run_html.append(
                                        f'<a href="{current_link_href}">{" ".join(current_link_buf)}</a>'
                                    )
                                    current_link_buf = []
                                current_link_href = href
                                current_link_buf.append(txt)
                                txt = ""
                except Exception:
                    pass
                if txt:
                    # flush any open link first when standalone text encountered
                    if current_link_href is not None and current_link_buf:
                        run_html.append(
                            f'<a href="{current_link_href}">{" ".join(current_link_buf)}</a>'
                        )
                        current_link_href = None
                        current_link_buf = []
                    run_html.append(txt)

        # Flush open hyperlink group
        if current_link_href is not None and current_link_buf:
            run_html.append(f'<a href="{current_link_href}">{" ".join(current_link_buf)}</a>')

        content = "".join(run_html).strip()
        if not content and not pending_img:
            continue
        # If previous paragraph was image-only and this is a caption, wrap in figure
        if "caption" in style and pending_img:
            flush_list()
            buf.append(f"<figure>{pending_img}<figcaption>{content}</figcaption></figure>")
            pending_img = None
            continue

        # If content is an image-only paragraph, hold to see if next is a caption
        if re.fullmatch(r"\s*<img[^>]+>\s*", content):
            pending_img = content
            continue

        # If we have a pending image and the current paragraph is not a caption, flush it first
        if pending_img and ("caption" not in style):
            flush_list()
            buf.append(f"<p>{pending_img}</p>")
            pending_img = None

        # Apply paragraph style mapping
        mapped_tag = paragraph_styles_map.get(p.style.name, "p")

        # Parse tag and class attributes
        tag_parts = mapped_tag.split(' class="')
        base_tag = tag_parts[0]
        css_class = tag_parts[1].rstrip('"') if len(tag_parts) > 1 else None

        # Build the opening and closing tags
        if css_class:
            opening_tag = f'<{base_tag} class="{css_class}">'
            closing_tag = f'</{base_tag}>'
        else:
            opening_tag = f'<{base_tag}>'
            closing_tag = f'</{base_tag}>'

        if base_tag == "h1":
            flush_list()
            flush_section(buf, current_notes)
            current_notes = []
            buf.append(f"{opening_tag}{content}{closing_tag}")
        elif base_tag in ["h2", "h3", "h4", "h5", "h6"]:
            flush_list()
            buf.append(f"{opening_tag}{content}{closing_tag}")
        elif base_tag == "li" or mapped_tag == "li":
            list_type = "ol" if is_num else "ul"
            if current_list_type and current_list_type != list_type:
                flush_list()
            current_list_type = list_type if current_list_type is None else current_list_type
            list_items.append(f"<li>{content}</li>")
        elif base_tag == "blockquote":
            flush_list()
            buf.append(f"{opening_tag}<p>{content}</p>{closing_tag}")
        elif base_tag == "figcaption":
            # Special handling for captions
            if pending_img:
                flush_list()
                buf.append(f"<figure>{pending_img}{opening_tag}{content}{closing_tag}</figure>")
                pending_img = None
            else:
                flush_list()
                buf.append(f"{opening_tag}{content}{closing_tag}")
        elif base_tag == "pre":
            flush_list()
            buf.append(f"{opening_tag}{content}{closing_tag}")
        else: # Default to paragraph or specified tag
            flush_list()
            buf.append(f"{opening_tag}{content}{closing_tag}")

    flush_list()
    flush_section(buf, current_notes)
    if not parts:
        parts = ["<section><p>(Empty document)</p></section>"]

    # Extract CSS from styles
    styles_css = extract_styles_css(styles_data)

    # Return extracted images as resources and styles CSS
    return parts, list(images.values()), styles_css


# Legacy alias for backward compatibility
def docx_to_html_chunks(docx_path: Path) -> list[str]:
    """Legacy function name for docx_to_html - returns only HTML chunks."""
    chunks, _, _ = docx_to_html(docx_path)
    return chunks
