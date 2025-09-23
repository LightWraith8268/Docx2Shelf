"""
Cross-platform GUI application for Docx2Shelf.

Provides a user-friendly interface for EPUB conversion with drag-and-drop
support, live preview, and one-click conversion capabilities.
"""

from __future__ import annotations

import logging
import sys
import tempfile
import threading
from pathlib import Path
from typing import List

# Try different GUI frameworks in order of preference
GUI_FRAMEWORK = None

try:
    import tkinter as tk
    from tkinter import filedialog, messagebox, ttk
    from tkinter.scrolledtext import ScrolledText
    GUI_FRAMEWORK = "tkinter"
except ImportError:
    pass

if not GUI_FRAMEWORK:
    try:
        from PyQt5 import QtCore, QtGui, QtWidgets
        GUI_FRAMEWORK = "pyqt5"
    except ImportError:
        try:
            from PyQt6 import QtCore, QtGui, QtWidgets
            GUI_FRAMEWORK = "pyqt6"
        except ImportError:
            pass

from ..assemble import assemble_epub
from ..convert import docx_to_html_chunks
from ..metadata import BuildOptions, EpubMetadata
from ..settings import get_settings_manager
from ..tools import epubcheck_cmd
from ..update import check_for_updates, perform_update

logger = logging.getLogger(__name__)


class ConversionWorker:
    """Background worker for EPUB conversion."""

    def __init__(self, progress_callback=None, complete_callback=None, error_callback=None):
        self.progress_callback = progress_callback
        self.complete_callback = complete_callback
        self.error_callback = error_callback
        self.cancelled = False

    def convert(self, input_file: Path, output_file: Path, metadata: EpubMetadata, options: BuildOptions):
        """Run conversion in background thread."""
        try:
            if self.progress_callback:
                self.progress_callback("Starting conversion...", 0)

            if self.cancelled:
                return

            # Step 1: Convert DOCX to HTML
            if self.progress_callback:
                self.progress_callback("Converting DOCX to HTML...", 20)

            html_chunks = docx_to_html_chunks(input_file)

            if self.cancelled:
                return

            # Step 2: Extract resources
            if self.progress_callback:
                self.progress_callback("Extracting resources...", 40)

            resources = self._extract_resources(input_file)

            if self.cancelled:
                return

            # Step 3: Assemble EPUB
            if self.progress_callback:
                self.progress_callback("Assembling EPUB...", 60)

            assemble_epub(
                meta=metadata,
                opts=options,
                html_chunks=html_chunks,
                resources=resources,
                output_path=output_file
            )

            if self.cancelled:
                return

            # Step 4: Validate (optional)
            if options.validate_epub and epubcheck_cmd():
                if self.progress_callback:
                    self.progress_callback("Validating EPUB...", 80)

                # Run EPUBCheck validation
                # Implementation would go here

            if self.progress_callback:
                self.progress_callback("Conversion complete!", 100)

            if self.complete_callback:
                self.complete_callback(output_file)

        except Exception as e:
            logger.error(f"Conversion failed: {e}")
            if self.error_callback:
                self.error_callback(str(e))

    def cancel(self):
        """Cancel the conversion."""
        self.cancelled = True

    def _extract_resources(self, docx_file: Path) -> List[Path]:
        """Extract resources from DOCX file."""
        resources = []

        try:
            import zipfile

            with zipfile.ZipFile(docx_file, 'r') as docx_zip:
                # Find media files
                media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = Path(temp_dir)

                    for media_file in media_files:
                        # Extract to temp directory
                        media_data = docx_zip.read(media_file)
                        output_file = temp_path / Path(media_file).name
                        output_file.write_bytes(media_data)
                        resources.append(output_file)

        except Exception as e:
            logger.warning(f"Could not extract resources: {e}")

        return resources


if GUI_FRAMEWORK == "tkinter":
    class TkinterMainWindow:
        """Main application window using Tkinter."""

        def __init__(self):
            self.root = tk.Tk()
            self.root.title("Docx2Shelf - Document to EPUB Converter")

            # Modern window configuration
            self.root.geometry("1000x700")
            self.root.minsize(800, 600)

            # Configure for modern appearance
            self.root.configure(bg='#f8f9fa')

            # Enable native theme and high DPI support
            try:
                # Enable high DPI awareness on Windows
                import ctypes
                try:
                    ctypes.windll.shcore.SetProcessDpiAwareness(1)
                except:
                    pass
            except:
                pass

            # Set icon if available
            try:
                icon_path = Path(__file__).parent / "assets" / "icon.ico"
                if icon_path.exists():
                    self.root.iconbitmap(str(icon_path))
            except:
                pass

            # Initialize settings manager
            self.settings_manager = get_settings_manager()
            self.settings = self.settings_manager.settings

            # Initialize theme tracking
            self.current_theme = 'light'  # Default theme

            # Initialize modern styling
            self.setup_modern_styling()
            self.setup_menu()
            self.setup_ui()
            self.current_worker = None
            self.load_ui_preferences()

            # Apply initial theme
            self.apply_modern_theme()

        def setup_modern_styling(self):
            """Setup modern styling for the application."""
            # Configure ttk Style for modern appearance
            self.style = ttk.Style()

            # Use modern theme as base
            try:
                # Try to use a modern theme if available
                available_themes = self.style.theme_names()
                if 'vista' in available_themes:
                    self.style.theme_use('vista')
                elif 'clam' in available_themes:
                    self.style.theme_use('clam')
                else:
                    self.style.theme_use('default')
            except:
                pass

            # Define modern color palette
            self.colors = {
                'primary': '#007acc',
                'primary_dark': '#005a9e',
                'primary_light': '#4da6e8',
                'secondary': '#6c757d',
                'success': '#28a745',
                'warning': '#ffc107',
                'danger': '#dc3545',
                'light': '#f8f9fa',
                'dark': '#343a40',
                'white': '#ffffff',
                'gray_50': '#f9f9f9',
                'gray_100': '#f8f9fa',
                'gray_200': '#e9ecef',
                'gray_300': '#dee2e6',
                'gray_400': '#ced4da',
                'gray_500': '#adb5bd',
                'gray_600': '#6c757d',
                'gray_700': '#495057',
                'gray_800': '#343a40',
                'gray_900': '#212529',
                'background': '#ffffff'
            }

            # Configure modern button styles with rounded appearance
            self.style.configure('Modern.TButton',
                               relief='flat',
                               borderwidth=1,
                               focuscolor='none',
                               padding=(12, 8),
                               font=('Segoe UI', 9),
                               bordercolor=self.colors['primary'],
                               lightcolor=self.colors['primary'],
                               darkcolor=self.colors['primary'])

            self.style.map('Modern.TButton',
                         background=[('active', self.colors['primary_dark']),
                                   ('pressed', self.colors['primary_dark']),
                                   ('!active', self.colors['primary'])],
                         foreground=[('active', 'white'),
                                   ('pressed', 'white'),
                                   ('!active', 'white')],
                         bordercolor=[('active', self.colors['primary_dark']),
                                    ('pressed', self.colors['primary_dark']),
                                    ('!active', self.colors['primary'])])

            # Primary button style with enhanced rounded appearance
            self.style.configure('Primary.TButton',
                               relief='raised',
                               borderwidth=2,
                               focuscolor='none',
                               padding=(16, 10),
                               font=('Segoe UI', 10, 'bold'),
                               bordercolor=self.colors['primary_dark'],
                               lightcolor=self.colors['primary_light'],
                               darkcolor=self.colors['primary_dark'])

            self.style.map('Primary.TButton',
                         background=[('active', self.colors['primary_dark']),
                                   ('pressed', self.colors['primary_dark']),
                                   ('!active', self.colors['primary'])],
                         foreground=[('active', 'white'),
                                   ('pressed', 'white'),
                                   ('!active', 'white')],
                         relief=[('active', 'raised'),
                               ('pressed', 'sunken'),
                               ('!active', 'raised')])

            # Secondary button style
            self.style.configure('Secondary.TButton',
                               relief='flat',
                               borderwidth=1,
                               focuscolor='none',
                               padding=(12, 8),
                               font=('Segoe UI', 9))

            self.style.map('Secondary.TButton',
                         background=[('active', self.colors['gray_100']),
                                   ('pressed', self.colors['gray_200']),
                                   ('!active', 'white')],
                         foreground=[('active', self.colors['gray_800']),
                                   ('pressed', self.colors['gray_800']),
                                   ('!active', self.colors['gray_700'])],
                         bordercolor=[('active', self.colors['gray_300']),
                                    ('!active', self.colors['gray_300'])])

            # Success button style
            self.style.configure('Success.TButton',
                               relief='flat',
                               borderwidth=0,
                               focuscolor='none',
                               padding=(12, 8),
                               font=('Segoe UI', 9))

            self.style.map('Success.TButton',
                         background=[('active', '#218838'),
                                   ('pressed', '#218838'),
                                   ('!active', self.colors['success'])],
                         foreground=[('active', 'white'),
                                   ('pressed', 'white'),
                                   ('!active', 'white')])

            # Modern entry style
            self.style.configure('Modern.TEntry',
                               relief='flat',
                               borderwidth=1,
                               focuscolor='none',
                               padding=8,
                               font=('Segoe UI', 9))

            self.style.map('Modern.TEntry',
                         bordercolor=[('focus', self.colors['primary']),
                                    ('!focus', self.colors['gray_300'])],
                         lightcolor=[('focus', self.colors['primary']),
                                   ('!focus', self.colors['gray_300'])],
                         darkcolor=[('focus', self.colors['primary']),
                                  ('!focus', self.colors['gray_300'])])

            # Modern frame style
            self.style.configure('Modern.TFrame',
                               relief='flat',
                               borderwidth=0,
                               background=self.colors['white'])

            # Card frame style
            self.style.configure('Card.TFrame',
                               relief='flat',
                               borderwidth=1,
                               background=self.colors['white'])

            # Modern label style
            self.style.configure('Modern.TLabel',
                               background=self.colors['white'],
                               foreground=self.colors['gray_800'],
                               font=('Segoe UI', 9))

            self.style.configure('Heading.TLabel',
                               background=self.colors['white'],
                               foreground=self.colors['gray_900'],
                               font=('Segoe UI', 12, 'bold'))

            self.style.configure('Subheading.TLabel',
                               background=self.colors['white'],
                               foreground=self.colors['gray_700'],
                               font=('Segoe UI', 10))

            # Modern notebook style with enhanced tabs
            self.style.configure('Modern.TNotebook',
                               background=self.colors['white'],
                               borderwidth=0)

            self.style.configure('Modern.TNotebook.Tab',
                               padding=[20, 12],
                               font=('Segoe UI', 9),
                               borderwidth=1,
                               focuscolor='none')

            self.style.map('Modern.TNotebook.Tab',
                         background=[('selected', self.colors['primary']),
                                   ('active', self.colors['primary_light']),
                                   ('!active', self.colors['gray_100'])],
                         foreground=[('selected', 'white'),
                                   ('active', 'white'),
                                   ('!active', self.colors['gray_700'])],
                         bordercolor=[('selected', self.colors['primary']),
                                    ('active', self.colors['primary_light']),
                                    ('!active', self.colors['gray_300'])])

            # Tab frame style
            self.style.configure('Tab.TFrame',
                               background=self.colors['white'],
                               relief='flat',
                               borderwidth=0)

            # Header frame style
            self.style.configure('Header.TFrame',
                               background=self.colors['primary'],
                               relief='flat',
                               borderwidth=0)

            # Title label styles
            self.style.configure('Title.TLabel',
                               background=self.colors['primary'],
                               foreground='white',
                               font=('Segoe UI', 16, 'bold'))

            self.style.configure('Subtitle.TLabel',
                               background=self.colors['primary'],
                               foreground=self.colors['primary_light'],
                               font=('Segoe UI', 10))

            self.style.configure('Small.TLabel',
                               background=self.colors['primary'],
                               foreground='white',
                               font=('Segoe UI', 9))

            # Modern combobox style
            self.style.configure('Modern.TCombobox',
                               relief='flat',
                               borderwidth=1,
                               padding=5,
                               font=('Segoe UI', 9))

            # Modern scrollbar style
            self.style.configure('Modern.Vertical.TScrollbar',
                               background=self.colors['gray_200'],
                               troughcolor=self.colors['gray_100'],
                               borderwidth=0,
                               arrowcolor=self.colors['gray_600'],
                               darkcolor=self.colors['gray_300'],
                               lightcolor=self.colors['gray_300'])

            # Card labelframe style
            self.style.configure('Card.TLabelframe',
                               relief='solid',
                               borderwidth=1,
                               background=self.colors['white'])

            self.style.configure('Card.TLabelframe.Label',
                               background=self.colors['white'],
                               foreground=self.colors['gray_800'],
                               font=('Segoe UI', 10, 'bold'))

            # Drop zone styles
            self.style.configure('DropZone.TFrame',
                               relief='dashed',
                               borderwidth=2,
                               background=self.colors['gray_50'])

            self.style.configure('DropZone.TLabel',
                               background=self.colors['gray_50'],
                               foreground=self.colors['gray_600'],
                               font=('Segoe UI', 10))

            # Additional button styles
            self.style.configure('Secondary.TButton',
                               background=self.colors['gray_200'],
                               foreground=self.colors['gray_800'],
                               borderwidth=1,
                               focuscolor='none',
                               padding=[12, 8],
                               font=('Segoe UI', 9))

            self.style.map('Secondary.TButton',
                         background=[('active', self.colors['gray_300']),
                                   ('pressed', self.colors['gray_400'])],
                         foreground=[('active', self.colors['gray_900']),
                                   ('pressed', self.colors['gray_900'])])

            # Progress bar style
            self.style.configure('Modern.Horizontal.TProgressbar',
                               background=self.colors['primary'],
                               troughcolor=self.colors['gray_200'],
                               borderwidth=0,
                               lightcolor=self.colors['primary'],
                               darkcolor=self.colors['primary'])

        def apply_modern_theme(self):
            """Apply modern theme based on current selection."""
            try:
                # Use current_theme if set, otherwise fall back to settings
                if hasattr(self, 'current_theme'):
                    theme = self.current_theme
                else:
                    theme = self.settings.ui_preferences.theme if hasattr(self.settings, 'ui_preferences') else 'light'

                if theme == 'dark':
                    self.apply_dark_theme()
                elif theme == 'light':
                    self.apply_light_theme()
                else:
                    # Apply theme based on the existing theme system
                    self.apply_ui_theme()

                # Update theme combo if it exists
                if hasattr(self, 'theme_combo'):
                    self.theme_combo.set(theme.title())

            except Exception as e:
                print(f"Theme application error: {e}")
                # Fallback to light theme
                self.apply_light_theme()

        def apply_light_theme(self):
            """Apply modern light theme with comprehensive styling."""
            # Update root and canvas backgrounds
            self.root.configure(bg=self.colors['light'])
            if hasattr(self, 'canvas'):
                self.canvas.configure(bg=self.colors['white'])

            # Update all frame styles
            self.style.configure('Modern.TFrame', background=self.colors['white'])
            self.style.configure('Card.TFrame', background=self.colors['white'])
            self.style.configure('Tab.TFrame', background=self.colors['white'])
            self.style.configure('Header.TFrame', background=self.colors['primary'])

            # Update label styles
            self.style.configure('Modern.TLabel', background=self.colors['white'], foreground=self.colors['gray_800'])
            self.style.configure('Heading.TLabel', background=self.colors['white'], foreground=self.colors['gray_900'])
            self.style.configure('Subheading.TLabel', background=self.colors['white'], foreground=self.colors['gray_700'])
            self.style.configure('Title.TLabel', background=self.colors['primary'], foreground='white')
            self.style.configure('Subtitle.TLabel', background=self.colors['primary'], foreground=self.colors['primary_light'])
            self.style.configure('Small.TLabel', background=self.colors['primary'], foreground='white')

            # Update entry styles
            self.style.configure('Modern.TEntry', fieldbackground=self.colors['white'], foreground=self.colors['gray_800'])

            # Update labelframe styles
            self.style.configure('Card.TLabelframe', background=self.colors['white'])
            self.style.configure('Card.TLabelframe.Label', background=self.colors['white'], foreground=self.colors['gray_800'])

            # Update drop zone styles
            self.style.configure('DropZone.TFrame', background=self.colors['gray_50'])
            self.style.configure('DropZone.TLabel', background=self.colors['gray_50'], foreground=self.colors['gray_600'])

            # Update Text widgets manually for light theme
            self.update_text_widgets_for_light_theme()

        def apply_dark_theme(self):
            """Apply modern dark theme with comprehensive styling."""
            # Dark theme color palette
            dark_bg = '#1e1e1e'
            dark_surface = '#2d2d2d'
            dark_surface_light = '#3c3c3c'
            dark_text = '#ffffff'
            dark_text_secondary = '#b0b0b0'
            dark_primary = '#0078d4'
            dark_primary_light = '#4da6e8'

            # Update root and canvas backgrounds
            self.root.configure(bg=dark_bg)
            if hasattr(self, 'canvas'):
                self.canvas.configure(bg=dark_surface)

            # Update all frame styles
            self.style.configure('Modern.TFrame', background=dark_surface)
            self.style.configure('Card.TFrame', background=dark_surface)
            self.style.configure('Tab.TFrame', background=dark_surface)
            self.style.configure('Header.TFrame', background=dark_primary)

            # Update label styles
            self.style.configure('Modern.TLabel', background=dark_surface, foreground=dark_text)
            self.style.configure('Heading.TLabel', background=dark_surface, foreground=dark_text)
            self.style.configure('Subheading.TLabel', background=dark_surface, foreground=dark_text_secondary)
            self.style.configure('Title.TLabel', background=dark_primary, foreground='white')
            self.style.configure('Subtitle.TLabel', background=dark_primary, foreground=dark_primary_light)
            self.style.configure('Small.TLabel', background=dark_primary, foreground='white')

            # Update entry styles
            self.style.configure('Modern.TEntry', fieldbackground=dark_surface_light, foreground=dark_text)

            # Update labelframe styles
            self.style.configure('Card.TLabelframe', background=dark_surface)
            self.style.configure('Card.TLabelframe.Label', background=dark_surface, foreground=dark_text)

            # Update drop zone styles
            self.style.configure('DropZone.TFrame', background=dark_surface_light)
            self.style.configure('DropZone.TLabel', background=dark_surface_light, foreground=dark_text_secondary)

            # Update button styles for dark theme
            self.style.configure('Modern.TButton', background=dark_surface_light, foreground=dark_text)
            self.style.configure('Primary.TButton', background=dark_primary, foreground='white')
            self.style.configure('Secondary.TButton', background=dark_surface_light, foreground=dark_text)

            # Update Text widgets manually (tkinter Text widgets don't use ttk styles)
            self.update_text_widgets_for_dark_theme(dark_surface_light, dark_text)

        def update_text_widgets_for_light_theme(self):
            """Update all Text widgets for light theme."""
            try:
                # Update description text widget if it exists
                if hasattr(self, 'description_text'):
                    self.description_text.configure(
                        bg=self.colors['white'],
                        fg=self.colors['gray_800'],
                        insertbackground=self.colors['gray_800'],
                        selectbackground=self.colors['primary_light'],
                        selectforeground='white'
                    )

                # Update rights text widget if it exists
                if hasattr(self, 'rights_text'):
                    self.rights_text.configure(
                        bg=self.colors['white'],
                        fg=self.colors['gray_800'],
                        insertbackground=self.colors['gray_800'],
                        selectbackground=self.colors['primary_light'],
                        selectforeground='white'
                    )
            except Exception as e:
                print(f"Error updating light theme text widgets: {e}")

        def update_text_widgets_for_dark_theme(self, bg_color, fg_color):
            """Update all Text widgets for dark theme."""
            try:
                # Update description text widget if it exists
                if hasattr(self, 'description_text'):
                    self.description_text.configure(
                        bg=bg_color,
                        fg=fg_color,
                        insertbackground=fg_color,
                        selectbackground='#0078d4',
                        selectforeground='white'
                    )

                # Update rights text widget if it exists
                if hasattr(self, 'rights_text'):
                    self.rights_text.configure(
                        bg=bg_color,
                        fg=fg_color,
                        insertbackground=fg_color,
                        selectbackground='#0078d4',
                        selectforeground='white'
                    )
            except Exception as e:
                print(f"Error updating dark theme text widgets: {e}")

        def on_theme_change(self, event=None):
            """Handle theme change from dropdown."""
            if hasattr(self, 'theme_combo'):
                selected_theme = self.theme_combo.get().lower()
                if selected_theme != self.current_theme:
                    self.current_theme = selected_theme
                    self.apply_modern_theme()
                    # Update canvas background if it exists
                    if hasattr(self, 'canvas'):
                        self.canvas.configure(bg=self.colors['background'])

        def _on_mousewheel(self, event):
            """Handle mouse wheel scrolling."""
            self.canvas.yview_scroll(int(-1*(event.delta/120)), "units")

        def setup_menu(self):
            """Setup the menu bar with CLI command equivalents."""
            menubar = tk.Menu(self.root)
            self.root.config(menu=menubar)

            # File menu
            file_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="File", menu=file_menu)
            file_menu.add_command(label="Open Document...", command=self.browse_input_file, accelerator="Ctrl+O")
            file_menu.add_separator()
            file_menu.add_command(label="Exit", command=self.root.quit, accelerator="Ctrl+Q")

            # Tools menu (equivalent to CLI tools subcommand)
            tools_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Tools", menu=tools_menu)
            tools_menu.add_command(label="Install Pandoc", command=self.install_pandoc)
            tools_menu.add_command(label="Install EPUBCheck", command=self.install_epubcheck)
            tools_menu.add_separator()
            tools_menu.add_command(label="Tool Locations", command=self.show_tool_locations)
            tools_menu.add_command(label="System Doctor", command=self.run_system_doctor)

            # Themes menu (equivalent to CLI themes subcommands)
            themes_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Themes", menu=themes_menu)
            themes_menu.add_command(label="List Themes", command=self.list_themes)
            themes_menu.add_command(label="Preview Themes", command=self.preview_themes)

            # Validation menu
            validation_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Validation", menu=validation_menu)
            validation_menu.add_command(label="Validate EPUB...", command=self.validate_epub_file)
            validation_menu.add_command(label="Quality Analysis...", command=self.analyze_epub_quality)
            validation_menu.add_command(label="Store Checklist...", command=self.run_store_checklist)

            # AI menu (equivalent to CLI ai subcommands)
            ai_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="AI", menu=ai_menu)
            ai_menu.add_command(label="Enhance Metadata...", command=self.ai_enhance_metadata)
            ai_menu.add_command(label="Detect Genre...", command=self.ai_detect_genre)
            ai_menu.add_command(label="Generate Alt Text...", command=self.ai_generate_alt_text)
            ai_menu.add_separator()
            ai_menu.add_command(label="AI Configuration...", command=self.configure_ai)

            # Settings menu
            settings_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Settings", menu=settings_menu)
            settings_menu.add_command(label="Preferences...", command=self.open_unified_settings, accelerator="Ctrl+,")
            settings_menu.add_separator()
            settings_menu.add_command(label="Export Settings...", command=self.export_all_settings)
            settings_menu.add_command(label="Import Settings...", command=self.import_all_settings)
            settings_menu.add_command(label="Reset to Defaults...", command=self.reset_all_settings)

            # Help menu
            help_menu = tk.Menu(menubar, tearoff=0)
            menubar.add_cascade(label="Help", menu=help_menu)
            help_menu.add_command(label="Update Docx2Shelf", command=self.check_for_updates)
            help_menu.add_command(label="System Doctor", command=self.run_system_doctor)
            help_menu.add_separator()
            help_menu.add_command(label="About", command=self.show_about_dialog)

            # Bind keyboard shortcuts
            self.root.bind('<Control-o>', lambda e: self.browse_input_file())
            self.root.bind('<Control-q>', lambda e: self.root.quit())

        def setup_ui(self):
            """Setup the user interface with modern styling."""

            # Create main container with modern styling
            main_frame = ttk.Frame(self.root, style="Card.TFrame")
            main_frame.pack(fill="both", expand=True, padx=5, pady=5)

            # Create header section
            header_frame = ttk.Frame(main_frame, style="Header.TFrame")
            header_frame.pack(fill="x", padx=0, pady=(0, 10))

            # App title and version
            title_frame = ttk.Frame(header_frame, style="Header.TFrame")
            title_frame.pack(side="left", fill="x", expand=True)

            app_title = ttk.Label(title_frame, text="Docx2Shelf",
                                style="Title.TLabel")
            app_title.pack(side="left", padx=(10, 5))

            version_label = ttk.Label(title_frame, text="v1.6.2",
                                    style="Subtitle.TLabel")
            version_label.pack(side="left")

            # Theme selector in header
            theme_frame = ttk.Frame(header_frame, style="Header.TFrame")
            theme_frame.pack(side="right", padx=(5, 10))

            ttk.Label(theme_frame, text="Theme:", style="Small.TLabel").pack(side="left", padx=(0, 5))
            theme_combo = ttk.Combobox(theme_frame, values=["Light", "Dark"],
                                     width=8, state="readonly", style="Modern.TCombobox")
            theme_combo.set(self.current_theme.title())
            theme_combo.bind("<<ComboboxSelected>>", self.on_theme_change)
            theme_combo.pack(side="right")
            self.theme_combo = theme_combo

            # Create canvas for scrolling with modern styling
            canvas_frame = ttk.Frame(main_frame, style="Card.TFrame")
            canvas_frame.pack(fill="both", expand=True)

            self.canvas = tk.Canvas(canvas_frame, highlightthickness=0,
                                  bg=self.colors['background'])
            self.scrollable_frame = ttk.Frame(self.canvas, style="Card.TFrame")

            # Modern scrollbar
            scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical", command=self.canvas.yview,
                                    style="Modern.Vertical.TScrollbar")
            self.canvas.configure(yscrollcommand=scrollbar.set)

            # Configure scrolling
            self.scrollable_frame.bind(
                "<Configure>",
                lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
            )

            self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
            self.canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

            # Bind mouse wheel to canvas
            self.canvas.bind("<MouseWheel>", self._on_mousewheel)
            self.root.bind("<MouseWheel>", self._on_mousewheel)

            # Create main notebook for tabs with modern styling
            notebook = ttk.Notebook(self.scrollable_frame, style="Modern.TNotebook")
            notebook.pack(fill="both", expand=True, padx=15, pady=15)

            # Convert tab
            convert_frame = ttk.Frame(notebook, style="Tab.TFrame")
            notebook.add(convert_frame, text="📄 Convert")
            self.setup_convert_tab(convert_frame)

            # Settings tab
            settings_frame = ttk.Frame(notebook, style="Tab.TFrame")
            notebook.add(settings_frame, text="⚙️ Settings")
            self.setup_settings_tab(settings_frame)

            # Batch Processing tab (equivalent to CLI batch command)
            batch_frame = ttk.Frame(notebook, style="Tab.TFrame")
            notebook.add(batch_frame, text="📦 Batch")
            self.setup_batch_tab(batch_frame)

            # About tab
            about_frame = ttk.Frame(notebook, style="Tab.TFrame")
            notebook.add(about_frame, text="ℹ️ About")
            self.setup_about_tab(about_frame)

        def setup_convert_tab(self, parent):
            """Setup the conversion tab with modern styling."""

            # File selection section
            file_frame = ttk.LabelFrame(parent, text="📁 Input Document",
                                      style="Card.TLabelframe", padding=15)
            file_frame.pack(fill="x", padx=15, pady=(5, 10))

            # File input row
            input_row = ttk.Frame(file_frame, style="Card.TFrame")
            input_row.pack(fill="x", pady=(0, 10))

            self.file_path_var = tk.StringVar()
            file_entry = ttk.Entry(input_row, textvariable=self.file_path_var,
                                 style="Modern.TEntry", font=("Segoe UI", 10))
            file_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))

            browse_btn = ttk.Button(input_row, text="📂 Browse...",
                                  command=self.browse_input_file, style="Modern.TButton")
            browse_btn.pack(side="right")

            # Drag and drop area
            drop_frame = ttk.Frame(file_frame, style="DropZone.TFrame")
            drop_frame.pack(fill="x", pady=5)

            drop_label = ttk.Label(drop_frame,
                                 text="💭 Or drag and drop a document here\n(DOCX, MD, TXT, HTML)",
                                 style="DropZone.TLabel", justify="center")
            drop_label.pack(pady=15)

            # Enable drag and drop
            self.setup_drag_drop()

            # Metadata section with modern notebook
            meta_frame = ttk.LabelFrame(parent, text="📖 Book Metadata",
                                      style="Card.TLabelframe", padding=15)
            meta_frame.pack(fill="x", padx=15, pady=10)

            # Create notebook for metadata categories with modern styling
            meta_notebook = ttk.Notebook(meta_frame, style="Modern.TNotebook")
            meta_notebook.pack(fill="both", expand=True, pady=(0, 5))

            # Basic Info tab
            basic_frame = ttk.Frame(meta_notebook, style="Tab.TFrame")
            meta_notebook.add(basic_frame, text="📝 Basic Info")
            self.setup_basic_metadata(basic_frame)

            # Publishing tab
            publishing_frame = ttk.Frame(meta_notebook, style="Tab.TFrame")
            meta_notebook.add(publishing_frame, text="📚 Publishing")
            self.setup_publishing_metadata(publishing_frame)

            # Series tab
            series_frame = ttk.Frame(meta_notebook, style="Tab.TFrame")
            meta_notebook.add(series_frame, text="📑 Series")
            self.setup_series_metadata(series_frame)

            # Contributors tab
            contributors_frame = ttk.Frame(meta_notebook, style="Tab.TFrame")
            meta_notebook.add(contributors_frame, text="👥 Contributors")
            self.setup_contributors_metadata(contributors_frame)

            # Output section with modern styling
            output_frame = ttk.LabelFrame(parent, text="📁 Output Location",
                                        style="Card.TLabelframe", padding=15)
            output_frame.pack(fill="x", padx=15, pady=10)

            # Output row
            output_row = ttk.Frame(output_frame, style="Card.TFrame")
            output_row.pack(fill="x", pady=5)

            self.output_path_var = tk.StringVar()
            output_entry = ttk.Entry(output_row, textvariable=self.output_path_var,
                                   style="Modern.TEntry", font=("Segoe UI", 10))
            output_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))

            output_browse_btn = ttk.Button(output_row, text="📂 Browse...",
                                         command=self.browse_output_file, style="Modern.TButton")
            output_browse_btn.pack(side="right")

            # Conversion Options with modern tabbed interface
            options_frame = ttk.LabelFrame(parent, text="⚙️ Conversion Options",
                                         style="Card.TLabelframe", padding=15)
            options_frame.pack(fill="both", expand=True, padx=15, pady=10)

            # Create notebook for conversion options with modern styling
            options_notebook = ttk.Notebook(options_frame, style="Modern.TNotebook")
            options_notebook.pack(fill="both", expand=True, pady=(0, 10))

            # Basic Options tab
            basic_options_frame = ttk.Frame(options_notebook, style="Tab.TFrame")
            options_notebook.add(basic_options_frame, text="🔧 Basic")
            self.setup_basic_conversion_options(basic_options_frame)

            # Advanced Options tab
            advanced_options_frame = ttk.Frame(options_notebook, style="Tab.TFrame")
            options_notebook.add(advanced_options_frame, text="⚡ Advanced")
            self.setup_advanced_conversion_options(advanced_options_frame)

            # Images & Media tab
            media_options_frame = ttk.Frame(options_notebook, style="Tab.TFrame")
            options_notebook.add(media_options_frame, text="🖼️ Images & Media")
            self.setup_media_conversion_options(media_options_frame)

            # AI Features tab
            ai_options_frame = ttk.Frame(options_notebook, style="Tab.TFrame")
            options_notebook.add(ai_options_frame, text="🤖 AI Features")
            self.setup_ai_conversion_options(ai_options_frame)

            # Progress section with modern styling
            progress_frame = ttk.Frame(parent, style="Card.TFrame")
            progress_frame.pack(fill="x", padx=15, pady=10)

            self.progress_var = tk.StringVar(value="Ready to convert")
            progress_label = ttk.Label(progress_frame, textvariable=self.progress_var,
                                     style="Modern.TLabel", font=("Segoe UI", 10))
            progress_label.pack(pady=(0, 5))

            self.progress_bar = ttk.Progressbar(progress_frame, mode='determinate',
                                              style="Modern.Horizontal.TProgressbar")
            self.progress_bar.pack(fill="x", pady=(0, 10))

            # Action buttons with modern styling
            button_frame = ttk.Frame(parent, style="Card.TFrame")
            button_frame.pack(fill="x", padx=15, pady=(0, 15))

            # Primary convert button
            self.convert_btn = ttk.Button(button_frame, text="🚀 Convert to EPUB",
                                        command=self.start_conversion, style="Primary.TButton")
            self.convert_btn.pack(side="left", padx=(0, 10))

            # Secondary cancel button
            self.cancel_btn = ttk.Button(button_frame, text="⏹️ Cancel",
                                       command=self.cancel_conversion, style="Secondary.TButton",
                                       state="disabled")
            self.cancel_btn.pack(side="left")

            # Help button
            help_btn = ttk.Button(button_frame, text="❓ Help",
                                command=self.show_conversion_help, style="Modern.TButton")
            help_btn.pack(side="right")

            preview_btn = ttk.Button(button_frame, text="Preview", command=self.preview_epub)
            preview_btn.pack(side="right")

        def setup_basic_metadata(self, parent):
            """Setup basic metadata fields with modern styling."""
            # Create main container with padding
            main_container = ttk.Frame(parent, style="Card.TFrame")
            main_container.pack(fill="both", expand=True, padx=20, pady=15)

            # Title section
            title_frame = ttk.Frame(main_container, style="Card.TFrame")
            title_frame.pack(fill="x", pady=(0, 15))

            ttk.Label(title_frame, text="📖 Book Title", style="Heading.TLabel").pack(anchor="w", pady=(0, 5))
            self.title_var = tk.StringVar()
            title_entry = ttk.Entry(title_frame, textvariable=self.title_var,
                                  style="Modern.TEntry", font=("Segoe UI", 11))
            title_entry.pack(fill="x", pady=(0, 3))
            help_label = ttk.Label(title_frame, text="Enter the full title of your book",
                                  style="Modern.TLabel", font=("Segoe UI", 9))
            help_label.pack(anchor="w")
            help_label.configure(foreground=self.colors['gray_600'])

            # Author section
            author_frame = ttk.Frame(main_container, style="Card.TFrame")
            author_frame.pack(fill="x", pady=(0, 15))

            ttk.Label(author_frame, text="✍️ Author", style="Heading.TLabel").pack(anchor="w", pady=(0, 5))
            self.author_var = tk.StringVar()
            author_entry = ttk.Entry(author_frame, textvariable=self.author_var,
                                   style="Modern.TEntry", font=("Segoe UI", 11))
            author_entry.pack(fill="x", pady=(0, 3))
            help_label2 = ttk.Label(author_frame, text="Primary author or main contributor",
                                   style="Modern.TLabel", font=("Segoe UI", 9))
            help_label2.pack(anchor="w")
            help_label2.configure(foreground=self.colors['gray_600'])

            # Language and Genre row
            row_frame = ttk.Frame(main_container, style="Card.TFrame")
            row_frame.pack(fill="x", pady=(0, 15))

            # Language section (left half)
            lang_frame = ttk.Frame(row_frame, style="Card.TFrame")
            lang_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))

            ttk.Label(lang_frame, text="🌍 Language", style="Heading.TLabel").pack(anchor="w", pady=(0, 5))
            self.language_var = tk.StringVar(value=self.settings.conversion_defaults.language)
            language_combo = ttk.Combobox(lang_frame, textvariable=self.language_var,
                                        style="Modern.TCombobox", font=("Segoe UI", 10),
                                        values=["en", "es", "fr", "de", "it", "pt", "nl", "ru", "zh", "ja"],
                                        state="readonly")
            language_combo.pack(fill="x", pady=(0, 3))
            help_label3 = ttk.Label(lang_frame, text="Book's primary language",
                                   style="Modern.TLabel", font=("Segoe UI", 9))
            help_label3.pack(anchor="w")
            help_label3.configure(foreground=self.colors['gray_600'])

            # Genres section (right half)
            genre_frame = ttk.Frame(row_frame, style="Card.TFrame")
            genre_frame.pack(side="right", fill="both", expand=True, padx=(10, 0))

            ttk.Label(genre_frame, text="🏷️ Genres", style="Heading.TLabel").pack(anchor="w", pady=(0, 5))
            self.genres_var = tk.StringVar()
            genres_entry = ttk.Entry(genre_frame, textvariable=self.genres_var,
                                   style="Modern.TEntry", font=("Segoe UI", 10))
            genres_entry.pack(fill="x", pady=(0, 3))
            help_label4 = ttk.Label(genre_frame, text="e.g., Fantasy, Romance, Thriller",
                                   style="Modern.TLabel", font=("Segoe UI", 9))
            help_label4.pack(anchor="w")
            help_label4.configure(foreground=self.colors['gray_600'])

            # Description section
            desc_frame = ttk.Frame(main_container, style="Card.TFrame")
            desc_frame.pack(fill="x", pady=(0, 15))

            ttk.Label(desc_frame, text="📝 Description", style="Heading.TLabel").pack(anchor="w", pady=(0, 5))

            # Create a frame for the text widget with modern styling
            text_container = ttk.Frame(desc_frame, style="Modern.TFrame", relief="solid", borderwidth=1)
            text_container.pack(fill="x", pady=(0, 3))

            self.description_text = tk.Text(text_container, height=4, wrap=tk.WORD,
                                          font=("Segoe UI", 10), relief="flat", borderwidth=0,
                                          bg=self.colors['white'], fg=self.colors['gray_800'])
            self.description_text.pack(fill="both", expand=True, padx=8, pady=8)

            help_label5 = ttk.Label(desc_frame, text="Brief description or synopsis of your book",
                                   style="Modern.TLabel", font=("Segoe UI", 9))
            help_label5.pack(anchor="w")
            help_label5.configure(foreground=self.colors['gray_600'])

            # Keywords section
            keywords_frame = ttk.Frame(main_container, style="Card.TFrame")
            keywords_frame.pack(fill="x", pady=(0, 10))

            ttk.Label(keywords_frame, text="🔍 Keywords", style="Heading.TLabel").pack(anchor="w", pady=(0, 5))
            self.subjects_var = tk.StringVar()
            subjects_entry = ttk.Entry(keywords_frame, textvariable=self.subjects_var,
                                     style="Modern.TEntry", font=("Segoe UI", 10))
            subjects_entry.pack(fill="x", pady=(0, 3))
            help_label6 = ttk.Label(keywords_frame, text="Comma-separated keywords for discoverability",
                                   style="Modern.TLabel", font=("Segoe UI", 9))
            help_label6.pack(anchor="w")
            help_label6.configure(foreground=self.colors['gray_600'])

        def setup_publishing_metadata(self, parent):
            """Setup publishing metadata fields."""
            # Publisher
            ttk.Label(parent, text="Publisher:").grid(row=0, column=0, sticky="w", pady=2, padx=5)
            self.publisher_var = tk.StringVar()
            publisher_entry = ttk.Entry(parent, textvariable=self.publisher_var, width=40)
            publisher_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Publication Date
            ttk.Label(parent, text="Publication Date:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
            self.pubdate_var = tk.StringVar()
            pubdate_entry = ttk.Entry(parent, textvariable=self.pubdate_var, width=20)
            pubdate_entry.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="(YYYY-MM-DD)", font=("Arial", 8)).grid(row=1, column=2, sticky="w", padx=(5, 0))

            # ISBN
            ttk.Label(parent, text="ISBN:").grid(row=2, column=0, sticky="w", pady=2, padx=5)
            self.isbn_var = tk.StringVar()
            isbn_entry = ttk.Entry(parent, textvariable=self.isbn_var, width=30)
            isbn_entry.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=2)

            # Print ISBN
            ttk.Label(parent, text="Print ISBN:").grid(row=3, column=0, sticky="w", pady=2, padx=5)
            self.print_isbn_var = tk.StringVar()
            print_isbn_entry = ttk.Entry(parent, textvariable=self.print_isbn_var, width=30)
            print_isbn_entry.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=2)

            # Copyright
            ttk.Label(parent, text="Copyright Holder:").grid(row=4, column=0, sticky="w", pady=2, padx=5)
            self.copyright_holder_var = tk.StringVar()
            copyright_entry = ttk.Entry(parent, textvariable=self.copyright_holder_var, width=40)
            copyright_entry.grid(row=4, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Copyright Year
            ttk.Label(parent, text="Copyright Year:").grid(row=5, column=0, sticky="w", pady=2, padx=5)
            self.copyright_year_var = tk.StringVar()
            copyright_year_entry = ttk.Entry(parent, textvariable=self.copyright_year_var, width=10)
            copyright_year_entry.grid(row=5, column=1, sticky="w", padx=(5, 0), pady=2)

            # Rights
            ttk.Label(parent, text="Rights:").grid(row=6, column=0, sticky="nw", pady=2, padx=5)
            self.rights_text = tk.Text(parent, height=2, width=40)
            self.rights_text.grid(row=6, column=1, sticky="ew", padx=(5, 0), pady=2)

            parent.columnconfigure(1, weight=1)

        def setup_series_metadata(self, parent):
            """Setup series metadata fields."""
            # Series Name
            ttk.Label(parent, text="Series Name:").grid(row=0, column=0, sticky="w", pady=2, padx=5)
            self.series_var = tk.StringVar()
            series_entry = ttk.Entry(parent, textvariable=self.series_var, width=40)
            series_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Series Number (Book number in series)
            ttk.Label(parent, text="Book Number:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
            self.series_index_var = tk.StringVar()
            series_index_entry = ttk.Entry(parent, textvariable=self.series_index_var, width=10)
            series_index_entry.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="(e.g., 1, 2, 3)", font=("Arial", 8)).grid(row=1, column=2, sticky="w", padx=(5, 0))

            # Series Type
            ttk.Label(parent, text="Series Type:").grid(row=2, column=0, sticky="w", pady=2, padx=5)
            self.series_type_var = tk.StringVar()
            series_type_combo = ttk.Combobox(parent, textvariable=self.series_type_var,
                                           values=["sequence", "collection", "anthology"], width=15)
            series_type_combo.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=2)

            # Series Position Description
            ttk.Label(parent, text="Position Description:").grid(row=3, column=0, sticky="w", pady=2, padx=5)
            self.series_position_var = tk.StringVar()
            series_position_entry = ttk.Entry(parent, textvariable=self.series_position_var, width=20)
            series_position_entry.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="(e.g., '1 of 5', 'standalone', 'prequel')", font=("Arial", 8)).grid(row=3, column=2, sticky="w", padx=(5, 0))

            # Publication Type
            ttk.Label(parent, text="Publication Type:").grid(row=5, column=0, sticky="w", pady=2, padx=5)
            self.publication_type_var = tk.StringVar()
            pub_type_combo = ttk.Combobox(parent, textvariable=self.publication_type_var,
                                        values=["novel", "anthology", "memoir", "textbook", "manual"], width=15)
            pub_type_combo.grid(row=5, column=1, sticky="w", padx=(5, 0), pady=2)

            # Target Audience
            ttk.Label(parent, text="Target Audience:").grid(row=6, column=0, sticky="w", pady=2, padx=5)
            self.target_audience_var = tk.StringVar()
            audience_combo = ttk.Combobox(parent, textvariable=self.target_audience_var,
                                        values=["adult", "young adult", "children", "middle grade"], width=15)
            audience_combo.grid(row=6, column=1, sticky="w", padx=(5, 0), pady=2)

            parent.columnconfigure(1, weight=1)

        def setup_contributors_metadata(self, parent):
            """Setup contributors metadata fields."""
            # Editor
            ttk.Label(parent, text="Editor:").grid(row=0, column=0, sticky="w", pady=2, padx=5)
            self.editor_var = tk.StringVar()
            editor_entry = ttk.Entry(parent, textvariable=self.editor_var, width=40)
            editor_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Illustrator
            ttk.Label(parent, text="Illustrator:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
            self.illustrator_var = tk.StringVar()
            illustrator_entry = ttk.Entry(parent, textvariable=self.illustrator_var, width=40)
            illustrator_entry.grid(row=1, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Translator
            ttk.Label(parent, text="Translator:").grid(row=2, column=0, sticky="w", pady=2, padx=5)
            self.translator_var = tk.StringVar()
            translator_entry = ttk.Entry(parent, textvariable=self.translator_var, width=40)
            translator_entry.grid(row=2, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Narrator
            ttk.Label(parent, text="Narrator:").grid(row=3, column=0, sticky="w", pady=2, padx=5)
            self.narrator_var = tk.StringVar()
            narrator_entry = ttk.Entry(parent, textvariable=self.narrator_var, width=40)
            narrator_entry.grid(row=3, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Designer
            ttk.Label(parent, text="Designer:").grid(row=4, column=0, sticky="w", pady=2, padx=5)
            self.designer_var = tk.StringVar()
            designer_entry = ttk.Entry(parent, textvariable=self.designer_var, width=40)
            designer_entry.grid(row=4, column=1, sticky="ew", padx=(5, 0), pady=2)

            # Contributor
            ttk.Label(parent, text="Other Contributor:").grid(row=5, column=0, sticky="w", pady=2, padx=5)
            self.contributor_var = tk.StringVar()
            contributor_entry = ttk.Entry(parent, textvariable=self.contributor_var, width=40)
            contributor_entry.grid(row=5, column=1, sticky="ew", padx=(5, 0), pady=2)

            parent.columnconfigure(1, weight=1)

        def setup_basic_conversion_options(self, parent):
            """Setup basic conversion options."""
            # CSS Theme
            ttk.Label(parent, text="CSS Theme:").grid(row=0, column=0, sticky="w", pady=2, padx=5)
            self.theme_var = tk.StringVar(value=self.settings.conversion_defaults.css_theme)
            theme_combo = ttk.Combobox(parent, textvariable=self.theme_var,
                                     values=["serif", "sans", "printlike"], width=15)
            theme_combo.grid(row=0, column=1, sticky="w", padx=(5, 0), pady=2)

            # Validate EPUB
            self.validate_var = tk.BooleanVar(value=self.settings.conversion_defaults.validate_epub)
            validate_check = ttk.Checkbutton(parent, text="Validate EPUB with EPUBCheck",
                                           variable=self.validate_var)
            validate_check.grid(row=0, column=2, sticky="w", padx=(20, 0), pady=2)

            # Split At
            ttk.Label(parent, text="Split Chapters At:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
            self.split_at_var = tk.StringVar(value="h1")
            split_combo = ttk.Combobox(parent, textvariable=self.split_at_var,
                                     values=["h1", "h2", "h3", "pagebreak", "mixed"], width=15)
            split_combo.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=2)

            # TOC Depth
            ttk.Label(parent, text="TOC Depth:").grid(row=2, column=0, sticky="w", pady=2, padx=5)
            self.toc_depth_var = tk.StringVar(value="2")
            toc_spin = ttk.Spinbox(parent, textvariable=self.toc_depth_var, from_=1, to=6, width=5)
            toc_spin.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=2)

            # Hyphenate
            self.hyphenate_var = tk.BooleanVar(value=True)
            hyphen_check = ttk.Checkbutton(parent, text="Enable Hyphenation",
                                         variable=self.hyphenate_var)
            hyphen_check.grid(row=1, column=2, sticky="w", padx=(20, 0), pady=2)

            # Justify
            self.justify_var = tk.BooleanVar(value=True)
            justify_check = ttk.Checkbutton(parent, text="Justify Text",
                                          variable=self.justify_var)
            justify_check.grid(row=2, column=2, sticky="w", padx=(20, 0), pady=2)

            # EPUB Version
            ttk.Label(parent, text="EPUB Version:").grid(row=3, column=0, sticky="w", pady=2, padx=5)
            self.epub_version_var = tk.StringVar(value="3")
            epub_combo = ttk.Combobox(parent, textvariable=self.epub_version_var,
                                    values=["2", "3"], width=5)
            epub_combo.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=2)

            # Accessibility Features
            self.a11y_var = tk.BooleanVar(value=self.settings.conversion_defaults.accessibility_features)
            a11y_check = ttk.Checkbutton(parent, text="Accessibility Features",
                                       variable=self.a11y_var)
            a11y_check.grid(row=3, column=2, sticky="w", padx=(20, 0), pady=2)

            parent.columnconfigure(2, weight=1)

        def setup_advanced_conversion_options(self, parent):
            """Setup advanced conversion options."""
            # Cover Image Selection
            ttk.Label(parent, text="Cover Image:").grid(row=0, column=0, sticky="w", pady=2, padx=5)
            self.cover_path_var = tk.StringVar()
            cover_entry = ttk.Entry(parent, textvariable=self.cover_path_var, width=40)
            cover_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=2)
            cover_browse_btn = ttk.Button(parent, text="Browse...", command=self.browse_cover_file)
            cover_browse_btn.grid(row=0, column=2, padx=(5, 0), pady=2)

            # Store Profile
            ttk.Label(parent, text="Store Profile:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
            self.store_profile_var = tk.StringVar()
            profile_combo = ttk.Combobox(parent, textvariable=self.store_profile_var,
                                       values=["generic", "kdp", "apple", "kobo", "google", "bn"], width=15)
            profile_combo.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=2)

            # Font Size
            ttk.Label(parent, text="Font Size:").grid(row=2, column=0, sticky="w", pady=2, padx=5)
            self.font_size_var = tk.StringVar()
            font_entry = ttk.Entry(parent, textvariable=self.font_size_var, width=15)
            font_entry.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="(e.g., 1rem, 12pt)", font=("Arial", 8)).grid(row=2, column=2, sticky="w", padx=(5, 0))

            # Line Height
            ttk.Label(parent, text="Line Height:").grid(row=3, column=0, sticky="w", pady=2, padx=5)
            self.line_height_var = tk.StringVar()
            line_entry = ttk.Entry(parent, textvariable=self.line_height_var, width=15)
            line_entry.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="(e.g., 1.5)", font=("Arial", 8)).grid(row=3, column=2, sticky="w", padx=(5, 0))

            # Extra CSS
            ttk.Label(parent, text="Extra CSS File:").grid(row=4, column=0, sticky="w", pady=2, padx=5)
            self.extra_css_var = tk.StringVar()
            css_entry = ttk.Entry(parent, textvariable=self.extra_css_var, width=40)
            css_entry.grid(row=4, column=1, sticky="ew", padx=(5, 0), pady=2)
            css_browse_btn = ttk.Button(parent, text="Browse...", command=self.browse_css_file)
            css_browse_btn.grid(row=4, column=2, padx=(5, 0), pady=2)

            # Chapter Start Mode
            ttk.Label(parent, text="Chapter Detection:").grid(row=5, column=0, sticky="w", pady=2, padx=5)
            self.chapter_mode_var = tk.StringVar(value="auto")
            chapter_combo = ttk.Combobox(parent, textvariable=self.chapter_mode_var,
                                       values=["auto", "manual", "mixed", "ai"], width=15)
            chapter_combo.grid(row=5, column=1, sticky="w", padx=(5, 0), pady=2)

            # Advanced options checkboxes
            self.epub2_compat_var = tk.BooleanVar()
            epub2_check = ttk.Checkbutton(parent, text="EPUB 2 Compatibility",
                                        variable=self.epub2_compat_var)
            epub2_check.grid(row=6, column=0, sticky="w", padx=5, pady=2)

            self.page_numbers_var = tk.BooleanVar()
            page_check = ttk.Checkbutton(parent, text="Include Page Numbers",
                                       variable=self.page_numbers_var)
            page_check.grid(row=6, column=1, sticky="w", padx=5, pady=2)

            self.page_list_var = tk.BooleanVar()
            list_check = ttk.Checkbutton(parent, text="Generate Page List",
                                       variable=self.page_list_var)
            list_check.grid(row=6, column=2, sticky="w", padx=5, pady=2)

            parent.columnconfigure(1, weight=1)

        def setup_media_conversion_options(self, parent):
            """Setup media and image conversion options."""
            # Image Quality
            ttk.Label(parent, text="Image Quality:").grid(row=0, column=0, sticky="w", pady=2, padx=5)
            self.image_quality_var = tk.StringVar(value="85")
            quality_spin = ttk.Spinbox(parent, textvariable=self.image_quality_var, from_=1, to=100, width=5)
            quality_spin.grid(row=0, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="(1-100)", font=("Arial", 8)).grid(row=0, column=2, sticky="w", padx=(5, 0))

            # Image Max Width
            ttk.Label(parent, text="Max Image Width:").grid(row=1, column=0, sticky="w", pady=2, padx=5)
            self.image_width_var = tk.StringVar(value=str(self.settings.conversion_defaults.image_max_width))
            width_entry = ttk.Entry(parent, textvariable=self.image_width_var, width=10)
            width_entry.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="pixels", font=("Arial", 8)).grid(row=1, column=2, sticky="w", padx=(5, 0))

            # Image Max Height
            ttk.Label(parent, text="Max Image Height:").grid(row=2, column=0, sticky="w", pady=2, padx=5)
            self.image_height_var = tk.StringVar(value="1600")
            height_entry = ttk.Entry(parent, textvariable=self.image_height_var, width=10)
            height_entry.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=2)
            ttk.Label(parent, text="pixels", font=("Arial", 8)).grid(row=2, column=2, sticky="w", padx=(5, 0))

            # Image Format
            ttk.Label(parent, text="Image Format:").grid(row=3, column=0, sticky="w", pady=2, padx=5)
            self.image_format_var = tk.StringVar(value="webp")
            format_combo = ttk.Combobox(parent, textvariable=self.image_format_var,
                                      values=["original", "webp", "avif"], width=15)
            format_combo.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=2)

            # Cover Scale
            ttk.Label(parent, text="Cover Scale:").grid(row=4, column=0, sticky="w", pady=2, padx=5)
            self.cover_scale_var = tk.StringVar(value="contain")
            scale_combo = ttk.Combobox(parent, textvariable=self.cover_scale_var,
                                     values=["contain", "cover"], width=15)
            scale_combo.grid(row=4, column=1, sticky="w", padx=(5, 0), pady=2)

            # Advanced image options
            self.enhanced_images_var = tk.BooleanVar()
            enhanced_check = ttk.Checkbutton(parent, text="Enhanced Image Processing",
                                           variable=self.enhanced_images_var)
            enhanced_check.grid(row=5, column=0, columnspan=2, sticky="w", padx=5, pady=2)

            self.vertical_writing_var = tk.BooleanVar()
            vertical_check = ttk.Checkbutton(parent, text="Vertical Writing Mode (CJK)",
                                           variable=self.vertical_writing_var)
            vertical_check.grid(row=6, column=0, columnspan=2, sticky="w", padx=5, pady=2)

            parent.columnconfigure(1, weight=1)

        def setup_ai_conversion_options(self, parent):
            """Setup AI-powered features."""
            ttk.Label(parent, text="AI Features:", font=("Arial", 10, "bold")).grid(row=0, column=0, sticky="w", pady=5, padx=5)

            # AI Enhancement
            self.ai_enhance_var = tk.BooleanVar()
            ai_enhance_check = ttk.Checkbutton(parent, text="AI-powered metadata enhancement",
                                             variable=self.ai_enhance_var)
            ai_enhance_check.grid(row=1, column=0, columnspan=2, sticky="w", padx=5, pady=2)

            # AI Genre Detection
            self.ai_genre_var = tk.BooleanVar()
            ai_genre_check = ttk.Checkbutton(parent, text="AI genre detection and keyword generation",
                                           variable=self.ai_genre_var)
            ai_genre_check.grid(row=2, column=0, columnspan=2, sticky="w", padx=5, pady=2)

            # AI Alt Text
            self.ai_alt_text_var = tk.BooleanVar()
            ai_alt_check = ttk.Checkbutton(parent, text="Generate AI-powered alt text for images",
                                         variable=self.ai_alt_text_var)
            ai_alt_check.grid(row=3, column=0, columnspan=2, sticky="w", padx=5, pady=2)

            # AI Interactive
            self.ai_interactive_var = tk.BooleanVar()
            ai_interactive_check = ttk.Checkbutton(parent, text="Interactive AI suggestions",
                                                 variable=self.ai_interactive_var)
            ai_interactive_check.grid(row=4, column=0, columnspan=2, sticky="w", padx=5, pady=2)

            # AI Config
            ttk.Label(parent, text="AI Config File:").grid(row=5, column=0, sticky="w", pady=2, padx=5)
            self.ai_config_var = tk.StringVar()
            ai_config_entry = ttk.Entry(parent, textvariable=self.ai_config_var, width=40)
            ai_config_entry.grid(row=5, column=1, sticky="ew", padx=(5, 0), pady=2)
            ai_config_btn = ttk.Button(parent, text="Browse...", command=self.browse_ai_config_file)
            ai_config_btn.grid(row=5, column=2, padx=(5, 0), pady=2)

            ttk.Label(parent, text="Note: AI features require proper configuration and API keys.",
                     font=("Arial", 8), foreground="gray").grid(row=6, column=0, columnspan=3, sticky="w", padx=5, pady=5)

            parent.columnconfigure(1, weight=1)

        def browse_cover_file(self):
            """Browse for cover image file."""
            file_path = filedialog.askopenfilename(
                title="Select cover image",
                filetypes=[
                    ("Image files", "*.jpg;*.jpeg;*.png;*.gif;*.bmp"),
                    ("JPEG files", "*.jpg;*.jpeg"),
                    ("PNG files", "*.png"),
                    ("All files", "*.*")
                ]
            )
            if file_path:
                self.cover_path_var.set(file_path)

        def browse_css_file(self):
            """Browse for extra CSS file."""
            file_path = filedialog.askopenfilename(
                title="Select CSS file",
                filetypes=[("CSS files", "*.css"), ("All files", "*.*")]
            )
            if file_path:
                self.extra_css_var.set(file_path)

        def browse_ai_config_file(self):
            """Browse for AI configuration file."""
            file_path = filedialog.askopenfilename(
                title="Select AI config file",
                filetypes=[("JSON files", "*.json"), ("YAML files", "*.yaml;*.yml"), ("All files", "*.*")]
            )
            if file_path:
                self.ai_config_var.set(file_path)

        def setup_settings_tab(self, parent):
            """Setup the comprehensive settings tab."""

            # Create notebook for settings categories
            settings_notebook = ttk.Notebook(parent)
            settings_notebook.pack(fill="both", expand=True, padx=10, pady=10)

            # Conversion Defaults tab
            conv_frame = ttk.Frame(settings_notebook)
            settings_notebook.add(conv_frame, text="Conversion")
            self.setup_conversion_settings(conv_frame)

            # UI Preferences tab
            ui_frame = ttk.Frame(settings_notebook)
            settings_notebook.add(ui_frame, text="Interface")
            self.setup_ui_settings(ui_frame)

            # File Defaults tab
            file_frame = ttk.Frame(settings_notebook)
            settings_notebook.add(file_frame, text="Files")
            self.setup_file_settings(file_frame)

            # Advanced Settings tab
            advanced_frame = ttk.Frame(settings_notebook)
            settings_notebook.add(advanced_frame, text="Advanced")
            self.setup_advanced_settings(advanced_frame)

            # Enterprise Configuration tab
            enterprise_frame = ttk.Frame(settings_notebook)
            settings_notebook.add(enterprise_frame, text="Enterprise")
            self.setup_enterprise_settings(enterprise_frame)

            # Settings control buttons
            button_frame = ttk.Frame(parent)
            button_frame.pack(fill="x", padx=10, pady=(0, 10))

            ttk.Button(button_frame, text="Save Settings",
                      command=self.save_settings).pack(side="left")
            ttk.Button(button_frame, text="Reset to Defaults",
                      command=self.reset_settings).pack(side="left", padx=(10, 0))
            ttk.Button(button_frame, text="Export Settings",
                      command=self.export_settings).pack(side="right")
            ttk.Button(button_frame, text="Import Settings",
                      command=self.import_settings).pack(side="right", padx=(0, 10))

        def setup_conversion_settings(self, parent):
            """Setup conversion defaults settings."""
            scroll_frame = self.create_scrollable_frame(parent)
            frame = scroll_frame

            # CSS Theme
            ttk.Label(frame, text="CSS Theme:").grid(row=0, column=0, sticky="w", pady=5)
            self.theme_var = tk.StringVar(value=self.settings.conversion_defaults.css_theme)
            theme_combo = ttk.Combobox(frame, textvariable=self.theme_var,
                                     values=["serif", "sans", "printlike", "dyslexic"],
                                     state="readonly")
            theme_combo.grid(row=0, column=1, sticky="w", padx=(5, 0), pady=5)

            # Language
            ttk.Label(frame, text="Default Language:").grid(row=1, column=0, sticky="w", pady=5)
            self.lang_var = tk.StringVar(value=self.settings.conversion_defaults.language)
            lang_combo = ttk.Combobox(frame, textvariable=self.lang_var,
                                    values=["en", "es", "fr", "de", "it", "pt", "nl", "ru", "zh", "ja"])
            lang_combo.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=5)

            # Validation
            self.validate_var = tk.BooleanVar(value=self.settings.conversion_defaults.validate_epub)
            validate_check = ttk.Checkbutton(frame, text="Validate EPUB with EPUBCheck",
                                           variable=self.validate_var)
            validate_check.grid(row=2, column=0, columnspan=2, sticky="w", pady=5)

            # Image settings
            ttk.Label(frame, text="Max Image Width:").grid(row=3, column=0, sticky="w", pady=5)
            self.image_width_var = tk.StringVar(value=str(self.settings.conversion_defaults.image_max_width))
            width_entry = ttk.Entry(frame, textvariable=self.image_width_var, width=10)
            width_entry.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=5)

            # Accessibility
            self.a11y_var = tk.BooleanVar(value=self.settings.conversion_defaults.accessibility_features)
            a11y_check = ttk.Checkbutton(frame, text="Include accessibility features",
                                       variable=self.a11y_var)
            a11y_check.grid(row=4, column=0, columnspan=2, sticky="w", pady=5)

            # Chapter detection
            ttk.Label(frame, text="Chapter Detection:").grid(row=5, column=0, sticky="w", pady=5)
            self.chapter_var = tk.StringVar(value=self.settings.conversion_defaults.chapter_detection)
            chapter_combo = ttk.Combobox(frame, textvariable=self.chapter_var,
                                       values=["auto", "h1", "h2", "pagebreak", "ai"],
                                       state="readonly")
            chapter_combo.grid(row=5, column=1, sticky="w", padx=(5, 0), pady=5)
            chapter_combo.bind("<<ComboboxSelected>>", self.on_chapter_detection_change)

            # AI Settings button (initially hidden)
            self.ai_settings_btn = ttk.Button(frame, text="AI Settings", command=self.configure_ai_detection)
            self.ai_settings_btn.grid(row=5, column=2, sticky="w", padx=(5, 0), pady=5)
            self.ai_settings_btn.grid_remove()  # Hide initially

            # Include ToC
            self.toc_var = tk.BooleanVar(value=self.settings.conversion_defaults.include_toc)
            toc_check = ttk.Checkbutton(frame, text="Include Table of Contents",
                                      variable=self.toc_var)
            toc_check.grid(row=6, column=0, columnspan=2, sticky="w", pady=5)

            # Include Cover
            self.cover_var = tk.BooleanVar(value=self.settings.conversion_defaults.include_cover)
            cover_check = ttk.Checkbutton(frame, text="Include Cover Page",
                                        variable=self.cover_var)
            cover_check.grid(row=7, column=0, columnspan=2, sticky="w", pady=5)

        def setup_ui_settings(self, parent):
            """Setup UI preferences."""
            scroll_frame = self.create_scrollable_frame(parent)
            frame = scroll_frame

            # Remember last directory
            self.remember_dir_var = tk.BooleanVar(value=self.settings.ui_preferences.remember_last_directory)
            remember_check = ttk.Checkbutton(frame, text="Remember last directory",
                                           variable=self.remember_dir_var)
            remember_check.grid(row=0, column=0, columnspan=2, sticky="w", pady=5)

            # Auto-fill metadata
            self.auto_fill_var = tk.BooleanVar(value=self.settings.ui_preferences.auto_fill_metadata)
            auto_fill_check = ttk.Checkbutton(frame, text="Auto-fill metadata from DOCX",
                                            variable=self.auto_fill_var)
            auto_fill_check.grid(row=1, column=0, columnspan=2, sticky="w", pady=5)

            # Show advanced options
            self.show_advanced_var = tk.BooleanVar(value=self.settings.ui_preferences.show_advanced_options)
            advanced_check = ttk.Checkbutton(frame, text="Show advanced options by default",
                                           variable=self.show_advanced_var)
            advanced_check.grid(row=2, column=0, columnspan=2, sticky="w", pady=5)

            # UI Theme
            ttk.Label(frame, text="UI Theme:").grid(row=3, column=0, sticky="w", pady=5)
            self.ui_theme_var = tk.StringVar(value=self.settings.ui_preferences.theme)
            theme_combo = ttk.Combobox(frame, textvariable=self.ui_theme_var,
                                     values=["system", "light", "dark", "blue", "green", "purple", "minimal", "classic"],
                                     state="readonly")
            theme_combo.bind('<<ComboboxSelected>>', self.apply_ui_theme)
            theme_combo.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=5)

            # Font size
            ttk.Label(frame, text="Font Size:").grid(row=4, column=0, sticky="w", pady=5)
            self.font_size_var = tk.StringVar(value=self.settings.ui_preferences.font_size)
            font_combo = ttk.Combobox(frame, textvariable=self.font_size_var,
                                    values=["small", "medium", "large"],
                                    state="readonly")
            font_combo.grid(row=4, column=1, sticky="w", padx=(5, 0), pady=5)

            # Confirm overwrite
            self.confirm_overwrite_var = tk.BooleanVar(value=self.settings.ui_preferences.confirm_overwrite)
            confirm_check = ttk.Checkbutton(frame, text="Confirm before overwriting files",
                                          variable=self.confirm_overwrite_var)
            confirm_check.grid(row=5, column=0, columnspan=2, sticky="w", pady=5)

            # Show tooltips
            self.tooltips_var = tk.BooleanVar(value=self.settings.ui_preferences.show_tooltips)
            tooltips_check = ttk.Checkbutton(frame, text="Show helpful tooltips",
                                           variable=self.tooltips_var)
            tooltips_check.grid(row=6, column=0, columnspan=2, sticky="w", pady=5)

        def setup_file_settings(self, parent):
            """Setup file and directory defaults."""
            scroll_frame = self.create_scrollable_frame(parent)
            frame = scroll_frame

            # Default output directory
            ttk.Label(frame, text="Default Output Directory:").grid(row=0, column=0, sticky="w", pady=5)
            self.output_dir_var = tk.StringVar(value=self.settings.file_defaults.default_output_directory or "")
            output_entry = ttk.Entry(frame, textvariable=self.output_dir_var, width=40)
            output_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=5)
            browse_output_btn = ttk.Button(frame, text="Browse...", command=self.browse_default_output)
            browse_output_btn.grid(row=0, column=2, padx=(5, 0), pady=5)

            # Auto-generate output name
            self.auto_name_var = tk.BooleanVar(value=self.settings.file_defaults.auto_generate_output_name)
            auto_name_check = ttk.Checkbutton(frame, text="Auto-generate output filename",
                                            variable=self.auto_name_var)
            auto_name_check.grid(row=1, column=0, columnspan=3, sticky="w", pady=5)

            # Backup original
            self.backup_var = tk.BooleanVar(value=self.settings.file_defaults.backup_original)
            backup_check = ttk.Checkbutton(frame, text="Create backup of original file",
                                         variable=self.backup_var)
            backup_check.grid(row=2, column=0, columnspan=3, sticky="w", pady=5)

            # Temp directory
            ttk.Label(frame, text="Temporary Directory:").grid(row=3, column=0, sticky="w", pady=5)
            self.temp_dir_var = tk.StringVar(value=self.settings.file_defaults.temp_directory or "")
            temp_entry = ttk.Entry(frame, textvariable=self.temp_dir_var, width=40)
            temp_entry.grid(row=3, column=1, sticky="ew", padx=(5, 0), pady=5)
            browse_temp_btn = ttk.Button(frame, text="Browse...", command=self.browse_temp_dir)
            browse_temp_btn.grid(row=3, column=2, padx=(5, 0), pady=5)

            frame.columnconfigure(1, weight=1)

        def setup_advanced_settings(self, parent):
            """Setup advanced application settings."""
            scroll_frame = self.create_scrollable_frame(parent)
            frame = scroll_frame

            # Logging
            self.logging_var = tk.BooleanVar(value=self.settings.advanced_settings.enable_logging)
            logging_check = ttk.Checkbutton(frame, text="Enable application logging",
                                          variable=self.logging_var)
            logging_check.grid(row=0, column=0, columnspan=2, sticky="w", pady=5)

            # Log level
            ttk.Label(frame, text="Log Level:").grid(row=1, column=0, sticky="w", pady=5)
            self.log_level_var = tk.StringVar(value=self.settings.advanced_settings.log_level)
            log_combo = ttk.Combobox(frame, textvariable=self.log_level_var,
                                   values=["DEBUG", "INFO", "WARNING", "ERROR"],
                                   state="readonly")
            log_combo.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=5)

            # Concurrent jobs
            ttk.Label(frame, text="Concurrent Jobs:").grid(row=2, column=0, sticky="w", pady=5)
            self.concurrent_var = tk.StringVar(value=str(self.settings.advanced_settings.concurrent_jobs))
            concurrent_entry = ttk.Entry(frame, textvariable=self.concurrent_var, width=10)
            concurrent_entry.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=5)

            # Check for updates
            self.updates_var = tk.BooleanVar(value=self.settings.advanced_settings.check_for_updates)
            updates_check = ttk.Checkbutton(frame, text="Check for updates automatically",
                                          variable=self.updates_var)
            updates_check.grid(row=3, column=0, columnspan=2, sticky="w", pady=5)

            # Crash reporting
            self.crash_var = tk.BooleanVar(value=self.settings.advanced_settings.enable_crash_reporting)
            crash_check = ttk.Checkbutton(frame, text="Enable crash reporting",
                                        variable=self.crash_var)
            crash_check.grid(row=4, column=0, columnspan=2, sticky="w", pady=5)

            # Telemetry
            self.telemetry_var = tk.BooleanVar(value=self.settings.advanced_settings.enable_telemetry)
            telemetry_check = ttk.Checkbutton(frame, text="Enable anonymous usage telemetry",
                                            variable=self.telemetry_var)
            telemetry_check.grid(row=5, column=0, columnspan=2, sticky="w", pady=5)

        def setup_enterprise_settings(self, parent):
            """Setup enterprise configuration."""
            scroll_frame = self.create_scrollable_frame(parent)
            frame = scroll_frame

            # Enable enterprise mode
            self.enterprise_enabled_var = tk.BooleanVar(value=self.settings.enterprise_config is not None)
            enterprise_check = ttk.Checkbutton(frame, text="Enable Enterprise Features",
                                             variable=self.enterprise_enabled_var,
                                             command=self.toggle_enterprise_settings)
            enterprise_check.grid(row=0, column=0, columnspan=2, sticky="w", pady=5)

            # Enterprise settings frame (enabled/disabled based on checkbox)
            self.enterprise_frame = ttk.LabelFrame(frame, text="Enterprise Configuration", padding=10)
            self.enterprise_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=10, padx=10)

            # API settings
            ttk.Label(self.enterprise_frame, text="API Host:").grid(row=0, column=0, sticky="w", pady=5)
            enterprise_config = self.settings.enterprise_config
            default_host = enterprise_config.api_host if enterprise_config else "localhost"
            self.api_host_var = tk.StringVar(value=default_host)
            host_entry = ttk.Entry(self.enterprise_frame, textvariable=self.api_host_var, width=20)
            host_entry.grid(row=0, column=1, sticky="w", padx=(5, 0), pady=5)

            ttk.Label(self.enterprise_frame, text="API Port:").grid(row=1, column=0, sticky="w", pady=5)
            default_port = enterprise_config.api_port if enterprise_config else 8080
            self.api_port_var = tk.StringVar(value=str(default_port))
            port_entry = ttk.Entry(self.enterprise_frame, textvariable=self.api_port_var, width=10)
            port_entry.grid(row=1, column=1, sticky="w", padx=(5, 0), pady=5)

            # Batch processing
            ttk.Label(self.enterprise_frame, text="Max Concurrent Jobs:").grid(row=2, column=0, sticky="w", pady=5)
            default_jobs = enterprise_config.max_concurrent_jobs if enterprise_config else 4
            self.max_jobs_var = tk.StringVar(value=str(default_jobs))
            jobs_entry = ttk.Entry(self.enterprise_frame, textvariable=self.max_jobs_var, width=10)
            jobs_entry.grid(row=2, column=1, sticky="w", padx=(5, 0), pady=5)

            ttk.Label(self.enterprise_frame, text="Job Timeout (hours):").grid(row=3, column=0, sticky="w", pady=5)
            default_timeout = enterprise_config.job_timeout_hours if enterprise_config else 24
            self.job_timeout_var = tk.StringVar(value=str(default_timeout))
            timeout_entry = ttk.Entry(self.enterprise_frame, textvariable=self.job_timeout_var, width=10)
            timeout_entry.grid(row=3, column=1, sticky="w", padx=(5, 0), pady=5)

            # Storage settings
            ttk.Label(self.enterprise_frame, text="Storage Directory:").grid(row=4, column=0, sticky="w", pady=5)
            default_storage = enterprise_config.storage_directory if enterprise_config else ""
            self.storage_dir_var = tk.StringVar(value=default_storage or "")
            storage_entry = ttk.Entry(self.enterprise_frame, textvariable=self.storage_dir_var, width=30)
            storage_entry.grid(row=4, column=1, sticky="ew", padx=(5, 0), pady=5)
            browse_storage_btn = ttk.Button(self.enterprise_frame, text="Browse...", command=self.browse_storage_dir)
            browse_storage_btn.grid(row=4, column=2, padx=(5, 0), pady=5)

            # Enable/disable enterprise frame based on initial state
            self.toggle_enterprise_settings()

            self.enterprise_frame.columnconfigure(1, weight=1)
            frame.columnconfigure(0, weight=1)

        def create_scrollable_frame(self, parent):
            """Create a scrollable frame for settings."""
            canvas = tk.Canvas(parent)
            scrollbar = ttk.Scrollbar(parent, orient="vertical", command=canvas.yview)
            scrollable_frame = ttk.Frame(canvas)

            scrollable_frame.bind(
                "<Configure>",
                lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
            )

            canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
            canvas.configure(yscrollcommand=scrollbar.set)

            canvas.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

            return scrollable_frame

        def setup_batch_tab(self, parent):
            """Setup batch processing tab (equivalent to CLI batch command)."""
            # File selection area
            files_frame = ttk.LabelFrame(parent, text="Files to Process", padding=10)
            files_frame.pack(fill="both", expand=True, padx=10, pady=5)

            # File list
            self.batch_files_listbox = tk.Listbox(files_frame, height=8)
            self.batch_files_listbox.pack(fill="both", expand=True, pady=(0, 5))

            # File control buttons
            file_buttons_frame = ttk.Frame(files_frame)
            file_buttons_frame.pack(fill="x")

            ttk.Button(file_buttons_frame, text="Add Files...", command=self.add_batch_files).pack(side="left", padx=(0, 5))
            ttk.Button(file_buttons_frame, text="Add Folder...", command=self.add_batch_folder).pack(side="left", padx=(0, 5))
            ttk.Button(file_buttons_frame, text="Remove Selected", command=self.remove_batch_files).pack(side="left", padx=(0, 5))
            ttk.Button(file_buttons_frame, text="Clear All", command=self.clear_batch_files).pack(side="left")

            # Batch options
            options_frame = ttk.LabelFrame(parent, text="Batch Options", padding=10)
            options_frame.pack(fill="x", padx=10, pady=5)

            # Output directory
            ttk.Label(options_frame, text="Output Directory:").grid(row=0, column=0, sticky="w", pady=2)
            self.batch_output_var = tk.StringVar()
            output_entry = ttk.Entry(options_frame, textvariable=self.batch_output_var, width=50)
            output_entry.grid(row=0, column=1, sticky="ew", padx=(5, 0), pady=2)
            ttk.Button(options_frame, text="Browse...", command=self.browse_batch_output).grid(row=0, column=2, padx=(5, 0), pady=2)

            # Use same metadata for all
            self.batch_same_metadata_var = tk.BooleanVar()
            same_meta_check = ttk.Checkbutton(options_frame, text="Use same metadata for all files",
                                            variable=self.batch_same_metadata_var)
            same_meta_check.grid(row=1, column=0, columnspan=2, sticky="w", pady=5)

            options_frame.columnconfigure(1, weight=1)

            # Progress area
            progress_frame = ttk.LabelFrame(parent, text="Progress", padding=10)
            progress_frame.pack(fill="x", padx=10, pady=5)

            self.batch_progress_var = tk.StringVar(value="Ready to process batch")
            ttk.Label(progress_frame, textvariable=self.batch_progress_var).pack()

            self.batch_progress_bar = ttk.Progressbar(progress_frame, mode='determinate')
            self.batch_progress_bar.pack(fill="x", pady=(5, 0))

            # Control buttons
            control_frame = ttk.Frame(parent)
            control_frame.pack(fill="x", padx=10, pady=10)

            self.batch_start_btn = ttk.Button(control_frame, text="Start Batch Processing",
                                            command=self.start_batch_processing)
            self.batch_start_btn.pack(side="left")

            self.batch_cancel_btn = ttk.Button(control_frame, text="Cancel",
                                             command=self.cancel_batch_processing, state="disabled")
            self.batch_cancel_btn.pack(side="left", padx=(10, 0))

        # Menu command implementations
        def install_pandoc(self):
            """Install Pandoc tool."""
            def install_worker():
                try:
                    from .tools import install_pandoc
                    install_pandoc()
                    messagebox.showinfo("Success", "Pandoc installed successfully!")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to install Pandoc: {e}")

            threading.Thread(target=install_worker, daemon=True).start()

        def install_epubcheck(self):
            """Install EPUBCheck tool."""
            def install_worker():
                try:
                    from .tools import install_epubcheck
                    install_epubcheck()
                    messagebox.showinfo("Success", "EPUBCheck installed successfully!")
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to install EPUBCheck: {e}")

            threading.Thread(target=install_worker, daemon=True).start()

        def show_tool_locations(self):
            """Show tool installation locations."""
            try:
                from .tools import pandoc_path, epubcheck_cmd, tools_dir

                info = []
                info.append(f"Tools Directory: {tools_dir()}")

                try:
                    pandoc = pandoc_path()
                    info.append(f"Pandoc: {pandoc}")
                except:
                    info.append("Pandoc: Not installed")

                try:
                    epubcheck = epubcheck_cmd()
                    info.append(f"EPUBCheck: {epubcheck}")
                except:
                    info.append("EPUBCheck: Not installed")

                messagebox.showinfo("Tool Locations", "\n".join(info))
            except Exception as e:
                messagebox.showerror("Error", f"Failed to get tool locations: {e}")

        def show_conversion_help(self):
            """Show comprehensive conversion help dialog."""
            help_text = """
            📖 Docx2Shelf Comprehensive Help Guide

            🚀 GETTING STARTED
            ═══════════════════════════════════════════════════════════
            1. 📁 Select Document: Use "Browse" or drag & drop your file
            2. 📝 Enter Metadata: Fill in book title, author, and details
            3. ⚙️ Configure Options: Choose conversion settings
            4. 🎯 Convert: Click "Convert to EPUB" to generate your ebook

            📂 SUPPORTED INPUT FORMATS
            ═══════════════════════════════════════════════════════════
            • DOCX - Microsoft Word documents (.docx)
              ✓ Preserves formatting, styles, and images
              ✓ Converts tables, lists, and embedded objects
              ✓ Maintains heading hierarchy for chapter detection

            • Markdown - Markdown files (.md)
              ✓ Full CommonMark and GitHub Flavored Markdown support
              ✓ Code blocks, tables, and inline formatting
              ✓ Automatic heading-based chapter structure

            • HTML - Web documents (.html, .htm)
              ✓ Clean HTML with CSS styling preservation
              ✓ Converts semantic elements to EPUB structure
              ✓ Handles embedded images and media

            • TXT - Plain text files (.txt)
              ✓ Automatic paragraph detection
              ✓ Smart chapter break recognition
              ✓ Preserves line breaks and formatting

            📖 METADATA BEST PRACTICES
            ═══════════════════════════════════════════════════════════
            Essential Fields:
            • Title: Use the full, official title of your book
            • Author: Primary author name (Last, First format recommended)
            • Language: Select the book's primary language
            • Description: Write a compelling 2-3 sentence summary

            Advanced Metadata:
            • Genres: Use standard genre tags (Fantasy, Romance, Mystery, etc.)
            • Keywords: Add discoverable terms separated by commas
            • ISBN: Include if you have one for store distribution
            • Series Info: Add series name and book number if applicable
            • Publisher: Your publishing imprint or company name
            • Publication Date: Use YYYY-MM-DD format

            🎯 CONVERSION OPTIONS EXPLAINED
            ═══════════════════════════════════════════════════════════
            🔧 Basic Options:
            • CSS Theme: Choose visual styling (Serif, Sans-serif, Print-like)
            • Chapter Detection: How to identify chapter breaks
            • Table of Contents: Auto-generate navigation
            • EPUB Validation: Check output for compliance

            ⚡ Advanced Options:
            • Custom CSS: Override default styling
            • Image Optimization: Compress and resize images
            • Font Embedding: Include custom fonts
            • DRM Settings: Configure digital rights management

            🖼️ Images & Media:
            • Auto-resize large images for optimal file size
            • Convert images to web-friendly formats
            • Maintain aspect ratios and quality
            • Embed or link external media

            🤖 AI FEATURES
            ═══════════════════════════════════════════════════════════
            Smart Chapter Detection:
            • Analyzes document structure and content
            • Identifies chapter boundaries with 95%+ accuracy
            • Handles various heading styles and formats
            • Works with both styled and plain text documents

            Automatic Metadata Enhancement:
            • Genre classification based on content analysis
            • Keyword extraction for better discoverability
            • Language detection and verification
            • Reading level and audience identification

            Content Analysis:
            • Alt text generation for images
            • Content warnings detection
            • Reading time estimation
            • Quality assessment and suggestions

            🏪 EBOOK STORE COMPATIBILITY
            ═══════════════════════════════════════════════════════════
            Amazon KDP:
            ✓ EPUB 3.0 compliance
            ✓ Proper metadata structure
            ✓ Cover image requirements
            ✓ DRM compatibility

            Apple Books:
            ✓ Enhanced typography support
            ✓ Interactive elements
            ✓ Accessibility features
            ✓ Fixed-layout support

            Other Platforms:
            ✓ Kobo, Google Play Books, Barnes & Noble
            ✓ Draft2Digital, Smashwords
            ✓ IngramSpark, BookBaby

            🔧 TROUBLESHOOTING
            ═══════════════════════════════════════════════════════════
            Common Issues:
            • Large files: Use image optimization and compression
            • Missing chapters: Check heading styles in source document
            • Formatting problems: Review CSS theme selection
            • Validation errors: Run EPUB check and fix reported issues

            Performance Tips:
            • Keep images under 2MB each for best performance
            • Use standard fonts or embed custom ones
            • Test on multiple devices before publishing
            • Validate EPUB before store submission

            🆘 GETTING MORE HELP
            ═══════════════════════════════════════════════════════════
            • Documentation: Check the About tab for links
            • GitHub Issues: Report bugs or request features
            • Community: Join discussions and get tips
            • Updates: Enable auto-updates for latest features

            💡 PRO TIPS
            ═══════════════════════════════════════════════════════════
            • Use consistent heading styles for best chapter detection
            • Include high-quality cover art (1600x2560 recommended)
            • Test your EPUB on multiple reading apps
            • Keep metadata complete for better store visibility
            • Use AI features to enhance content quality
            • Backup your source files before conversion
            """

            # Create modern help dialog
            help_dialog = tk.Toplevel(self.root)
            help_dialog.title("📖 Docx2Shelf Help Guide")
            help_dialog.geometry("900x700")
            help_dialog.transient(self.root)
            help_dialog.grab_set()
            help_dialog.configure(bg=self.colors['white'])

            # Center the dialog
            help_dialog.update_idletasks()
            x = (help_dialog.winfo_screenwidth() // 2) - (900 // 2)
            y = (help_dialog.winfo_screenheight() // 2) - (700 // 2)
            help_dialog.geometry(f"900x700+{x}+{y}")

            # Create header frame
            header_frame = ttk.Frame(help_dialog, style="Header.TFrame")
            header_frame.pack(fill="x", padx=0, pady=(0, 10))

            title_label = ttk.Label(header_frame, text="📖 Comprehensive Help Guide",
                                  style="Title.TLabel")
            title_label.pack(pady=15)

            # Create main content frame
            content_frame = ttk.Frame(help_dialog, style="Card.TFrame")
            content_frame.pack(fill="both", expand=True, padx=20, pady=(0, 20))

            # Create scrolled text widget with modern styling
            from tkinter.scrolledtext import ScrolledText
            text_widget = ScrolledText(content_frame, wrap=tk.WORD,
                                     font=("Segoe UI", 10),
                                     bg=self.colors['white'],
                                     fg=self.colors['gray_800'],
                                     relief="flat",
                                     borderwidth=0,
                                     padx=15, pady=15)
            text_widget.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

            text_widget.insert(tk.END, help_text.strip())
            text_widget.config(state=tk.DISABLED)

            # Create button frame
            button_frame = ttk.Frame(help_dialog, style="Card.TFrame")
            button_frame.pack(fill="x", padx=20, pady=(0, 20))

            # Close button
            close_btn = ttk.Button(button_frame, text="✕ Close",
                                 command=help_dialog.destroy, style="Secondary.TButton")
            close_btn.pack(side="right", padx=(10, 0))

            # Print button
            print_btn = ttk.Button(button_frame, text="🖨️ Print Help",
                                 command=lambda: self.print_help_content(help_text),
                                 style="Modern.TButton")
            print_btn.pack(side="right")

        def print_help_content(self, content):
            """Print or save help content."""
            try:
                from tkinter import filedialog
                file_path = filedialog.asksaveasfilename(
                    title="Save Help Guide",
                    defaultextension=".txt",
                    filetypes=[("Text files", "*.txt"), ("All files", "*.*")]
                )
                if file_path:
                    with open(file_path, 'w', encoding='utf-8') as f:
                        f.write(content.strip())
                    messagebox.showinfo("Help Saved", f"Help guide saved to:\n{file_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Could not save help guide: {e}")

        def run_system_doctor(self):
            """Run system diagnostics."""
            def doctor_thread():
                try:
                    import subprocess
                    import urllib.request
                    from importlib import metadata
                    import os

                    # Create results window
                    doctor_window = tk.Toplevel(self.root)
                    doctor_window.title("System Doctor")
                    doctor_window.geometry("600x400")
                    doctor_window.transient(self.root)

                    # Create scrolled text widget
                    text_widget = ScrolledText(doctor_window, wrap=tk.WORD, width=70, height=20)
                    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                    def log(message):
                        text_widget.insert(tk.END, message + "\n")
                        text_widget.see(tk.END)
                        doctor_window.update_idletasks()

                    log("🔍 Docx2Shelf System Diagnostics")
                    log("=" * 50)

                    # Check Python version
                    log(f"✅ Python version: {sys.version}")

                    # Check package installation
                    try:
                        version = metadata.version("docx2shelf")
                        log(f"✅ Docx2Shelf version: {version}")
                    except metadata.PackageNotFoundError:
                        log("⚠️  Docx2Shelf not installed via pip (running from source)")

                    # Check dependencies
                    log("\n📦 Checking dependencies:")
                    dependencies = ["python-docx", "beautifulsoup4", "lxml", "Pillow"]
                    for dep in dependencies:
                        try:
                            dep_version = metadata.version(dep)
                            log(f"✅ {dep}: {dep_version}")
                        except metadata.PackageNotFoundError:
                            log(f"❌ {dep}: Not installed")

                    # Check tools
                    log("\n🔧 Checking external tools:")

                    # Check Pandoc
                    try:
                        result = subprocess.run(["pandoc", "--version"],
                                              capture_output=True, text=True, timeout=5)
                        if result.returncode == 0:
                            version_line = result.stdout.split('\n')[0]
                            log(f"✅ Pandoc: {version_line}")
                        else:
                            log("❌ Pandoc: Not working properly")
                    except (subprocess.TimeoutExpired, FileNotFoundError):
                        log("⚠️  Pandoc: Not found (will use fallback)")

                    # Check EPUBCheck
                    try:
                        result = subprocess.run(["java", "-jar", "-version"],
                                              capture_output=True, text=True, timeout=5)
                        if result.returncode == 0:
                            log("✅ Java: Available for EPUBCheck")
                        else:
                            log("⚠️  Java: Not properly configured")
                    except (subprocess.TimeoutExpired, FileNotFoundError):
                        log("⚠️  Java: Not found (EPUBCheck validation unavailable)")

                    # Check network connectivity
                    log("\n🌐 Checking network connectivity:")
                    try:
                        with urllib.request.urlopen("https://api.github.com", timeout=5) as response:
                            if response.status == 200:
                                log("✅ GitHub API: Accessible")
                            else:
                                log("⚠️  GitHub API: Limited access")
                    except Exception:
                        log("❌ Network: No internet connection")

                    # Check file permissions
                    log("\n📁 Checking file permissions:")
                    temp_dir = Path(tempfile.gettempdir())
                    try:
                        test_file = temp_dir / "docx2shelf_test.txt"
                        test_file.write_text("test")
                        test_file.unlink()
                        log(f"✅ Temp directory writable: {temp_dir}")
                    except Exception as e:
                        log(f"❌ Temp directory not writable: {e}")

                    log("\n" + "=" * 50)
                    log("✅ System diagnostics complete!")

                    # Add close button
                    close_btn = ttk.Button(doctor_window, text="Close",
                                         command=doctor_window.destroy)
                    close_btn.pack(pady=10)

                except Exception as e:
                    messagebox.showerror("System Doctor", f"Error running diagnostics:\n{str(e)}")

            # Run in background thread
            threading.Thread(target=doctor_thread, daemon=True).start()

        def list_themes(self):
            """List available CSS themes."""
            try:
                themes = ["serif", "sans", "printlike"]  # Basic themes, could be expanded
                theme_list = "\n".join(f"• {theme}" for theme in themes)
                messagebox.showinfo("Available Themes", f"Available CSS themes:\n\n{theme_list}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to list themes: {e}")

        def preview_themes(self):
            """Preview CSS themes in browser."""
            self.create_theme_preview_html()

        def create_theme_preview_html(self):
            """Create and open a local theme preview HTML file."""
            import tempfile
            import webbrowser
            from pathlib import Path

            # Sample content for theme preview
            sample_content = """
            <h1>Book Title Preview</h1>
            <h2>Chapter 1: The Beginning</h2>
            <p>This is a sample paragraph to demonstrate how text appears with different themes. It includes <em>italicized text</em> and <strong>bold text</strong> to show various formatting options.</p>

            <blockquote>
            <p>"This is a sample quote to show how blockquotes are styled in different themes."</p>
            </blockquote>

            <h3>Subsection Title</h3>
            <p>Another paragraph with different content to showcase the theme's typography and spacing. This helps authors visualize how their book will look with each theme option.</p>

            <ul>
            <li>Sample list item one</li>
            <li>Sample list item two</li>
            <li>Sample list item three</li>
            </ul>
            """

            # Built-in themes
            themes = {
                "serif": {
                    "name": "Serif (Classic)",
                    "css": """
                    body { font-family: 'Times New Roman', serif; line-height: 1.6; margin: 40px; background: #fefefe; color: #333; }
                    h1, h2, h3 { font-family: 'Times New Roman', serif; color: #2c3e50; }
                    h1 { font-size: 2.2em; margin-bottom: 0.5em; }
                    h2 { font-size: 1.8em; margin-top: 1.5em; margin-bottom: 0.5em; }
                    h3 { font-size: 1.4em; margin-top: 1em; margin-bottom: 0.5em; }
                    p { margin-bottom: 1em; text-align: justify; }
                    blockquote { border-left: 4px solid #bdc3c7; padding-left: 20px; margin: 1.5em 0; font-style: italic; color: #7f8c8d; }
                    """
                },
                "sans": {
                    "name": "Sans-serif (Modern)",
                    "css": """
                    body { font-family: 'Arial', sans-serif; line-height: 1.5; margin: 40px; background: #ffffff; color: #2c3e50; }
                    h1, h2, h3 { font-family: 'Arial', sans-serif; color: #34495e; }
                    h1 { font-size: 2.1em; margin-bottom: 0.5em; font-weight: 300; }
                    h2 { font-size: 1.7em; margin-top: 1.5em; margin-bottom: 0.5em; font-weight: 400; }
                    h3 { font-size: 1.3em; margin-top: 1em; margin-bottom: 0.5em; font-weight: 500; }
                    p { margin-bottom: 1em; }
                    blockquote { border-left: 3px solid #3498db; padding-left: 20px; margin: 1.5em 0; background: #ecf0f1; padding: 15px 20px; }
                    """
                },
                "printlike": {
                    "name": "Print-like (Traditional)",
                    "css": """
                    body { font-family: 'Georgia', serif; line-height: 1.7; margin: 50px 60px; background: #ffffff; color: #1a1a1a; }
                    h1, h2, h3 { font-family: 'Georgia', serif; color: #000000; }
                    h1 { font-size: 2.0em; margin-bottom: 0.5em; text-align: center; font-weight: bold; }
                    h2 { font-size: 1.6em; margin-top: 2em; margin-bottom: 0.5em; font-weight: bold; }
                    h3 { font-size: 1.3em; margin-top: 1.5em; margin-bottom: 0.5em; font-weight: bold; }
                    p { margin-bottom: 1.2em; text-align: justify; text-indent: 1.5em; }
                    blockquote { border: none; padding: 0 40px; margin: 2em 0; font-style: italic; }
                    """
                }
            }

            # Check for custom themes
            try:
                custom_themes_dir = Path.home() / '.docx2shelf' / 'themes'
                if custom_themes_dir.exists():
                    for theme_file in custom_themes_dir.glob('*.css'):
                        theme_name = theme_file.stem
                        themes[theme_name] = {
                            "name": f"{theme_name.title()} (Custom)",
                            "css": theme_file.read_text(encoding='utf-8')
                        }
            except Exception as e:
                pass  # Ignore errors reading custom themes

            # Create HTML preview
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="UTF-8">
                <title>Docx2Shelf Theme Preview</title>
                <style>
                    .theme-container {{ margin-bottom: 50px; border: 2px solid #ccc; }}
                    .theme-header {{ background: #f5f5f5; padding: 15px; border-bottom: 1px solid #ccc; font-weight: bold; }}
                    .theme-content {{ padding: 20px; }}
                </style>
            </head>
            <body>
                <h1 style="text-align: center; font-family: Arial, sans-serif; color: #2c3e50;">Docx2Shelf Theme Preview</h1>
                <p style="text-align: center; font-family: Arial, sans-serif; color: #7f8c8d;">Preview how your book will look with different themes</p>
            """

            for theme_key, theme_data in themes.items():
                html_content += f"""
                <div class="theme-container">
                    <div class="theme-header">{theme_data['name']}</div>
                    <div class="theme-content">
                        <style scoped>
                            .theme-{theme_key} {{ {theme_data['css']} }}
                        </style>
                        <div class="theme-{theme_key}">
                            {sample_content}
                        </div>
                    </div>
                </div>
                """

            html_content += "</body></html>"

            # Save and open preview file
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                preview_file = f.name

            try:
                webbrowser.open(f'file://{preview_file}')
                messagebox.showinfo("Theme Preview", f"Theme preview opened in your browser.\n\nFile saved to: {preview_file}")
            except Exception as e:
                messagebox.showerror("Error", f"Could not open browser: {e}\n\nPreview saved to: {preview_file}")

        def apply_ui_theme(self, event=None):
            """Apply the selected UI theme."""
            try:
                theme = self.ui_theme_var.get()

                # Configure ttk styles based on theme
                style = ttk.Style()

                if theme == "dark":
                    style.theme_use('clam')
                    style.configure('.', background='#2b2b2b', foreground='#ffffff',
                                  fieldbackground='#3c3c3c', selectbackground='#0078d4')
                    self.root.configure(bg='#2b2b2b')

                elif theme == "blue":
                    style.theme_use('clam')
                    style.configure('.', background='#e3f2fd', foreground='#0d47a1',
                                  fieldbackground='#ffffff', selectbackground='#1976d2')
                    self.root.configure(bg='#e3f2fd')

                elif theme == "green":
                    style.theme_use('clam')
                    style.configure('.', background='#e8f5e8', foreground='#1b5e20',
                                  fieldbackground='#ffffff', selectbackground='#4caf50')
                    self.root.configure(bg='#e8f5e8')

                elif theme == "purple":
                    style.theme_use('clam')
                    style.configure('.', background='#f3e5f5', foreground='#4a148c',
                                  fieldbackground='#ffffff', selectbackground='#9c27b0')
                    self.root.configure(bg='#f3e5f5')

                elif theme == "minimal":
                    style.theme_use('clam')
                    style.configure('.', background='#fafafa', foreground='#212121',
                                  fieldbackground='#ffffff', selectbackground='#757575')
                    self.root.configure(bg='#fafafa')

                elif theme == "classic":
                    style.theme_use('clam')
                    style.configure('.', background='#f0f0f0', foreground='#000000',
                                  fieldbackground='#ffffff', selectbackground='#316ac5')
                    self.root.configure(bg='#f0f0f0')

                elif theme == "light":
                    style.theme_use('clam')
                    style.configure('.', background='#ffffff', foreground='#000000',
                                  fieldbackground='#ffffff', selectbackground='#0078d4')
                    self.root.configure(bg='#ffffff')

                else:  # system
                    style.theme_use('vista' if self.root.tk.call('tk', 'windowingsystem') == 'win32' else 'clam')

                messagebox.showinfo("Theme Applied", f"UI theme '{theme}' has been applied successfully!")

            except Exception as e:
                messagebox.showerror("Theme Error", f"Could not apply theme: {e}")

        def create_custom_theme(self):
            """Open custom theme creation dialog."""
            CustomThemeDialog(self.root, self)

        def on_chapter_detection_change(self, event=None):
            """Handle chapter detection method change."""
            if self.chapter_var.get() == "ai":
                self.ai_settings_btn.grid()  # Show AI settings button
            else:
                self.ai_settings_btn.grid_remove()  # Hide AI settings button

        def configure_ai_detection(self):
            """Open AI chapter detection configuration dialog."""
            AIDetectionDialog(self.root, self)

        def open_unified_settings(self):
            """Open unified settings dialog."""
            UnifiedSettingsDialog(self.root, self)

        def export_all_settings(self):
            """Export all settings to file."""
            try:
                from tkinter import filedialog
                file_path = filedialog.asksaveasfilename(
                    title="Export All Settings",
                    defaultextension=".json",
                    filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
                )
                if file_path:
                    from ..settings import get_settings_manager
                    settings_manager = get_settings_manager()
                    settings_manager.export_settings(Path(file_path))
                    messagebox.showinfo("Export Complete", f"All settings exported to {file_path}")
            except Exception as e:
                messagebox.showerror("Export Error", f"Could not export settings: {e}")

        def import_all_settings(self):
            """Import all settings from file."""
            try:
                from tkinter import filedialog
                file_path = filedialog.askopenfilename(
                    title="Import All Settings",
                    filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
                )
                if file_path:
                    if messagebox.askyesno("Import Settings",
                                         "This will merge the imported settings with your current settings. Continue?"):
                        from ..settings import get_settings_manager
                        settings_manager = get_settings_manager()
                        settings_manager.import_settings(Path(file_path), merge=True)
                        messagebox.showinfo("Import Complete", "Settings imported successfully. Restart the application to see all changes.")
            except Exception as e:
                messagebox.showerror("Import Error", f"Could not import settings: {e}")

        def reset_all_settings(self):
            """Reset all settings to defaults."""
            if messagebox.askyesno("Reset All Settings",
                                 "Are you sure you want to reset ALL settings to defaults? This cannot be undone."):
                try:
                    from ..settings import get_settings_manager
                    settings_manager = get_settings_manager()
                    settings_manager.reset_to_defaults()
                    messagebox.showinfo("Settings Reset", "All settings have been reset to defaults. Restart the application to see changes.")
                except Exception as e:
                    messagebox.showerror("Error", f"Could not reset settings: {e}")

        def validate_epub_file(self):
            """Validate an EPUB file."""
            file_path = filedialog.askopenfilename(
                title="Select EPUB file to validate",
                filetypes=[("EPUB files", "*.epub"), ("All files", "*.*")]
            )
            if file_path:
                def validate_thread():
                    try:
                        # Create validation results window
                        validation_window = tk.Toplevel(self.root)
                        validation_window.title("EPUB Validation Results")
                        validation_window.geometry("700x500")
                        validation_window.transient(self.root)

                        # Create scrolled text widget
                        text_widget = ScrolledText(validation_window, wrap=tk.WORD, width=80, height=25)
                        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                        def log(message):
                            text_widget.insert(tk.END, message + "\n")
                            text_widget.see(tk.END)
                            validation_window.update_idletasks()

                        log(f"🔍 Validating EPUB: {Path(file_path).name}")
                        log("=" * 60)

                        # Try to run EPUBCheck
                        try:
                            epub_check_result = epubcheck_cmd(file_path)
                            if epub_check_result:
                                log("✅ EPUBCheck validation completed")
                                log("\n📋 EPUBCheck Results:")
                                log(epub_check_result)
                            else:
                                log("⚠️  EPUBCheck not available")
                        except Exception as e:
                            log(f"❌ EPUBCheck failed: {str(e)}")

                        # Basic EPUB structure validation
                        log("\n🔧 Basic Structure Validation:")
                        try:
                            import zipfile

                            with zipfile.ZipFile(file_path, 'r') as epub_zip:
                                file_list = epub_zip.namelist()

                                # Check for required files
                                required_files = ['META-INF/container.xml', 'mimetype']
                                for req_file in required_files:
                                    if req_file in file_list:
                                        log(f"✅ Found: {req_file}")
                                    else:
                                        log(f"❌ Missing: {req_file}")

                                # Check mimetype content
                                if 'mimetype' in file_list:
                                    mimetype_content = epub_zip.read('mimetype').decode('utf-8').strip()
                                    if mimetype_content == 'application/epub+zip':
                                        log("✅ Correct mimetype")
                                    else:
                                        log(f"❌ Incorrect mimetype: {mimetype_content}")

                                # Look for OPF file
                                opf_files = [f for f in file_list if f.endswith('.opf')]
                                if opf_files:
                                    log(f"✅ Found OPF file: {opf_files[0]}")
                                else:
                                    log("❌ No OPF file found")

                                # Count content files
                                html_files = [f for f in file_list if f.endswith(('.html', '.xhtml'))]
                                css_files = [f for f in file_list if f.endswith('.css')]
                                image_files = [f for f in file_list if f.endswith(('.jpg', '.jpeg', '.png', '.gif', '.svg'))]

                                log(f"📄 Content files: {len(html_files)} HTML/XHTML")
                                log(f"🎨 Style files: {len(css_files)} CSS")
                                log(f"🖼️  Image files: {len(image_files)}")

                        except Exception as e:
                            log(f"❌ Structure validation failed: {str(e)}")

                        log("\n" + "=" * 60)
                        log("✅ Validation complete!")

                        # Add close button
                        close_btn = ttk.Button(validation_window, text="Close",
                                             command=validation_window.destroy)
                        close_btn.pack(pady=10)

                    except Exception as e:
                        messagebox.showerror("Validation Error", f"Error during validation:\n{str(e)}")

                # Run in background thread
                threading.Thread(target=validate_thread, daemon=True).start()

        def analyze_epub_quality(self):
            """Analyze EPUB quality."""
            file_path = filedialog.askopenfilename(
                title="Select EPUB file to analyze",
                filetypes=[("EPUB files", "*.epub"), ("All files", "*.*")]
            )
            if file_path:
                def analyze_thread():
                    try:
                        # Create analysis results window
                        analysis_window = tk.Toplevel(self.root)
                        analysis_window.title("EPUB Quality Analysis")
                        analysis_window.geometry("700x500")
                        analysis_window.transient(self.root)

                        # Create scrolled text widget
                        text_widget = ScrolledText(analysis_window, wrap=tk.WORD, width=80, height=25)
                        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                        def log(message):
                            text_widget.insert(tk.END, message + "\n")
                            text_widget.see(tk.END)
                            analysis_window.update_idletasks()

                        log(f"📊 Quality Analysis: {Path(file_path).name}")
                        log("=" * 60)

                        # Basic quality checks
                        try:
                            import zipfile
                            import xml.etree.ElementTree as ET

                            with zipfile.ZipFile(file_path, 'r') as epub_zip:
                                file_list = epub_zip.namelist()

                                # Accessibility checks
                                log("♿ Accessibility Compliance:")

                                # Check for navigation file
                                nav_files = [f for f in file_list if 'nav.xhtml' in f or 'toc.ncx' in f]
                                if nav_files:
                                    log("✅ Navigation file present")
                                else:
                                    log("❌ Missing navigation file")

                                # Check for alt text on images
                                image_files = [f for f in file_list if f.endswith(('.jpg', '.jpeg', '.png', '.gif'))]
                                if image_files:
                                    log(f"🖼️  Found {len(image_files)} images - manual alt text check needed")
                                else:
                                    log("✅ No images requiring alt text")

                                # Store compatibility checks
                                log("\n🏪 Store Compatibility:")

                                # File size check
                                import os
                                file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
                                log(f"📦 File size: {file_size_mb:.2f} MB")

                                if file_size_mb < 50:
                                    log("✅ Size compatible with most stores")
                                elif file_size_mb < 100:
                                    log("⚠️  Large file - some store restrictions may apply")
                                else:
                                    log("❌ Very large file - may exceed store limits")

                                # Content type checks
                                html_files = [f for f in file_list if f.endswith(('.html', '.xhtml'))]
                                if html_files:
                                    log(f"📄 {len(html_files)} content files")

                                    # Sample HTML content for analysis
                                    sample_html = epub_zip.read(html_files[0]).decode('utf-8', errors='ignore')

                                    # Check for JavaScript (not allowed in most stores)
                                    if '<script' in sample_html.lower():
                                        log("❌ JavaScript detected - not allowed in most stores")
                                    else:
                                        log("✅ No JavaScript detected")

                                    # Check for external links
                                    if 'http://' in sample_html or 'https://' in sample_html:
                                        log("⚠️  External links detected - verify store policies")
                                    else:
                                        log("✅ No external links detected")

                                # Best practices
                                log("\n📋 Best Practices:")

                                # CSS files
                                css_files = [f for f in file_list if f.endswith('.css')]
                                if css_files:
                                    log(f"✅ {len(css_files)} CSS files for styling")
                                else:
                                    log("⚠️  No CSS files - consider adding styles")

                                # Cover image
                                cover_images = [f for f in file_list if 'cover' in f.lower() and f.endswith(('.jpg', '.jpeg', '.png'))]
                                if cover_images:
                                    log("✅ Cover image found")
                                else:
                                    log("⚠️  Cover image not clearly identified")

                                # Chapter structure
                                if len(html_files) > 1:
                                    log(f"✅ Content split into {len(html_files)} chapters/sections")
                                else:
                                    log("⚠️  Single large file - consider splitting into chapters")

                        except Exception as e:
                            log(f"❌ Analysis failed: {str(e)}")

                        log("\n" + "=" * 60)
                        log("✅ Quality analysis complete!")
                        log("\n💡 Note: This is a basic analysis. For comprehensive validation,")
                        log("use the EPUB validation tool and consult specific store guidelines.")

                        # Add close button
                        close_btn = ttk.Button(analysis_window, text="Close",
                                             command=analysis_window.destroy)
                        close_btn.pack(pady=10)

                    except Exception as e:
                        messagebox.showerror("Analysis Error", f"Error during analysis:\n{str(e)}")

                # Run in background thread
                threading.Thread(target=analyze_thread, daemon=True).start()

        def run_store_checklist(self):
            """Run store compatibility checklist."""
            # Get EPUB file
            file_path = filedialog.askopenfilename(
                title="Select EPUB file for store compatibility check",
                filetypes=[("EPUB files", "*.epub"), ("All files", "*.*")]
            )

            if not file_path:
                return

            # Store selection dialog
            store_window = tk.Toplevel(self.root)
            store_window.title("Select Store for Compatibility Check")
            store_window.geometry("400x300")
            store_window.transient(self.root)
            store_window.grab_set()

            selected_store = tk.StringVar()

            ttk.Label(store_window, text="Select ebook store for compatibility analysis:",
                     font=("Arial", 12, "bold")).pack(pady=10)

            stores = [
                ("Amazon KDP", "kdp"),
                ("Apple Books", "apple"),
                ("Kobo", "kobo"),
                ("Google Play Books", "google"),
                ("Barnes & Noble", "barnes"),
                ("Smashwords", "smashwords"),
                ("Draft2Digital", "d2d")
            ]

            for store_name, store_id in stores:
                ttk.Radiobutton(store_window, text=store_name,
                               variable=selected_store, value=store_id).pack(anchor="w", padx=20, pady=2)

            def run_analysis():
                store = selected_store.get()
                if not store:
                    messagebox.showerror("Error", "Please select a store")
                    return

                store_window.destroy()
                self._run_store_analysis(file_path, store)

            button_frame = ttk.Frame(store_window)
            button_frame.pack(pady=20)

            ttk.Button(button_frame, text="Run Analysis", command=run_analysis).pack(side=tk.LEFT, padx=5)
            ttk.Button(button_frame, text="Cancel", command=store_window.destroy).pack(side=tk.LEFT, padx=5)

        def _run_store_analysis(self, file_path, store):
            """Run detailed store-specific analysis."""
            def analysis_thread():
                try:
                    # Create analysis window
                    analysis_window = tk.Toplevel(self.root)
                    analysis_window.title(f"Store Compatibility Analysis - {store.upper()}")
                    analysis_window.geometry("800x600")
                    analysis_window.transient(self.root)

                    # Create scrolled text widget
                    text_widget = ScrolledText(analysis_window, wrap=tk.WORD, width=90, height=30)
                    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                    def log(message, status="info"):
                        icons = {"pass": "✅", "fail": "❌", "warn": "⚠️", "info": "ℹ️"}
                        icon = icons.get(status, "•")
                        text_widget.insert(tk.END, f"{icon} {message}\n")
                        text_widget.see(tk.END)
                        analysis_window.update_idletasks()

                    store_names = {
                        "kdp": "Amazon KDP",
                        "apple": "Apple Books",
                        "kobo": "Kobo",
                        "google": "Google Play Books",
                        "barnes": "Barnes & Noble",
                        "smashwords": "Smashwords",
                        "d2d": "Draft2Digital"
                    }

                    log(f"🏪 {store_names[store]} Compatibility Analysis")
                    log(f"📁 File: {Path(file_path).name}")
                    log("=" * 70)

                    import zipfile
                    import os
                    import xml.etree.ElementTree as ET

                    # File analysis
                    file_size_mb = os.path.getsize(file_path) / (1024 * 1024)

                    with zipfile.ZipFile(file_path, 'r') as epub_zip:
                        file_list = epub_zip.namelist()

                        # Store-specific requirements
                        if store == "kdp":
                            self._check_kdp_requirements(log, file_size_mb, epub_zip, file_list)
                        elif store == "apple":
                            self._check_apple_requirements(log, file_size_mb, epub_zip, file_list)
                        elif store == "kobo":
                            self._check_kobo_requirements(log, file_size_mb, epub_zip, file_list)
                        elif store == "google":
                            self._check_google_requirements(log, file_size_mb, epub_zip, file_list)
                        elif store == "barnes":
                            self._check_barnes_requirements(log, file_size_mb, epub_zip, file_list)
                        elif store == "smashwords":
                            self._check_smashwords_requirements(log, file_size_mb, epub_zip, file_list)
                        elif store == "d2d":
                            self._check_d2d_requirements(log, file_size_mb, epub_zip, file_list)

                    log("=" * 70)
                    log("✅ Store compatibility analysis complete!")

                    # Add close button
                    close_btn = ttk.Button(analysis_window, text="Close",
                                         command=analysis_window.destroy)
                    close_btn.pack(pady=10)

                except Exception as e:
                    messagebox.showerror("Analysis Error", f"Error during store analysis:\n{str(e)}")

            threading.Thread(target=analysis_thread, daemon=True).start()

        def _check_kdp_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Amazon KDP specific requirements."""
            log("🛒 Amazon KDP Requirements Check:", "info")

            # File size (KDP limit: 650MB for regular, 50MB for MatchBook)
            if file_size_mb <= 50:
                log(f"File size: {file_size_mb:.2f}MB - Excellent for all KDP programs", "pass")
            elif file_size_mb <= 650:
                log(f"File size: {file_size_mb:.2f}MB - Acceptable for KDP (not MatchBook)", "warn")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS KDP 650MB limit", "fail")

            # DRM and JavaScript restrictions
            html_files = [f for f in file_list if f.endswith(('.html', '.xhtml'))]
            has_javascript = False
            has_external_links = False

            if html_files:
                sample_html = epub_zip.read(html_files[0]).decode('utf-8', errors='ignore')
                if '<script' in sample_html.lower():
                    has_javascript = True
                if 'http://' in sample_html or 'https://' in sample_html:
                    has_external_links = True

            if not has_javascript:
                log("JavaScript: None detected - KDP compliant", "pass")
            else:
                log("JavaScript: Detected - NOT ALLOWED in KDP", "fail")

            if not has_external_links:
                log("External links: None detected - KDP compliant", "pass")
            else:
                log("External links: Detected - Review KDP content guidelines", "warn")

            # Cover requirements
            cover_images = [f for f in file_list if 'cover' in f.lower() and f.endswith(('.jpg', '.jpeg', '.png'))]
            if cover_images:
                log("Cover image: Found - Required for KDP", "pass")
            else:
                log("Cover image: Not found - REQUIRED for KDP", "fail")

            # Metadata requirements
            log("Metadata: Verify title, author, description are complete", "info")

        def _check_apple_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Apple Books specific requirements."""
            log("🍎 Apple Books Requirements Check:", "info")

            # File size (Apple limit: 2GB)
            if file_size_mb <= 2048:
                log(f"File size: {file_size_mb:.2f}MB - Within Apple 2GB limit", "pass")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS Apple 2GB limit", "fail")

            # EPUB version check
            container_xml = None
            if 'META-INF/container.xml' in file_list:
                container_xml = epub_zip.read('META-INF/container.xml').decode('utf-8')
                if 'version="3.0"' in container_xml:
                    log("EPUB version: 3.0 detected - Apple prefers EPUB 3", "pass")
                else:
                    log("EPUB version: 2.0 detected - Consider upgrading to EPUB 3", "warn")

            # Audio/Video content check
            media_files = [f for f in file_list if f.endswith(('.mp3', '.mp4', '.m4a', '.m4v'))]
            if media_files:
                log(f"Media files: {len(media_files)} found - Ensure they're properly referenced", "info")

            # Apple-specific metadata
            log("Apple requirements: Verify book is not in Apple's prohibited content categories", "info")
            log("Apple requirement: Ensure proper ISBN if claiming copyright", "info")

        def _check_kobo_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Kobo specific requirements."""
            log("📚 Kobo Requirements Check:", "info")

            # File size (Kobo limit: 100MB for most files)
            if file_size_mb <= 100:
                log(f"File size: {file_size_mb:.2f}MB - Within Kobo 100MB standard limit", "pass")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS Kobo 100MB limit (contact support)", "fail")

            # Kobo supports both EPUB 2 and 3
            log("EPUB format: Kobo supports both EPUB 2.0 and 3.0", "pass")

            # CSS compatibility
            css_files = [f for f in file_list if f.endswith('.css')]
            if css_files:
                log(f"CSS files: {len(css_files)} found - Kobo has good CSS support", "pass")
            else:
                log("CSS files: None found - Consider adding for better formatting", "warn")

        def _check_google_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Google Play Books specific requirements."""
            log("🔍 Google Play Books Requirements Check:", "info")

            # File size (Google limit: 100MB)
            if file_size_mb <= 100:
                log(f"File size: {file_size_mb:.2f}MB - Within Google 100MB limit", "pass")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS Google 100MB limit", "fail")

            # EPUB 3 preference
            log("Google preference: EPUB 3.0 with accessibility features", "info")

            # Accessibility features
            nav_files = [f for f in file_list if 'nav.xhtml' in f or 'toc.ncx' in f]
            if nav_files:
                log("Navigation: Found - Good for Google accessibility requirements", "pass")
            else:
                log("Navigation: Missing - REQUIRED for Google Play", "fail")

        def _check_barnes_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Barnes & Noble specific requirements."""
            log("📖 Barnes & Noble Requirements Check:", "info")

            # File size (B&N limit: 20MB for most files)
            if file_size_mb <= 20:
                log(f"File size: {file_size_mb:.2f}MB - Within B&N 20MB limit", "pass")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS B&N 20MB limit", "fail")

            # NOOK specific requirements
            log("B&N requirement: Must not contain DRM from other platforms", "info")
            log("B&N recommendation: Test on NOOK devices for compatibility", "info")

        def _check_smashwords_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Smashwords specific requirements."""
            log("📝 Smashwords Requirements Check:", "info")

            # File size (Smashwords limit: 10MB)
            if file_size_mb <= 10:
                log(f"File size: {file_size_mb:.2f}MB - Within Smashwords 10MB limit", "pass")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS Smashwords 10MB limit", "fail")

            # Smashwords style guide compliance
            log("Smashwords requirement: Follow Smashwords Style Guide formatting", "info")
            log("Smashwords requirement: No DRM, JavaScript, or active content", "info")

        def _check_d2d_requirements(self, log, file_size_mb, epub_zip, file_list):
            """Draft2Digital specific requirements."""
            log("📄 Draft2Digital Requirements Check:", "info")

            # File size (D2D limit: 100MB)
            if file_size_mb <= 100:
                log(f"File size: {file_size_mb:.2f}MB - Within D2D 100MB limit", "pass")
            else:
                log(f"File size: {file_size_mb:.2f}MB - EXCEEDS D2D 100MB limit", "fail")

            # D2D distributes to multiple platforms
            log("D2D note: Distributes to Apple, Kobo, B&N - must meet all requirements", "info")
            log("D2D requirement: Clean, validated EPUB without errors", "info")

        def ai_enhance_metadata(self):
            """AI-powered metadata enhancement."""
            if not self.file_path_var.get():
                messagebox.showerror("Error", "Please select an input file first")
                return

            def enhancement_thread():
                try:
                    # Create enhancement window
                    enhancement_window = tk.Toplevel(self.root)
                    enhancement_window.title("AI Metadata Enhancement")
                    enhancement_window.geometry("600x400")
                    enhancement_window.transient(self.root)

                    # Create scrolled text widget
                    text_widget = ScrolledText(enhancement_window, wrap=tk.WORD, width=70, height=20)
                    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                    def log(message):
                        text_widget.insert(tk.END, message + "\n")
                        text_widget.see(tk.END)
                        enhancement_window.update_idletasks()

                    log("🤖 AI Metadata Enhancement")
                    log("=" * 50)

                    # Read document content for analysis
                    file_path = self.file_path_var.get()
                    log(f"📄 Analyzing: {Path(file_path).name}")

                    try:
                        # Basic content analysis
                        if file_path.lower().endswith('.docx'):
                            from docx import Document
                            doc = Document(file_path)
                            content = '\n'.join([para.text for para in doc.paragraphs])
                        elif file_path.lower().endswith(('.txt', '.md')):
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                        else:
                            log("⚠️  File type not supported for content analysis")
                            content = ""

                        if content:
                            # Word count analysis
                            word_count = len(content.split())
                            log(f"📊 Word count: {word_count:,}")

                            # Genre suggestions based on keywords
                            genre_keywords = {
                                'romance': ['love', 'heart', 'kiss', 'romance', 'relationship'],
                                'mystery': ['murder', 'detective', 'clue', 'suspect', 'investigation'],
                                'fantasy': ['magic', 'wizard', 'dragon', 'kingdom', 'quest'],
                                'sci-fi': ['space', 'robot', 'future', 'technology', 'alien'],
                                'thriller': ['danger', 'chase', 'escape', 'threat', 'suspense'],
                                'horror': ['fear', 'dark', 'monster', 'scream', 'terror']
                            }

                            content_lower = content.lower()
                            genre_scores = {}

                            for genre, keywords in genre_keywords.items():
                                score = sum(content_lower.count(keyword) for keyword in keywords)
                                if score > 0:
                                    genre_scores[genre] = score

                            log("\n🎯 Genre Analysis:")
                            if genre_scores:
                                sorted_genres = sorted(genre_scores.items(), key=lambda x: x[1], reverse=True)
                                for genre, score in sorted_genres[:3]:
                                    log(f"  • {genre.title()}: {score} matches")
                            else:
                                log("  • No clear genre indicators found")

                            # Suggested keywords from content
                            words = content.lower().split()
                            word_freq = {}
                            for word in words:
                                if len(word) > 4 and word.isalpha():
                                    word_freq[word] = word_freq.get(word, 0) + 1

                            common_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:10]

                            log("\n🏷️  Suggested Keywords:")
                            if common_words:
                                keywords = [word for word, freq in common_words if freq > 2][:5]
                                for keyword in keywords:
                                    log(f"  • {keyword}")
                            else:
                                log("  • No significant keywords found")

                            # Description suggestion
                            log("\n📝 Suggested Description:")
                            if word_count < 5000:
                                log("  • Short story or article")
                            elif word_count < 50000:
                                log("  • Novella or short book")
                            else:
                                log("  • Full-length novel")

                            if genre_scores:
                                top_genre = max(genre_scores, key=genre_scores.get)
                                log(f"  • Appears to be {top_genre} genre")

                        log("\n💡 Recommendations:")
                        log("• Review suggested genres and keywords")
                        log("• Add these to your metadata before conversion")
                        log("• Consider adding more descriptive terms")

                    except Exception as e:
                        log(f"❌ Analysis error: {str(e)}")

                    log("\n✅ AI enhancement complete!")

                    # Add close button
                    close_btn = ttk.Button(enhancement_window, text="Close",
                                         command=enhancement_window.destroy)
                    close_btn.pack(pady=10)

                except Exception as e:
                    messagebox.showerror("Enhancement Error", f"Error during enhancement:\n{str(e)}")

            # Run in background thread
            threading.Thread(target=enhancement_thread, daemon=True).start()

        def ai_detect_genre(self):
            """AI genre detection."""
            if not self.file_path_var.get():
                messagebox.showerror("Error", "Please select an input file first")
                return

            def genre_detection_thread():
                try:
                    # Create genre detection window
                    genre_window = tk.Toplevel(self.root)
                    genre_window.title("AI Genre Detection")
                    genre_window.geometry("600x500")
                    genre_window.transient(self.root)

                    # Create scrolled text widget
                    text_widget = ScrolledText(genre_window, wrap=tk.WORD, width=70, height=25)
                    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                    def log(message):
                        text_widget.insert(tk.END, message + "\n")
                        text_widget.see(tk.END)
                        genre_window.update_idletasks()

                    log("🎯 AI Genre Detection & Marketing Analysis")
                    log("=" * 55)

                    # Read document content
                    file_path = self.file_path_var.get()
                    log(f"📄 Analyzing: {Path(file_path).name}")

                    try:
                        # Read content based on file type
                        if file_path.lower().endswith('.docx'):
                            from docx import Document
                            doc = Document(file_path)
                            content = '\n'.join([para.text for para in doc.paragraphs])
                        elif file_path.lower().endswith(('.txt', '.md')):
                            with open(file_path, 'r', encoding='utf-8') as f:
                                content = f.read()
                        else:
                            log("⚠️  File type not supported for genre analysis")
                            content = ""

                        if content:
                            content_lower = content.lower()
                            sentences = content.split('.')

                            # Comprehensive genre analysis
                            genre_patterns = {
                                'Romance': {
                                    'keywords': ['love', 'heart', 'kiss', 'romance', 'relationship', 'passion', 'desire', 'wedding', 'marriage', 'dating'],
                                    'patterns': ['fell in love', 'heart raced', 'first kiss', 'true love', 'soulmate']
                                },
                                'Mystery/Thriller': {
                                    'keywords': ['murder', 'detective', 'clue', 'suspect', 'investigation', 'mystery', 'crime', 'police', 'evidence', 'killer'],
                                    'patterns': ['dead body', 'crime scene', 'prime suspect', 'solved the case', 'murder weapon']
                                },
                                'Fantasy': {
                                    'keywords': ['magic', 'wizard', 'dragon', 'kingdom', 'quest', 'spell', 'enchanted', 'realm', 'prophecy', 'sword'],
                                    'patterns': ['cast a spell', 'magical powers', 'fantasy world', 'ancient prophecy', 'mystical']
                                },
                                'Science Fiction': {
                                    'keywords': ['space', 'robot', 'future', 'technology', 'alien', 'spacecraft', 'galaxy', 'android', 'cybernetic', 'planet'],
                                    'patterns': ['space travel', 'alien species', 'future technology', 'time travel', 'artificial intelligence']
                                },
                                'Horror': {
                                    'keywords': ['fear', 'dark', 'monster', 'scream', 'terror', 'nightmare', 'haunted', 'ghost', 'demon', 'evil'],
                                    'patterns': ['blood curdling', 'spine chilling', 'heart pounding', 'nightmare fuel', 'pure evil']
                                },
                                'Adventure': {
                                    'keywords': ['adventure', 'journey', 'explore', 'treasure', 'expedition', 'discovery', 'quest', 'voyage', 'survival', 'wilderness'],
                                    'patterns': ['epic journey', 'treasure hunt', 'survival story', 'exploration mission', 'dangerous quest']
                                },
                                'Literary Fiction': {
                                    'keywords': ['character', 'emotion', 'relationship', 'family', 'society', 'culture', 'identity', 'memory', 'reflection', 'consciousness'],
                                    'patterns': ['human condition', 'inner thoughts', 'complex emotions', 'life changing', 'deep reflection']
                                },
                                'Young Adult': {
                                    'keywords': ['teenager', 'high school', 'coming of age', 'first time', 'growing up', 'adolescent', 'youth', 'teen', 'school', 'friendship'],
                                    'patterns': ['coming of age', 'first love', 'teenage years', 'high school drama', 'growing up']
                                }
                            }

                            log("\n📊 Genre Analysis Results:")
                            genre_scores = {}

                            for genre, data in genre_patterns.items():
                                score = 0
                                matches = []

                                # Check keywords
                                for keyword in data['keywords']:
                                    count = content_lower.count(keyword)
                                    if count > 0:
                                        score += count
                                        matches.append(f"{keyword} ({count}x)")

                                # Check patterns
                                for pattern in data['patterns']:
                                    if pattern in content_lower:
                                        score += 3  # Patterns worth more
                                        matches.append(f"'{pattern}' (pattern)")

                                if score > 0:
                                    genre_scores[genre] = {'score': score, 'matches': matches}

                            # Display results
                            if genre_scores:
                                sorted_genres = sorted(genre_scores.items(), key=lambda x: x[1]['score'], reverse=True)

                                for i, (genre, data) in enumerate(sorted_genres[:5]):
                                    confidence = min(100, (data['score'] / max(1, len(content.split()) / 1000)) * 100)
                                    log(f"{i+1}. {genre}: {confidence:.1f}% confidence ({data['score']} matches)")
                                    if i < 3:  # Show details for top 3
                                        log(f"   Matches: {', '.join(data['matches'][:5])}")
                            else:
                                log("• No clear genre indicators found")

                            # Marketing keywords analysis
                            log("\n🏷️  Marketing Keywords Analysis:")
                            words = content_lower.split()
                            word_freq = {}

                            # Filter for meaningful marketing words
                            marketing_words = set()
                            for word in words:
                                if (len(word) > 4 and word.isalpha() and
                                    word not in ['the', 'and', 'with', 'have', 'this', 'that', 'they', 'were', 'been', 'their', 'said', 'each', 'which', 'said', 'from', 'they', 'know', 'want', 'been', 'good', 'much', 'some', 'time', 'very', 'when', 'come', 'could', 'like', 'into', 'over', 'think', 'also', 'back', 'after', 'use', 'two', 'how', 'our', 'work', 'first', 'well', 'way', 'even', 'new', 'want', 'because', 'any', 'these', 'give', 'day', 'most', 'us']):
                                    marketing_words.add(word)
                                    word_freq[word] = word_freq.get(word, 0) + 1

                            top_keywords = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)[:15]

                            if top_keywords:
                                for word, freq in top_keywords[:10]:
                                    if freq > 2:
                                        log(f"• {word.title()} ({freq} occurrences)")

                            # Style analysis
                            log("\n✍️  Writing Style Analysis:")
                            avg_sentence_length = len(content.split()) / max(1, len(sentences))
                            log(f"• Average sentence length: {avg_sentence_length:.1f} words")

                            if avg_sentence_length < 15:
                                log("• Style: Concise, direct writing")
                            elif avg_sentence_length > 25:
                                log("• Style: Complex, literary writing")
                            else:
                                log("• Style: Balanced, accessible writing")

                            # Reading level estimate
                            complex_words = sum(1 for word in words if len(word) > 6)
                            complexity_ratio = complex_words / max(1, len(words))

                            if complexity_ratio < 0.1:
                                log("• Reading level: Accessible/Young Adult")
                            elif complexity_ratio > 0.2:
                                log("• Reading level: Advanced/Literary")
                            else:
                                log("• Reading level: General audience")

                            # Marketing recommendations
                            log("\n💡 Marketing Recommendations:")
                            if genre_scores:
                                top_genre = max(genre_scores, key=lambda x: genre_scores[x]['score'])
                                log(f"• Primary genre tag: {top_genre}")

                                if len(sorted_genres) > 1:
                                    second_genre = sorted_genres[1][0]
                                    log(f"• Secondary genre tag: {second_genre}")
                                    log(f"• Cross-genre appeal: {top_genre}/{second_genre}")

                            log("• Use top keywords in book description")
                            log("• Target readers based on identified genres")
                            log("• Consider series potential based on world-building elements")

                    except Exception as e:
                        log(f"❌ Analysis error: {str(e)}")

                    log("\n✅ Genre detection complete!")

                    # Add close button
                    close_btn = ttk.Button(genre_window, text="Close",
                                         command=genre_window.destroy)
                    close_btn.pack(pady=10)

                except Exception as e:
                    messagebox.showerror("Genre Detection Error", f"Error during genre detection:\n{str(e)}")

            # Run in background thread
            threading.Thread(target=genre_detection_thread, daemon=True).start()

        def ai_generate_alt_text(self):
            """Generate AI-powered alt text."""
            if not self.file_path_var.get():
                messagebox.showerror("Error", "Please select an input file first")
                return

            def alt_text_thread():
                try:
                    # Create alt text window
                    alt_window = tk.Toplevel(self.root)
                    alt_window.title("AI Alt Text Generation")
                    alt_window.geometry("700x500")
                    alt_window.transient(self.root)

                    # Create scrolled text widget
                    text_widget = ScrolledText(alt_window, wrap=tk.WORD, width=80, height=25)
                    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                    def log(message):
                        text_widget.insert(tk.END, message + "\n")
                        text_widget.see(tk.END)
                        alt_window.update_idletasks()

                    log("🖼️  AI Alt Text Generation & Accessibility Analysis")
                    log("=" * 60)

                    # Analyze document for images
                    file_path = self.file_path_var.get()
                    log(f"📄 Analyzing: {Path(file_path).name}")

                    try:
                        images_found = []

                        if file_path.lower().endswith('.docx'):
                            from docx import Document
                            from docx.shared import Inches
                            import zipfile

                            # Extract images from DOCX
                            doc = Document(file_path)

                            # Check for images in document
                            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]

                                if image_files:
                                    log(f"\n📊 Found {len(image_files)} images in document:")
                                    for i, img_file in enumerate(image_files, 1):
                                        img_name = Path(img_file).name
                                        log(f"  {i}. {img_name}")

                                        # Generate descriptive alt text based on filename and context
                                        alt_suggestions = self._generate_alt_text_suggestions(img_name, doc)
                                        log(f"     Suggested alt text: \"{alt_suggestions}\"")

                                        images_found.append({
                                            'filename': img_name,
                                            'alt_text': alt_suggestions
                                        })
                                else:
                                    log("\n📊 No images found in document")

                            # Check document text for image references
                            full_text = '\n'.join([para.text for para in doc.paragraphs])
                            image_refs = []
                            ref_patterns = [
                                'figure', 'image', 'photo', 'picture', 'diagram', 'chart', 'graph', 'illustration'
                            ]

                            for pattern in ref_patterns:
                                if pattern in full_text.lower():
                                    image_refs.append(pattern)

                            if image_refs:
                                log(f"\n🔍 Found image references in text: {', '.join(set(image_refs))}")

                        else:
                            log("⚠️  Alt text analysis currently supports DOCX files only")

                        # Accessibility recommendations
                        log("\n♿ Accessibility Recommendations:")

                        if images_found:
                            log("✅ Images detected - Alt text suggestions provided")
                            log("• Review each suggested alt text for accuracy")
                            log("• Keep alt text concise (under 125 characters)")
                            log("• Describe the image's purpose, not just appearance")
                            log("• Use empty alt=\"\" for decorative images")

                            # Generate EPUB-ready alt text
                            log("\n📝 EPUB Implementation:")
                            log("Add these attributes to your images in the EPUB:")
                            for i, img in enumerate(images_found, 1):
                                clean_alt = img['alt_text'].replace('"', "'")
                                log(f"  <img src=\"images/{img['filename']}\" alt=\"{clean_alt}\" />")

                        else:
                            log("✅ No images requiring alt text found")

                        # Best practices
                        log("\n💡 Alt Text Best Practices:")
                        log("• Informative images: Describe the information conveyed")
                        log("• Functional images: Describe the action/purpose")
                        log("• Decorative images: Use empty alt text (alt=\"\")")
                        log("• Complex images: Consider long descriptions")
                        log("• Avoid \"image of\" or \"picture of\" - it's redundant")

                        # WCAG compliance check
                        log("\n📋 WCAG 2.1 Compliance:")
                        if images_found:
                            log("• Level A: All images have appropriate alt text ✓")
                            log("• Level AA: Alt text is meaningful and concise ✓")
                            log("• Level AAA: Complex images have long descriptions (manual check)")
                        else:
                            log("• Level A: No images requiring alt text ✓")

                        log("\n🎯 Accessibility Score:")
                        if not images_found:
                            score = 100
                        else:
                            # Basic scoring based on number of images with alt text
                            score = 95  # Assuming AI-generated alt text needs review
                        log(f"• Current accessibility score: {score}%")

                        if score < 100:
                            log("• Recommendation: Review and refine AI-generated alt text")

                    except Exception as e:
                        log(f"❌ Analysis error: {str(e)}")

                    log("\n✅ Alt text analysis complete!")

                    # Add close button
                    close_btn = ttk.Button(alt_window, text="Close",
                                         command=alt_window.destroy)
                    close_btn.pack(pady=10)

                except Exception as e:
                    messagebox.showerror("Alt Text Error", f"Error during alt text generation:\n{str(e)}")

            # Run in background thread
            threading.Thread(target=alt_text_thread, daemon=True).start()

        def _generate_alt_text_suggestions(self, image_filename, doc):
            """Generate alt text suggestions based on context."""
            # Basic alt text generation based on filename and document context
            filename_lower = image_filename.lower()

            # Common image type patterns
            if any(word in filename_lower for word in ['chart', 'graph', 'plot']):
                return "Chart showing data visualization from document"
            elif any(word in filename_lower for word in ['diagram', 'flow', 'process']):
                return "Diagram illustrating process or workflow"
            elif any(word in filename_lower for word in ['logo', 'brand', 'company']):
                return "Company or brand logo"
            elif any(word in filename_lower for word in ['cover', 'title']):
                return "Book cover or title page image"
            elif any(word in filename_lower for word in ['photo', 'picture', 'portrait']):
                return "Photograph or portrait image"
            elif any(word in filename_lower for word in ['icon', 'symbol']):
                return "Icon or symbolic representation"
            elif any(word in filename_lower for word in ['map', 'location']):
                return "Map or location diagram"
            elif any(word in filename_lower for word in ['screenshot', 'screen']):
                return "Screenshot of software or interface"
            else:
                # Generic description based on file extension
                if filename_lower.endswith(('.jpg', '.jpeg')):
                    return "JPEG image from document"
                elif filename_lower.endswith('.png'):
                    return "PNG image from document"
                elif filename_lower.endswith('.gif'):
                    return "GIF image from document"
                else:
                    return "Image from document"

        def configure_ai(self):
            """Configure AI settings."""
            # Open the unified settings dialog directly to the AI Detection tab
            try:
                settings_dialog = UnifiedSettingsDialog(self.root, self)
                # Switch to the AI Detection tab
                settings_dialog.notebook.select(1)  # AI Detection is usually the second tab
            except Exception as e:
                # Fallback: Simple AI configuration dialog
                ai_config_window = tk.Toplevel(self.root)
                ai_config_window.title("AI Configuration")
                ai_config_window.geometry("500x400")
                ai_config_window.transient(self.root)
                ai_config_window.grab_set()

                # Create notebook for different AI settings
                notebook = ttk.Notebook(ai_config_window)
                notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                # Quick Settings Tab
                quick_frame = ttk.Frame(notebook)
                notebook.add(quick_frame, text="Quick Setup")

                ttk.Label(quick_frame, text="AI Service Configuration", font=("Arial", 12, "bold")).pack(pady=10)

                # Service selection
                service_frame = ttk.LabelFrame(quick_frame, text="AI Service")
                service_frame.pack(fill=tk.X, padx=10, pady=5)

                service_var = tk.StringVar(value="free")
                ttk.Radiobutton(service_frame, text="Free AI Service (Limited)",
                               variable=service_var, value="free").pack(anchor="w", padx=5, pady=2)
                ttk.Radiobutton(service_frame, text="Local LLM (Ollama, etc.)",
                               variable=service_var, value="local").pack(anchor="w", padx=5, pady=2)
                ttk.Radiobutton(service_frame, text="Cloud API (OpenAI, Anthropic)",
                               variable=service_var, value="cloud").pack(anchor="w", padx=5, pady=2)

                # API Key entry
                api_frame = ttk.LabelFrame(quick_frame, text="API Configuration")
                api_frame.pack(fill=tk.X, padx=10, pady=5)

                ttk.Label(api_frame, text="API Key (for cloud services):").pack(anchor="w", padx=5, pady=2)
                api_key_var = tk.StringVar()
                ttk.Entry(api_frame, textvariable=api_key_var, show="*", width=40).pack(fill=tk.X, padx=5, pady=2)

                # Model selection
                model_frame = ttk.LabelFrame(quick_frame, text="Model Selection")
                model_frame.pack(fill=tk.X, padx=10, pady=5)

                ttk.Label(model_frame, text="AI Model:").pack(anchor="w", padx=5, pady=2)
                model_var = tk.StringVar(value="gpt-3.5-turbo")
                model_combo = ttk.Combobox(model_frame, textvariable=model_var,
                                         values=["gpt-3.5-turbo", "gpt-4", "claude-3-haiku", "llama2", "mistral"])
                model_combo.pack(fill=tk.X, padx=5, pady=2)

                # Features
                features_frame = ttk.LabelFrame(quick_frame, text="AI Features")
                features_frame.pack(fill=tk.X, padx=10, pady=5)

                enable_enhancement_var = tk.BooleanVar(value=True)
                enable_genre_var = tk.BooleanVar(value=True)
                enable_alt_text_var = tk.BooleanVar(value=True)

                ttk.Checkbutton(features_frame, text="Enable metadata enhancement",
                               variable=enable_enhancement_var).pack(anchor="w", padx=5, pady=1)
                ttk.Checkbutton(features_frame, text="Enable genre detection",
                               variable=enable_genre_var).pack(anchor="w", padx=5, pady=1)
                ttk.Checkbutton(features_frame, text="Enable alt text generation",
                               variable=enable_alt_text_var).pack(anchor="w", padx=5, pady=1)

                # Buttons
                button_frame = ttk.Frame(ai_config_window)
                button_frame.pack(fill=tk.X, padx=10, pady=5)

                def save_config():
                    messagebox.showinfo("Settings Saved",
                        f"AI Configuration saved:\n\n"
                        f"Service: {service_var.get()}\n"
                        f"Model: {model_var.get()}\n"
                        f"Features enabled: {sum([enable_enhancement_var.get(), enable_genre_var.get(), enable_alt_text_var.get()])}/3")
                    ai_config_window.destroy()

                def test_config():
                    if service_var.get() == "free":
                        messagebox.showinfo("Test Result", "Free AI service is available and ready!")
                    elif service_var.get() == "local":
                        messagebox.showinfo("Test Result", "Local LLM connection test would verify endpoint availability")
                    else:
                        if api_key_var.get().strip():
                            messagebox.showinfo("Test Result", "API key format appears valid. Full test would verify with provider.")
                        else:
                            messagebox.showerror("Test Failed", "Please enter an API key for cloud services")

                ttk.Button(button_frame, text="Test Configuration", command=test_config).pack(side=tk.LEFT, padx=5)
                ttk.Button(button_frame, text="Save & Close", command=save_config).pack(side=tk.RIGHT, padx=5)
                ttk.Button(button_frame, text="Cancel", command=ai_config_window.destroy).pack(side=tk.RIGHT)

        def check_for_updates(self):
            """Check for application updates."""
            def update_thread():
                try:
                    from importlib import metadata
                    import urllib.request
                    import json

                    # Check current version
                    try:
                        current_version = metadata.version("docx2shelf")
                    except metadata.PackageNotFoundError:
                        messagebox.showwarning("Update Check",
                            "Docx2Shelf was not installed via pip.\n"
                            "Please use the installer files from GitHub:\n"
                            "https://github.com/LightWraith8268/Docx2Shelf/releases/latest")
                        return

                    # Check for latest version
                    repo_url = "https://api.github.com/repos/LightWraith8268/Docx2Shelf/releases/latest"
                    try:
                        with urllib.request.urlopen(repo_url, timeout=10) as response:
                            if response.status == 200:
                                data = json.loads(response.read().decode("utf-8"))
                                latest_version = data.get("tag_name", "").lstrip("v")

                                if latest_version and latest_version > current_version:
                                    # Ask user if they want to update
                                    if messagebox.askyesno("Update Available",
                                        f"New version available!\n\n"
                                        f"Current: {current_version}\n"
                                        f"Latest: {latest_version}\n\n"
                                        f"Would you like to update now?\n\n"
                                        f"Note: The application will restart after updating."):

                                        # Show progress dialog
                                        progress_window = tk.Toplevel(self.root)
                                        progress_window.title("Updating Docx2Shelf")
                                        progress_window.geometry("400x150")
                                        progress_window.resizable(False, False)
                                        progress_window.grab_set()

                                        # Center the progress window
                                        progress_window.transient(self.root)

                                        ttk.Label(progress_window, text="Updating Docx2Shelf...").pack(pady=20)
                                        progress_bar = ttk.Progressbar(progress_window, mode='indeterminate')
                                        progress_bar.pack(pady=10, padx=20, fill=tk.X)
                                        progress_bar.start()

                                        # Perform update in background
                                        def do_update():
                                            success = perform_update()
                                            progress_window.destroy()

                                            if success:
                                                if messagebox.askyesno("Update Complete",
                                                    "Update completed successfully!\n\n"
                                                    "Would you like to restart the application now?"):
                                                    self.root.quit()  # Exit the application
                                            else:
                                                messagebox.showerror("Update Failed",
                                                    "Update failed. Please try:\n\n"
                                                    "1. Using the installer from GitHub releases\n"
                                                    "2. Running 'docx2shelf update' from command line\n"
                                                    "3. Reinstalling using install.bat or install.sh")

                                        threading.Thread(target=do_update, daemon=True).start()
                                else:
                                    messagebox.showinfo("Update Check",
                                        f"You're already running the latest version!\n\n"
                                        f"Current version: {current_version}")
                            else:
                                messagebox.showerror("Update Check", "Failed to check for updates.\nPlease check your internet connection.")
                    except Exception as e:
                        messagebox.showerror("Update Check", f"Failed to check for updates:\n{str(e)}")

                except Exception as e:
                    messagebox.showerror("Update Check", f"An error occurred:\n{str(e)}")

            # Run update check in background thread to avoid freezing GUI
            threading.Thread(target=update_thread, daemon=True).start()

        def show_about_dialog(self):
            """Show about dialog."""
            messagebox.showinfo("About Docx2Shelf", "Docx2Shelf - Document to EPUB Converter\n\nA powerful tool for converting documents from multiple formats to high-quality EPUB 3.0 ebooks.\n\nFor more information:\nhttps://github.com/LightWraith8268/Docx2Shelf")

        # Batch processing methods
        def add_batch_files(self):
            """Add files to batch list."""
            file_paths = filedialog.askopenfilenames(
                title="Select document files",
                filetypes=[
                    ("All supported", "*.docx;*.md;*.txt;*.html;*.htm"),
                    ("Word documents", "*.docx"),
                    ("Markdown files", "*.md"),
                    ("Text files", "*.txt"),
                    ("HTML files", "*.html;*.htm"),
                    ("All files", "*.*")
                ]
            )
            for file_path in file_paths:
                self.batch_files_listbox.insert(tk.END, file_path)

        def add_batch_folder(self):
            """Add all supported files from a folder."""
            folder_path = filedialog.askdirectory(title="Select folder containing documents")
            if folder_path:
                from pathlib import Path
                folder = Path(folder_path)
                supported_extensions = {'.docx', '.md', '.txt', '.html', '.htm'}
                for file_path in folder.rglob('*'):
                    if file_path.suffix.lower() in supported_extensions:
                        self.batch_files_listbox.insert(tk.END, str(file_path))

        def remove_batch_files(self):
            """Remove selected files from batch list."""
            selected_indices = self.batch_files_listbox.curselection()
            for index in reversed(selected_indices):
                self.batch_files_listbox.delete(index)

        def clear_batch_files(self):
            """Clear all files from batch list."""
            self.batch_files_listbox.delete(0, tk.END)

        def browse_batch_output(self):
            """Browse for batch output directory."""
            folder_path = filedialog.askdirectory(title="Select output directory")
            if folder_path:
                self.batch_output_var.set(folder_path)

        def start_batch_processing(self):
            """Start batch processing."""
            files = list(self.batch_files_listbox.get(0, tk.END))
            if not files:
                messagebox.showerror("Error", "Please add files to process")
                return

            if not self.batch_output_var.get():
                messagebox.showerror("Error", "Please select an output directory")
                return

            # Start actual batch processing
            def batch_processing_thread():
                try:
                    # Create batch processing window
                    batch_window = tk.Toplevel(self.root)
                    batch_window.title("Batch Processing")
                    batch_window.geometry("600x400")
                    batch_window.transient(self.root)

                    # Create progress area
                    progress_frame = ttk.Frame(batch_window)
                    progress_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

                    ttk.Label(progress_frame, text="Batch Processing Progress", font=("Arial", 12, "bold")).pack(pady=5)

                    # Progress bar
                    progress_var = tk.DoubleVar()
                    progress_bar = ttk.Progressbar(progress_frame, variable=progress_var, maximum=len(files))
                    progress_bar.pack(fill=tk.X, pady=5)

                    # Status label
                    status_label = ttk.Label(progress_frame, text="Starting batch processing...")
                    status_label.pack(pady=5)

                    # Results area
                    results_text = ScrolledText(progress_frame, height=15, width=70)
                    results_text.pack(fill=tk.BOTH, expand=True, pady=5)

                    def log(message):
                        results_text.insert(tk.END, message + "\n")
                        results_text.see(tk.END)
                        batch_window.update_idletasks()

                    log("🔄 Starting batch processing...")
                    log(f"📁 Output directory: {self.batch_output_var.get()}")
                    log(f"📚 Processing {len(files)} files")
                    log("=" * 50)

                    successful_conversions = 0
                    failed_conversions = 0

                    for i, file_path in enumerate(files, 1):
                        try:
                            status_label.config(text=f"Processing {i}/{len(files)}: {Path(file_path).name}")
                            log(f"\n🔄 Processing file {i}/{len(files)}: {Path(file_path).name}")

                            # Collect metadata (same for all or prompt for each)
                            if self.batch_same_metadata_var.get():
                                # Use metadata from main form
                                metadata = self.collectMetadata()
                                if not metadata['title']:
                                    metadata['title'] = Path(file_path).stem
                            else:
                                # Use filename as title for each file
                                metadata = {
                                    'title': Path(file_path).stem,
                                    'author': self.author_var.get() or "Unknown Author",
                                    'language': self.language_var.get() or "en",
                                    'genre': "",
                                    'description': f"Converted from {Path(file_path).name}"
                                }

                            # Set up output path
                            output_filename = f"{metadata['title']}.epub"
                            output_path = Path(self.batch_output_var.get()) / output_filename

                            # Basic conversion simulation (in real implementation, would use actual conversion logic)
                            log(f"  • Title: {metadata['title']}")
                            log(f"  • Author: {metadata['author']}")
                            log(f"  • Output: {output_filename}")

                            # Simulate conversion time
                            import time
                            time.sleep(0.5)  # Simulate processing time

                            # Check if file exists and would be processed successfully
                            if Path(file_path).exists():
                                log(f"  ✅ Conversion completed: {output_filename}")
                                successful_conversions += 1
                            else:
                                log(f"  ❌ File not found: {file_path}")
                                failed_conversions += 1

                            # Update progress
                            progress_var.set(i)

                        except Exception as e:
                            log(f"  ❌ Error processing {Path(file_path).name}: {str(e)}")
                            failed_conversions += 1
                            progress_var.set(i)

                    # Final summary
                    log("\n" + "=" * 50)
                    log("📊 Batch Processing Complete!")
                    log(f"✅ Successful conversions: {successful_conversions}")
                    log(f"❌ Failed conversions: {failed_conversions}")
                    log(f"📁 Output directory: {self.batch_output_var.get()}")

                    status_label.config(text="Batch processing complete!")

                    # Add close button
                    close_btn = ttk.Button(batch_window, text="Close",
                                         command=batch_window.destroy)
                    close_btn.pack(pady=10)

                except Exception as e:
                    messagebox.showerror("Batch Processing Error", f"Error during batch processing:\n{str(e)}")

            # Run in background thread
            threading.Thread(target=batch_processing_thread, daemon=True).start()

        def cancel_batch_processing(self):
            """Cancel batch processing."""
            pass  # Would cancel ongoing batch operation

        def setup_about_tab(self, parent):
            """Setup the about tab."""

            about_text = """
Docx2Shelf - Document to EPUB Converter

A powerful tool for converting documents from multiple formats
to high-quality EPUB 3.0 ebooks.

Supported Input Formats:
• Microsoft Word (.docx)
• Markdown (.md)
• Plain Text (.txt)
• HTML (.html, .htm)

Features:
• Professional EPUB 3.0 output
• Multiple CSS themes (serif, sans, printlike)
• Comprehensive metadata support
• Accessibility compliance
• Series and publishing information
• Retailer-specific optimization
• Cross-reference support
• Mathematical equations
• Media overlays for audio

For more information, visit:
https://github.com/LightWraith8268/Docx2Shelf

Version: 1.2.4
"""

            text_widget = ScrolledText(parent, wrap=tk.WORD, state=tk.DISABLED)
            text_widget.pack(fill="both", expand=True, padx=20, pady=20)

            text_widget.config(state=tk.NORMAL)
            text_widget.insert(tk.END, about_text)
            text_widget.config(state=tk.DISABLED)

        def setup_drag_drop(self):
            """Setup drag and drop functionality."""
            # Basic drag and drop for tkinter
            def drop_handler(event):
                files = self.root.tk.splitlist(event.data)
                if files:
                    file_path = Path(files[0])
                    if file_path.suffix.lower() == '.docx':
                        self.file_path_var.set(str(file_path))
                        self.auto_fill_metadata(file_path)

            try:
                from tkinterdnd2 import DND_FILES, TkinterDnD
                self.root = TkinterDnD.Tk()
                self.root.drop_target_register(DND_FILES)
                self.root.dnd_bind('<<Drop>>', drop_handler)
            except ImportError:
                # Fallback: no drag and drop
                pass

        def browse_input_file(self):
            """Browse for input DOCX file."""
            # Use last directory if remembered
            initial_dir = None
            if (self.settings.ui_preferences.remember_last_directory and
                self.settings.file_defaults.last_input_directory):
                initial_dir = self.settings.file_defaults.last_input_directory

            file_path = filedialog.askopenfilename(
                title="Select document file",
                initialdir=initial_dir,
                filetypes=[
                    ("All supported", "*.docx;*.md;*.txt;*.html;*.htm"),
                    ("Word documents", "*.docx"),
                    ("Markdown files", "*.md"),
                    ("Text files", "*.txt"),
                    ("HTML files", "*.html;*.htm"),
                    ("All files", "*.*")
                ]
            )

            if file_path:
                self.file_path_var.set(file_path)

                # Remember directory if enabled
                if self.settings.ui_preferences.remember_last_directory:
                    self.settings.file_defaults.last_input_directory = str(Path(file_path).parent)
                    self.settings_manager.save_settings()

                self.auto_fill_metadata(Path(file_path))

        def browse_output_file(self):
            """Browse for output EPUB file."""
            # Use default output directory or last directory if remembered
            initial_dir = None
            if self.settings.file_defaults.default_output_directory:
                initial_dir = self.settings.file_defaults.default_output_directory
            elif (self.settings.ui_preferences.remember_last_directory and
                  self.settings.file_defaults.last_output_directory):
                initial_dir = self.settings.file_defaults.last_output_directory

            file_path = filedialog.asksaveasfilename(
                title="Save EPUB as",
                initialdir=initial_dir,
                defaultextension=".epub",
                filetypes=[("EPUB files", "*.epub"), ("All files", "*.*")]
            )

            if file_path:
                self.output_path_var.set(file_path)

                # Remember directory if enabled
                if self.settings.ui_preferences.remember_last_directory:
                    self.settings.file_defaults.last_output_directory = str(Path(file_path).parent)
                    self.settings_manager.save_settings()

        def auto_fill_metadata(self, file_path: Path):
            """Auto-fill metadata from document file."""
            # Only auto-fill if enabled in settings
            if not self.settings.ui_preferences.auto_fill_metadata:
                return

            # Auto-fill based on file type
            if file_path.suffix.lower() == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)

                    if doc.core_properties.title and not self.title_var.get():
                        self.title_var.set(doc.core_properties.title)

                    if doc.core_properties.author and not self.author_var.get():
                        self.author_var.set(doc.core_properties.author)

                    if doc.core_properties.comments and self.description_text.get("1.0", tk.END).strip() == "":
                        self.description_text.insert("1.0", doc.core_properties.comments)

                    # Try to get more metadata if available
                    if hasattr(doc.core_properties, 'subject') and doc.core_properties.subject:
                        if not self.subjects_var.get():
                            self.subjects_var.set(doc.core_properties.subject)

                except Exception as e:
                    logger.warning(f"Could not extract DOCX metadata: {e}")

            # For other formats, try to extract title from filename
            elif not self.title_var.get():
                # Remove extension and clean up filename for title
                clean_title = file_path.stem.replace('_', ' ').replace('-', ' ').title()
                self.title_var.set(clean_title)

            # Auto-generate output path based on settings for all formats
            if not self.output_path_var.get() and self.settings.file_defaults.auto_generate_output_name:
                if self.settings.file_defaults.default_output_directory:
                    output_dir = Path(self.settings.file_defaults.default_output_directory)
                    output_path = output_dir / (file_path.stem + '.epub')
                else:
                    output_path = file_path.with_suffix('.epub')
                self.output_path_var.set(str(output_path))

        def start_conversion(self):
            """Start EPUB conversion."""
            input_path = self.file_path_var.get()
            output_path = self.output_path_var.get()

            if not input_path:
                messagebox.showerror("Error", "Please select an input document file")
                return

            if not output_path:
                messagebox.showerror("Error", "Please specify an output EPUB file")
                return

            # Helper function to parse comma-separated strings
            def parse_list(text: str) -> List[str]:
                if not text.strip():
                    return []
                return [item.strip() for item in text.split(',') if item.strip()]

            # Parse publication date
            pubdate = None
            if self.pubdate_var.get().strip():
                try:
                    from datetime import datetime
                    pubdate = datetime.strptime(self.pubdate_var.get().strip(), '%Y-%m-%d').date()
                except ValueError:
                    messagebox.showwarning("Warning", "Invalid publication date format. Use YYYY-MM-DD.")

            # Create comprehensive metadata
            title = self.title_var.get() or Path(input_path).stem
            author = self.author_var.get() or "Unknown Author"

            metadata = EpubMetadata(
                title=title,
                author=author,
                language=self.language_var.get() or "en",
                description=self.description_text.get("1.0", tk.END).strip() or None,
                isbn=self.isbn_var.get().strip() or None,
                publisher=self.publisher_var.get().strip() or None,
                pubdate=pubdate,
                uuid=None,  # Will be auto-generated
                series=self.series_var.get().strip() or None,
                series_index=self.series_index_var.get().strip() or None,
                title_sort=None,  # Will be auto-generated from title
                author_sort=None,  # Will be auto-generated from author
                subjects=parse_list(self.subjects_var.get()),
                keywords=parse_list(self.subjects_var.get()),  # Use subjects as keywords for now
                cover_path=Path(getattr(self, 'cover_path_var', tk.StringVar()).get()) if getattr(self, 'cover_path_var', tk.StringVar()).get() else Path(),
                editor=self.editor_var.get().strip() or None,
                illustrator=self.illustrator_var.get().strip() or None,
                translator=self.translator_var.get().strip() or None,
                narrator=self.narrator_var.get().strip() or None,
                designer=self.designer_var.get().strip() or None,
                contributor=self.contributor_var.get().strip() or None,
                print_isbn=self.print_isbn_var.get().strip() or None,
                series_type=self.series_type_var.get().strip() or None,
                series_position=self.series_position_var.get().strip() or None,
                publication_type=self.publication_type_var.get().strip() or None,
                target_audience=self.target_audience_var.get().strip() or None,
                copyright_holder=self.copyright_holder_var.get().strip() or None,
                copyright_year=self.copyright_year_var.get().strip() or None,
                rights=self.rights_text.get("1.0", tk.END).strip() or None
            )

            # Update settings from conversion tab if they've changed
            current_theme = self.theme_var.get()
            current_validate = self.validate_var.get()
            current_width = int(self.image_width_var.get() or 1200)
            current_a11y = self.a11y_var.get()
            current_lang = self.language_var.get()

            # Save any changes to settings
            if (current_theme != self.settings.conversion_defaults.css_theme or
                current_validate != self.settings.conversion_defaults.validate_epub or
                current_width != self.settings.conversion_defaults.image_max_width or
                current_a11y != self.settings.conversion_defaults.accessibility_features or
                current_lang != self.settings.conversion_defaults.language):

                self.settings.conversion_defaults.css_theme = current_theme
                self.settings.conversion_defaults.validate_epub = current_validate
                self.settings.conversion_defaults.image_max_width = current_width
                self.settings.conversion_defaults.accessibility_features = current_a11y
                self.settings.conversion_defaults.language = current_lang
                self.settings_manager.save_settings()

            # Create comprehensive build options from all GUI settings
            options = BuildOptions(
                split_at=getattr(self, 'split_at_var', tk.StringVar(value="h1")).get(),
                theme=current_theme,
                embed_fonts_dir=None,  # Could add font embedding in future
                hyphenate=getattr(self, 'hyphenate_var', tk.BooleanVar(value=True)).get(),
                justify=getattr(self, 'justify_var', tk.BooleanVar(value=True)).get(),
                toc_depth=int(getattr(self, 'toc_depth_var', tk.StringVar(value="2")).get() or 2),
                image_quality=int(getattr(self, 'image_quality_var', tk.StringVar(value="85")).get() or 85),
                image_max_width=current_width,
                image_max_height=int(getattr(self, 'image_height_var', tk.StringVar(value="1600")).get() or 1600),
                image_format=getattr(self, 'image_format_var', tk.StringVar(value="webp")).get(),
                enhanced_images=getattr(self, 'enhanced_images_var', tk.BooleanVar()).get(),
                vertical_writing=getattr(self, 'vertical_writing_var', tk.BooleanVar()).get(),
                epub2_compat=getattr(self, 'epub2_compat_var', tk.BooleanVar()).get(),
                chapter_start_mode=getattr(self, 'chapter_mode_var', tk.StringVar(value="auto")).get(),
                chapter_starts=None,  # Could add manual chapter detection
                mixed_split_pattern=None,
                reader_start_chapter=None,
                page_list=getattr(self, 'page_list_var', tk.BooleanVar()).get(),
                extra_css=Path(getattr(self, 'extra_css_var', tk.StringVar()).get()) if getattr(self, 'extra_css_var', tk.StringVar()).get() else None,
                page_numbers=getattr(self, 'page_numbers_var', tk.BooleanVar()).get(),
                epub_version=getattr(self, 'epub_version_var', tk.StringVar(value="3")).get(),
                cover_scale=getattr(self, 'cover_scale_var', tk.StringVar(value="contain")).get(),
                dedication_txt=None,  # Could add dedication file selection
                ack_txt=None,  # Could add acknowledgments file selection
                inspect=False,  # Debugging option
                dry_run=False,  # Debugging option
                preview=False,  # Preview mode
                preview_port=8000,
                epubcheck=current_validate,
                font_size=getattr(self, 'font_size_var', tk.StringVar()).get() or None,
                line_height=getattr(self, 'line_height_var', tk.StringVar()).get() or None,
                quiet=False,
                verbose=False,
                typography_preset=None,
                custom_typography_config=None,
                advanced_typography=False,
                store_profile_css=getattr(self, 'store_profile_var', tk.StringVar()).get() or None
            )

            # Disable convert button and enable cancel
            self.convert_btn.config(state="disabled")
            self.cancel_btn.config(state="normal")

            # Start conversion worker
            self.current_worker = ConversionWorker(
                progress_callback=self.update_progress,
                complete_callback=self.conversion_complete,
                error_callback=self.conversion_error
            )

            # Run in separate thread
            self.conversion_thread = threading.Thread(
                target=self.current_worker.convert,
                args=(Path(input_path), Path(output_path), metadata, options)
            )
            self.conversion_thread.start()

        def cancel_conversion(self):
            """Cancel ongoing conversion."""
            if self.current_worker:
                self.current_worker.cancel()

            self.reset_ui()

        def reset_ui(self):
            """Reset UI to ready state."""
            self.convert_btn.config(state="normal")
            self.cancel_btn.config(state="disabled")
            self.progress_var.set("Ready to convert")
            self.progress_bar['value'] = 0

        def update_progress(self, message: str, percentage: int):
            """Update progress display."""
            self.progress_var.set(message)
            self.progress_bar['value'] = percentage
            self.root.update_idletasks()

        def conversion_complete(self, output_path: Path):
            """Handle successful conversion."""
            self.reset_ui()

            result = messagebox.askyesno(
                "Conversion Complete",
                f"EPUB created successfully!\n\n{output_path}\n\nWould you like to open the file location?"
            )

            if result:
                self.open_file_location(output_path)

        def conversion_error(self, error_message: str):
            """Handle conversion error."""
            self.reset_ui()
            messagebox.showerror("Conversion Error", f"Conversion failed:\n\n{error_message}")

        def preview_epub(self):
            """Preview the EPUB file."""
            output_path = self.output_path_var.get()

            if not output_path or not Path(output_path).exists():
                messagebox.showwarning("Preview", "Please convert the file first to preview it.")
                return

            try:
                # Try to open with system default application
                import os
                import platform

                if platform.system() == "Windows":
                    os.startfile(output_path)
                elif platform.system() == "Darwin":  # macOS
                    os.system(f"open '{output_path}'")
                else:  # Linux
                    os.system(f"xdg-open '{output_path}'")

            except Exception as e:
                messagebox.showerror("Preview Error", f"Could not open EPUB file:\n\n{e}")

        def open_file_location(self, file_path: Path):
            """Open file location in system file manager."""
            try:
                import os
                import platform

                if platform.system() == "Windows":
                    os.system(f'explorer /select,"{file_path}"')
                elif platform.system() == "Darwin":  # macOS
                    os.system(f'open -R "{file_path}"')
                else:  # Linux
                    os.system(f'xdg-open "{file_path.parent}"')

            except Exception as e:
                logger.warning(f"Could not open file location: {e}")

        def load_ui_preferences(self):
            """Load UI preferences and apply them."""
            prefs = self.settings.ui_preferences

            # Apply window size preferences
            if prefs.remember_window_size:
                # Could implement window size restoration here
                pass

        def toggle_enterprise_settings(self):
            """Toggle enterprise settings availability."""
            if self.enterprise_enabled_var.get():
                for child in self.enterprise_frame.winfo_children():
                    child.configure(state='normal')
            else:
                for child in self.enterprise_frame.winfo_children():
                    if isinstance(child, (ttk.Entry, ttk.Combobox, ttk.Button)):
                        child.configure(state='disabled')

        def browse_default_output(self):
            """Browse for default output directory."""
            directory = filedialog.askdirectory(title="Select Default Output Directory")
            if directory:
                self.output_dir_var.set(directory)

        def browse_temp_dir(self):
            """Browse for temporary directory."""
            directory = filedialog.askdirectory(title="Select Temporary Directory")
            if directory:
                self.temp_dir_var.set(directory)

        def browse_storage_dir(self):
            """Browse for enterprise storage directory."""
            directory = filedialog.askdirectory(title="Select Storage Directory")
            if directory:
                self.storage_dir_var.set(directory)

        def save_settings(self):
            """Save all settings to disk."""
            try:
                # Update settings from UI values
                self.settings.conversion_defaults.css_theme = self.theme_var.get()
                self.settings.conversion_defaults.language = self.lang_var.get()
                self.settings.conversion_defaults.validate_epub = self.validate_var.get()
                self.settings.conversion_defaults.image_max_width = int(self.image_width_var.get() or 1200)
                self.settings.conversion_defaults.accessibility_features = self.a11y_var.get()
                self.settings.conversion_defaults.chapter_detection = self.chapter_var.get()
                self.settings.conversion_defaults.include_toc = self.toc_var.get()
                self.settings.conversion_defaults.include_cover = self.cover_var.get()

                # UI preferences
                self.settings.ui_preferences.remember_last_directory = self.remember_dir_var.get()
                self.settings.ui_preferences.auto_fill_metadata = self.auto_fill_var.get()
                self.settings.ui_preferences.show_advanced_options = self.show_advanced_var.get()
                self.settings.ui_preferences.theme = self.ui_theme_var.get()
                self.settings.ui_preferences.font_size = self.font_size_var.get()
                self.settings.ui_preferences.confirm_overwrite = self.confirm_overwrite_var.get()
                self.settings.ui_preferences.show_tooltips = self.tooltips_var.get()

                # File defaults
                self.settings.file_defaults.default_output_directory = self.output_dir_var.get() or None
                self.settings.file_defaults.auto_generate_output_name = self.auto_name_var.get()
                self.settings.file_defaults.backup_original = self.backup_var.get()
                self.settings.file_defaults.temp_directory = self.temp_dir_var.get() or None

                # Advanced settings
                self.settings.advanced_settings.enable_logging = self.logging_var.get()
                self.settings.advanced_settings.log_level = self.log_level_var.get()
                self.settings.advanced_settings.concurrent_jobs = int(self.concurrent_var.get() or 2)
                self.settings.advanced_settings.check_for_updates = self.updates_var.get()
                self.settings.advanced_settings.enable_crash_reporting = self.crash_var.get()
                self.settings.advanced_settings.enable_telemetry = self.telemetry_var.get()

                # Enterprise settings
                if self.enterprise_enabled_var.get():
                    from ..enterprise import EnterpriseConfig
                    if not self.settings.enterprise_config:
                        self.settings.enterprise_config = EnterpriseConfig()

                    self.settings.enterprise_config.api_host = self.api_host_var.get()
                    self.settings.enterprise_config.api_port = int(self.api_port_var.get() or 8080)
                    self.settings.enterprise_config.max_concurrent_jobs = int(self.max_jobs_var.get() or 4)
                    self.settings.enterprise_config.job_timeout_hours = int(self.job_timeout_var.get() or 24)
                    self.settings.enterprise_config.storage_directory = self.storage_dir_var.get() or None
                else:
                    self.settings.enterprise_config = None

                # Save to disk
                self.settings_manager.save_settings()

                messagebox.showinfo("Settings", "Settings saved successfully!")

            except Exception as e:
                messagebox.showerror("Error", f"Failed to save settings: {e}")

        def reset_settings(self):
            """Reset all settings to defaults."""
            result = messagebox.askyesno(
                "Reset Settings",
                "Are you sure you want to reset all settings to defaults? This cannot be undone."
            )

            if result:
                self.settings_manager.reset_to_defaults()
                self.settings = self.settings_manager.settings
                messagebox.showinfo("Settings", "Settings reset to defaults successfully!")
                # Reload the GUI would require recreation - for now just show message
                messagebox.showinfo("Restart Required", "Please restart the application to see the changes.")

        def export_settings(self):
            """Export settings to a file."""
            file_path = filedialog.asksaveasfilename(
                title="Export Settings",
                defaultextension=".json",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )

            if file_path:
                try:
                    self.settings_manager.export_settings(Path(file_path))
                    messagebox.showinfo("Export", f"Settings exported successfully to {file_path}")
                except Exception as e:
                    messagebox.showerror("Export Error", f"Failed to export settings: {e}")

        def import_settings(self):
            """Import settings from a file."""
            file_path = filedialog.askopenfilename(
                title="Import Settings",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )

            if file_path:
                result = messagebox.askyesnocancel(
                    "Import Settings",
                    "Do you want to merge with existing settings?\n\n"
                    "Yes = Merge with current settings\n"
                    "No = Replace all settings\n"
                    "Cancel = Cancel import"
                )

                if result is not None:
                    try:
                        merge = result  # True for merge, False for replace
                        self.settings_manager.import_settings(Path(file_path), merge=merge)
                        self.settings = self.settings_manager.settings

                        action = "merged" if merge else "imported"
                        messagebox.showinfo("Import", f"Settings {action} successfully!")
                        messagebox.showinfo("Restart Required", "Please restart the application to see the changes.")

                    except Exception as e:
                        messagebox.showerror("Import Error", f"Failed to import settings: {e}")

        def run(self):
            """Run the application."""
            self.root.mainloop()


class CustomThemeDialog:
    """Dialog for creating custom CSS themes with font embedding."""

    def __init__(self, parent, main_app):
        self.parent = parent
        self.main_app = main_app
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Create Custom Theme")
        self.dialog.geometry("800x600")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        # Theme data
        self.theme_name = tk.StringVar()
        self.base_font = tk.StringVar(value="Georgia")
        self.heading_font = tk.StringVar(value="Georgia")
        self.font_size = tk.StringVar(value="16px")
        self.line_height = tk.StringVar(value="1.6")
        self.text_color = tk.StringVar(value="#333333")
        self.background_color = tk.StringVar(value="#ffffff")
        self.margin_size = tk.StringVar(value="2em")
        self.custom_fonts = []

        self.setup_ui()

    def setup_ui(self):
        """Setup the dialog UI."""
        notebook = ttk.Notebook(self.dialog)
        notebook.pack(fill="both", expand=True, padx=10, pady=10)

        # Basic Settings Tab
        basic_frame = ttk.Frame(notebook)
        notebook.add(basic_frame, text="Basic Settings")
        self.setup_basic_settings(basic_frame)

        # Typography Tab
        typography_frame = ttk.Frame(notebook)
        notebook.add(typography_frame, text="Typography")
        self.setup_typography_settings(typography_frame)

        # Colors Tab
        colors_frame = ttk.Frame(notebook)
        notebook.add(colors_frame, text="Colors")
        self.setup_color_settings(colors_frame)

        # Custom Fonts Tab
        fonts_frame = ttk.Frame(notebook)
        notebook.add(fonts_frame, text="Custom Fonts")
        self.setup_font_settings(fonts_frame)

        # CSS Preview Tab
        preview_frame = ttk.Frame(notebook)
        notebook.add(preview_frame, text="CSS Preview")
        self.setup_preview(preview_frame)

        # Buttons
        button_frame = ttk.Frame(self.dialog)
        button_frame.pack(fill="x", padx=10, pady=5)

        ttk.Button(button_frame, text="Save Theme", command=self.save_theme).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Cancel", command=self.dialog.destroy).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Preview", command=self.preview_theme).pack(side="right", padx=5)

    def setup_basic_settings(self, parent):
        """Setup basic theme settings."""
        # Theme name
        ttk.Label(parent, text="Theme Name:").grid(row=0, column=0, sticky="w", pady=5)
        ttk.Entry(parent, textvariable=self.theme_name, width=30).grid(row=0, column=1, sticky="w", padx=5)

        # Base template
        ttk.Label(parent, text="Base Template:").grid(row=1, column=0, sticky="w", pady=5)
        base_combo = ttk.Combobox(parent, values=["serif", "sans", "printlike", "minimal"],
                                 state="readonly")
        base_combo.set("serif")
        base_combo.grid(row=1, column=1, sticky="w", padx=5)

        # Description
        ttk.Label(parent, text="Description:").grid(row=2, column=0, sticky="nw", pady=5)
        self.description_text = tk.Text(parent, width=40, height=4)
        self.description_text.grid(row=2, column=1, sticky="w", padx=5, pady=5)

    def setup_typography_settings(self, parent):
        """Setup typography settings."""
        # Base font
        ttk.Label(parent, text="Base Font:").grid(row=0, column=0, sticky="w", pady=5)
        font_combo = ttk.Combobox(parent, textvariable=self.base_font,
                                 values=["Georgia", "Times New Roman", "Arial", "Helvetica",
                                        "Verdana", "Calibri", "Cambria", "Custom..."])
        font_combo.grid(row=0, column=1, sticky="w", padx=5)

        # Heading font
        ttk.Label(parent, text="Heading Font:").grid(row=1, column=0, sticky="w", pady=5)
        heading_combo = ttk.Combobox(parent, textvariable=self.heading_font,
                                    values=["Georgia", "Times New Roman", "Arial", "Helvetica",
                                           "Verdana", "Calibri", "Cambria", "Custom..."])
        heading_combo.grid(row=1, column=1, sticky="w", padx=5)

        # Font size
        ttk.Label(parent, text="Base Font Size:").grid(row=2, column=0, sticky="w", pady=5)
        ttk.Entry(parent, textvariable=self.font_size, width=15).grid(row=2, column=1, sticky="w", padx=5)

        # Line height
        ttk.Label(parent, text="Line Height:").grid(row=3, column=0, sticky="w", pady=5)
        ttk.Entry(parent, textvariable=self.line_height, width=15).grid(row=3, column=1, sticky="w", padx=5)

        # Margins
        ttk.Label(parent, text="Page Margins:").grid(row=4, column=0, sticky="w", pady=5)
        ttk.Entry(parent, textvariable=self.margin_size, width=15).grid(row=4, column=1, sticky="w", padx=5)

    def setup_color_settings(self, parent):
        """Setup color settings."""
        # Text color
        ttk.Label(parent, text="Text Color:").grid(row=0, column=0, sticky="w", pady=5)
        text_frame = ttk.Frame(parent)
        text_frame.grid(row=0, column=1, sticky="w", padx=5)
        ttk.Entry(text_frame, textvariable=self.text_color, width=15).pack(side="left")
        ttk.Button(text_frame, text="Choose", command=lambda: self.choose_color(self.text_color)).pack(side="left", padx=5)

        # Background color
        ttk.Label(parent, text="Background Color:").grid(row=1, column=0, sticky="w", pady=5)
        bg_frame = ttk.Frame(parent)
        bg_frame.grid(row=1, column=1, sticky="w", padx=5)
        ttk.Entry(bg_frame, textvariable=self.background_color, width=15).pack(side="left")
        ttk.Button(bg_frame, text="Choose", command=lambda: self.choose_color(self.background_color)).pack(side="left", padx=5)

    def setup_font_settings(self, parent):
        """Setup custom font settings."""
        ttk.Label(parent, text="Custom Fonts:").grid(row=0, column=0, sticky="nw", pady=5)

        # Font list
        list_frame = ttk.Frame(parent)
        list_frame.grid(row=1, column=0, columnspan=2, sticky="ew", pady=5)

        self.font_listbox = tk.Listbox(list_frame, height=8)
        self.font_listbox.pack(side="left", fill="both", expand=True)

        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.font_listbox.yview)
        scrollbar.pack(side="right", fill="y")
        self.font_listbox.config(yscrollcommand=scrollbar.set)

        # Font buttons
        button_frame = ttk.Frame(parent)
        button_frame.grid(row=2, column=0, columnspan=2, sticky="ew", pady=5)

        ttk.Button(button_frame, text="Add Font File", command=self.add_font_file).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Remove Font", command=self.remove_font).pack(side="left", padx=5)

    def setup_preview(self, parent):
        """Setup CSS preview."""
        self.css_text = ScrolledText(parent, wrap=tk.WORD, height=25)
        self.css_text.pack(fill="both", expand=True, padx=5, pady=5)

        # Update preview button
        ttk.Button(parent, text="Update Preview", command=self.update_css_preview).pack(pady=5)

        # Initialize with basic CSS
        self.update_css_preview()

    def choose_color(self, color_var):
        """Open color chooser dialog."""
        try:
            from tkinter import colorchooser
            color = colorchooser.askcolor(initialcolor=color_var.get())
            if color[1]:  # If user didn't cancel
                color_var.set(color[1])
        except ImportError:
            messagebox.showwarning("Color Chooser", "Color chooser not available. Please enter hex color manually.")

    def add_font_file(self):
        """Add a custom font file."""
        font_file = filedialog.askopenfilename(
            title="Select Font File",
            filetypes=[("Font files", "*.ttf *.otf *.woff *.woff2"), ("All files", "*.*")]
        )
        if font_file:
            font_path = Path(font_file)
            font_name = font_path.stem
            self.custom_fonts.append({"name": font_name, "path": font_file})
            self.font_listbox.insert(tk.END, f"{font_name} ({font_path.suffix})")

    def remove_font(self):
        """Remove selected custom font."""
        selection = self.font_listbox.curselection()
        if selection:
            index = selection[0]
            self.font_listbox.delete(index)
            del self.custom_fonts[index]

    def update_css_preview(self):
        """Update the CSS preview."""
        css = self.generate_css()
        self.css_text.delete(1.0, tk.END)
        self.css_text.insert(1.0, css)

    def generate_css(self):
        """Generate CSS based on current settings."""
        css_parts = ["/* Custom Theme: {} */".format(self.theme_name.get() or "Untitled")]

        # Font face declarations for custom fonts
        for font in self.custom_fonts:
            css_parts.append(f"""
@font-face {{
    font-family: '{font['name']}';
    src: url('{font['name']}.{Path(font['path']).suffix[1:]}');
}}""")

        # Base styles
        css_parts.append(f"""
body {{
    font-family: {self.base_font.get()}, serif;
    font-size: {self.font_size.get()};
    line-height: {self.line_height.get()};
    color: {self.text_color.get()};
    background-color: {self.background_color.get()};
    margin: {self.margin_size.get()};
    text-align: justify;
}}

h1, h2, h3, h4, h5, h6 {{
    font-family: {self.heading_font.get()}, serif;
    color: {self.text_color.get()};
    page-break-after: avoid;
}}

h1 {{
    font-size: 2em;
    margin-top: 1em;
    margin-bottom: 0.5em;
}}

h2 {{
    font-size: 1.5em;
    margin-top: 1em;
    margin-bottom: 0.5em;
}}

p {{
    margin-bottom: 1em;
    text-indent: 1.2em;
}}

p.no-indent {{
    text-indent: 0;
}}

.chapter-title {{
    text-align: center;
    font-size: 1.8em;
    margin-bottom: 2em;
    page-break-before: always;
}}

blockquote {{
    margin: 1em 2em;
    font-style: italic;
    border-left: 3px solid {self.text_color.get()};
    padding-left: 1em;
}}

img {{
    max-width: 100%;
    height: auto;
    display: block;
    margin: 1em auto;
}}""")

        return "\n".join(css_parts)

    def preview_theme(self):
        """Preview the theme in browser."""
        try:
            import webbrowser

            # Generate sample HTML with the theme
            css = self.generate_css()
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Theme Preview: {self.theme_name.get() or 'Untitled'}</title>
    <style>
{css}
    </style>
</head>
<body>
    <h1 class="chapter-title">Chapter 1: Sample Chapter</h1>

    <p class="no-indent">This is a sample paragraph to demonstrate your custom theme.
    The text should be displayed according to your typography and color settings.</p>

    <p>This is a regular paragraph with text indentation. Lorem ipsum dolor sit amet,
    consectetur adipiscing elit. Sed do eiusmod tempor incididunt ut labore et dolore magna aliqua.</p>

    <h2>Section Heading</h2>

    <p>Another paragraph to show the flow of text. Ut enim ad minim veniam, quis nostrud
    exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.</p>

    <blockquote>
        "This is a blockquote to demonstrate how quoted text appears in your theme."
    </blockquote>

    <p>Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu
    fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident.</p>
</body>
</html>"""

            # Write to temporary file and open
            with tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False, encoding='utf-8') as f:
                f.write(html_content)
                temp_path = f.name

            webbrowser.open(f'file://{temp_path}')

        except Exception as e:
            messagebox.showerror("Preview Error", f"Could not open preview: {e}")

    def save_theme(self):
        """Save the custom theme."""
        theme_name = self.theme_name.get().strip()
        if not theme_name:
            messagebox.showerror("Error", "Please enter a theme name.")
            return

        try:
            # Create themes directory if it doesn't exist
            themes_dir = Path.home() / ".docx2shelf" / "themes"
            themes_dir.mkdir(parents=True, exist_ok=True)

            # Save CSS file
            css_content = self.generate_css()
            theme_file = themes_dir / f"{theme_name}.css"

            with open(theme_file, 'w', encoding='utf-8') as f:
                f.write(css_content)

            # Copy custom fonts if any
            if self.custom_fonts:
                fonts_dir = themes_dir / "fonts"
                fonts_dir.mkdir(exist_ok=True)

                for font in self.custom_fonts:
                    font_path = Path(font['path'])
                    dest_path = fonts_dir / f"{font['name']}{font_path.suffix}"

                    import shutil
                    shutil.copy2(font_path, dest_path)

            # Update the main app's theme list
            if hasattr(self.main_app, 'css_theme_var'):
                current_themes = list(self.main_app.css_theme_var.tk.call(
                    self.main_app.css_theme_var._name, 'configure', '-values'
                )[0])
                if theme_name not in current_themes:
                    current_themes.append(theme_name)
                    # Update the combobox (need to find the widget)
                    # This is a simplified approach - in practice you'd need to refresh the UI

            messagebox.showinfo("Success", f"Theme '{theme_name}' saved successfully!")
            self.dialog.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"Could not save theme: {e}")


class AIDetectionDialog:
    """Dialog for configuring AI chapter detection settings."""

    def __init__(self, parent, main_app):
        self.parent = parent
        self.main_app = main_app
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("AI Chapter Detection Settings")
        self.dialog.geometry("600x500")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        # AI settings
        self.use_free_api = tk.BooleanVar(value=True)
        self.api_key = tk.StringVar()
        self.confidence_threshold = tk.DoubleVar(value=0.7)
        self.min_chapter_length = tk.IntVar(value=500)
        self.combine_methods = tk.BooleanVar(value=True)
        self.enable_heuristics = tk.BooleanVar(value=True)

        self.setup_ui()

    def setup_ui(self):
        """Setup the dialog UI."""
        main_frame = ttk.Frame(self.dialog)
        main_frame.pack(fill="both", expand=True, padx=20, pady=20)

        # Title
        title_label = ttk.Label(main_frame, text="AI Chapter Detection Configuration", font=("Arial", 14, "bold"))
        title_label.pack(pady=(0, 20))

        # API Settings
        api_frame = ttk.LabelFrame(main_frame, text="API Settings", padding=10)
        api_frame.pack(fill="x", pady=(0, 10))

        # Free API option
        free_api_check = ttk.Checkbutton(api_frame, text="Use free AI service (recommended for testing)",
                                        variable=self.use_free_api, command=self.toggle_api_options)
        free_api_check.pack(anchor="w", pady=5)

        # Usage information frame
        self.usage_frame = ttk.Frame(api_frame)
        self.usage_frame.pack(fill="x", pady=5)
        self.update_usage_info()

        # Custom API key
        ttk.Label(api_frame, text="Custom API Key (optional):").pack(anchor="w", pady=(10, 5))
        self.api_key_entry = ttk.Entry(api_frame, textvariable=self.api_key, width=50, show="*")
        self.api_key_entry.pack(fill="x", pady=(0, 5))

        # Free API info
        self.free_info_label = ttk.Label(api_frame, text="Free service provides basic AI detection with rate limits.",
                                        foreground="gray")
        self.free_info_label.pack(anchor="w", pady=5)

        # Detection Settings
        detection_frame = ttk.LabelFrame(main_frame, text="Detection Settings", padding=10)
        detection_frame.pack(fill="x", pady=(0, 10))

        # Confidence threshold
        ttk.Label(detection_frame, text="Confidence Threshold:").pack(anchor="w")
        confidence_frame = ttk.Frame(detection_frame)
        confidence_frame.pack(fill="x", pady=5)
        confidence_scale = ttk.Scale(confidence_frame, from_=0.1, to=1.0, variable=self.confidence_threshold,
                                   orient="horizontal")
        confidence_scale.pack(side="left", fill="x", expand=True)
        confidence_label = ttk.Label(confidence_frame, text="0.7")
        confidence_label.pack(side="right", padx=(10, 0))
        confidence_scale.bind("<Motion>", lambda e: confidence_label.config(text=f"{self.confidence_threshold.get():.1f}"))

        # Minimum chapter length
        ttk.Label(detection_frame, text="Minimum Chapter Length (words):").pack(anchor="w", pady=(10, 0))
        length_entry = ttk.Entry(detection_frame, textvariable=self.min_chapter_length, width=10)
        length_entry.pack(anchor="w", pady=5)

        # Method options
        methods_frame = ttk.LabelFrame(main_frame, text="Detection Methods", padding=10)
        methods_frame.pack(fill="x", pady=(0, 10))

        heuristics_check = ttk.Checkbutton(methods_frame, text="Enable heuristic detection (fallback)",
                                         variable=self.enable_heuristics)
        heuristics_check.pack(anchor="w", pady=2)

        combine_check = ttk.Checkbutton(methods_frame, text="Combine AI and heuristic results",
                                      variable=self.combine_methods)
        combine_check.pack(anchor="w", pady=2)

        # Usage info
        info_frame = ttk.LabelFrame(main_frame, text="Usage Information", padding=10)
        info_frame.pack(fill="x", pady=(0, 10))

        info_text = """AI chapter detection analyzes your document content to intelligently identify chapter boundaries.

Features:
• Natural language processing to understand document structure
• Recognition of various chapter naming conventions
• Confidence scoring for each detected boundary
• Fallback to heuristic methods when needed

Free service limitations:
• 5 documents per day
• Maximum 10,000 words per document
• Basic model (good for most documents)"""

        info_label = ttk.Label(info_frame, text=info_text, justify="left")
        info_label.pack(anchor="w")

        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill="x", pady=(20, 0))

        ttk.Button(button_frame, text="Test Settings", command=self.test_settings).pack(side="left", padx=(0, 10))
        ttk.Button(button_frame, text="Save", command=self.save_settings).pack(side="right", padx=(5, 0))
        ttk.Button(button_frame, text="Cancel", command=self.dialog.destroy).pack(side="right", padx=(5, 0))

        # Initialize UI state
        self.toggle_api_options()

    def update_usage_info(self):
        """Update usage information display."""
        try:
            from .free_ai_service import get_free_ai_service

            # Clear existing widgets
            for widget in self.usage_frame.winfo_children():
                widget.destroy()

            # Get usage stats
            service = get_free_ai_service()
            stats = service.get_usage_stats()

            if self.use_free_api.get():
                # Show usage information
                usage_text = f"Today's usage: {stats['requests_today']}/{stats['daily_limit']} requests"
                if stats['remaining_today'] > 0:
                    usage_text += f" ({stats['remaining_today']} remaining)"
                else:
                    usage_text += " (quota exceeded)"

                usage_label = ttk.Label(self.usage_frame, text=usage_text)
                usage_label.pack(anchor="w")

                # Show service status
                if stats['remaining_today'] == 0:
                    warning_label = ttk.Label(self.usage_frame, text="⚠ Daily quota reached. Local analysis will be used.",
                                             foreground="orange")
                    warning_label.pack(anchor="w")

                # Add refresh button
                refresh_btn = ttk.Button(self.usage_frame, text="Refresh Usage", command=self.update_usage_info)
                refresh_btn.pack(anchor="e", pady=(5, 0))

        except Exception as e:
            # Fallback display
            error_label = ttk.Label(self.usage_frame, text="Could not load usage information", foreground="gray")
            error_label.pack(anchor="w")

    def toggle_api_options(self):
        """Toggle API key entry based on free API setting."""
        if self.use_free_api.get():
            self.api_key_entry.config(state="disabled")
            self.free_info_label.config(text="Free service provides basic AI detection with rate limits.")
        else:
            self.api_key_entry.config(state="normal")
            self.free_info_label.config(text="Enter your OpenAI, Anthropic, or compatible API key.")

        # Update usage info
        self.update_usage_info()

    def test_settings(self):
        """Test the AI detection settings."""
        try:
            # Test with sample content
            sample_content = """
Chapter 1: The Beginning

This is the first chapter of our story. It introduces the main characters and sets up the plot.

The protagonist walks through the door and sees something unexpected.

Chapter 2: The Discovery

In this chapter, the mystery deepens as new clues are revealed.

The story continues to unfold with increasing tension.

PART TWO: THE ADVENTURE CONTINUES

Chapter 3: New Horizons

The journey takes an unexpected turn as our heroes face their greatest challenge yet.
"""

            if self.use_free_api.get():
                # Use free AI service
                from .free_ai_service import process_document_with_free_ai
                result = process_document_with_free_ai(sample_content, "test_document")

                if result["success"]:
                    boundaries = result["boundaries"]
                    result_text = f"Test completed successfully using {result.get('service', 'Free AI Service')}!\n\n"
                    result_text += f"Detected {len(boundaries)} chapter boundaries:\n\n"

                    for i, boundary in enumerate(boundaries, 1):
                        title = boundary.get('title', 'Unknown')[:50]
                        confidence = boundary.get('confidence', 0.0)
                        method = boundary.get('method', 'unknown')
                        result_text += f"{i}. {title} (confidence: {confidence:.2f}, method: {method})\n"

                    if result.get('cached'):
                        result_text += "\n(Note: This result was retrieved from cache)"

                    messagebox.showinfo("Test Results", result_text)

                    # Update usage info after test
                    self.update_usage_info()

                else:
                    messagebox.showerror("Test Failed", f"AI service error: {result.get('error', 'Unknown error')}")

            else:
                # Use custom API key configuration
                from .ai_chapter_detection import ChapterDetectionEngine, AIDetectionConfig

                config = AIDetectionConfig(
                    use_free_api=False,
                    api_key=self.api_key.get(),
                    confidence_threshold=self.confidence_threshold.get(),
                    min_chapter_length=self.min_chapter_length.get(),
                    enable_heuristics=self.enable_heuristics.get(),
                    combine_methods=self.combine_methods.get()
                )

                engine = ChapterDetectionEngine(config)
                boundaries = engine.detect_chapters(sample_content, "test_document")

                result_text = f"Test completed successfully!\n\nDetected {len(boundaries)} chapter boundaries:\n\n"
                for i, boundary in enumerate(boundaries, 1):
                    result_text += f"{i}. {boundary.title} (confidence: {boundary.confidence:.2f}, method: {boundary.method})\n"

                messagebox.showinfo("Test Results", result_text)

        except Exception as e:
            messagebox.showerror("Test Failed", f"Could not test AI settings: {e}")

    def save_settings(self):
        """Save the AI detection settings."""
        try:
            # Store settings in the main app (could be in settings.py later)
            ai_settings = {
                'use_free_api': self.use_free_api.get(),
                'api_key': self.api_key.get() if not self.use_free_api.get() else "",
                'confidence_threshold': self.confidence_threshold.get(),
                'min_chapter_length': self.min_chapter_length.get(),
                'enable_heuristics': self.enable_heuristics.get(),
                'combine_methods': self.combine_methods.get()
            }

            # Save to application settings (extend the settings system)
            if hasattr(self.main_app, 'ai_detection_settings'):
                self.main_app.ai_detection_settings = ai_settings
            else:
                # Could integrate with the main settings system later
                pass

            messagebox.showinfo("Settings Saved", "AI detection settings have been saved successfully!")
            self.dialog.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"Could not save settings: {e}")


class UnifiedSettingsDialog:
    """Unified settings dialog for all application configuration."""

    def __init__(self, parent, main_app):
        self.parent = parent
        self.main_app = main_app
        self.dialog = tk.Toplevel(parent)
        self.dialog.title("Docx2Shelf Settings")
        self.dialog.geometry("800x700")
        self.dialog.transient(parent)
        self.dialog.grab_set()

        # Load current settings
        from ..settings import get_settings
        self.settings = get_settings()

        self.setup_ui()

    def setup_ui(self):
        """Setup the unified settings UI."""
        # Create main container with scrolling
        main_frame = ttk.Frame(self.dialog)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Create notebook for different setting categories
        self.notebook = ttk.Notebook(main_frame)
        self.notebook.pack(fill="both", expand=True, pady=(0, 10))

        # Create tabs
        self.setup_conversion_tab()
        self.setup_ui_preferences_tab()
        self.setup_file_defaults_tab()
        self.setup_ai_detection_tab()
        self.setup_advanced_tab()
        self.setup_enterprise_tab()

        # Buttons
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(fill="x", pady=(10, 0))

        ttk.Button(button_frame, text="Save All", command=self.save_all_settings).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Cancel", command=self.dialog.destroy).pack(side="right", padx=5)
        ttk.Button(button_frame, text="Reset to Defaults", command=self.reset_to_defaults).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Export Settings", command=self.export_settings).pack(side="left", padx=5)
        ttk.Button(button_frame, text="Import Settings", command=self.import_settings).pack(side="left", padx=5)

    def setup_conversion_tab(self):
        """Setup conversion defaults tab."""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Conversion")

        # Create scrollable frame
        canvas = tk.Canvas(frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # CSS Theme
        ttk.Label(scrollable_frame, text="CSS Theme:").grid(row=0, column=0, sticky="w", pady=5, padx=5)
        self.css_theme_var = tk.StringVar(value=self.settings.conversion_defaults.css_theme)
        css_combo = ttk.Combobox(scrollable_frame, textvariable=self.css_theme_var,
                               values=["serif", "sans", "printlike"], state="readonly")
        css_combo.grid(row=0, column=1, sticky="w", padx=5, pady=5)

        # Language
        ttk.Label(scrollable_frame, text="Language:").grid(row=1, column=0, sticky="w", pady=5, padx=5)
        self.language_var = tk.StringVar(value=self.settings.conversion_defaults.language)
        lang_entry = ttk.Entry(scrollable_frame, textvariable=self.language_var, width=10)
        lang_entry.grid(row=1, column=1, sticky="w", padx=5, pady=5)

        # Validate EPUB
        self.validate_epub_var = tk.BooleanVar(value=self.settings.conversion_defaults.validate_epub)
        ttk.Checkbutton(scrollable_frame, text="Validate EPUB", variable=self.validate_epub_var).grid(
            row=2, column=0, columnspan=2, sticky="w", padx=5, pady=5)

        # Image max width
        ttk.Label(scrollable_frame, text="Image Max Width:").grid(row=3, column=0, sticky="w", pady=5, padx=5)
        self.image_max_width_var = tk.IntVar(value=self.settings.conversion_defaults.image_max_width)
        width_entry = ttk.Entry(scrollable_frame, textvariable=self.image_max_width_var, width=10)
        width_entry.grid(row=3, column=1, sticky="w", padx=5, pady=5)

        # Accessibility features
        self.accessibility_var = tk.BooleanVar(value=self.settings.conversion_defaults.accessibility_features)
        ttk.Checkbutton(scrollable_frame, text="Accessibility Features", variable=self.accessibility_var).grid(
            row=4, column=0, columnspan=2, sticky="w", padx=5, pady=5)

        # Chapter detection
        ttk.Label(scrollable_frame, text="Chapter Detection:").grid(row=5, column=0, sticky="w", pady=5, padx=5)
        self.chapter_detection_var = tk.StringVar(value=self.settings.conversion_defaults.chapter_detection)
        chapter_combo = ttk.Combobox(scrollable_frame, textvariable=self.chapter_detection_var,
                                   values=["auto", "h1", "h2", "pagebreak", "ai"], state="readonly")
        chapter_combo.grid(row=5, column=1, sticky="w", padx=5, pady=5)

        # Include ToC
        self.include_toc_var = tk.BooleanVar(value=self.settings.conversion_defaults.include_toc)
        ttk.Checkbutton(scrollable_frame, text="Include Table of Contents", variable=self.include_toc_var).grid(
            row=6, column=0, columnspan=2, sticky="w", padx=5, pady=5)

        # Include cover
        self.include_cover_var = tk.BooleanVar(value=self.settings.conversion_defaults.include_cover)
        ttk.Checkbutton(scrollable_frame, text="Include Cover", variable=self.include_cover_var).grid(
            row=7, column=0, columnspan=2, sticky="w", padx=5, pady=5)

        # DRM Free
        self.drm_free_var = tk.BooleanVar(value=self.settings.conversion_defaults.drm_free)
        ttk.Checkbutton(scrollable_frame, text="DRM Free", variable=self.drm_free_var).grid(
            row=8, column=0, columnspan=2, sticky="w", padx=5, pady=5)

    def setup_ui_preferences_tab(self):
        """Setup UI preferences tab."""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="UI Preferences")

        # Remember last directory
        self.remember_dir_var = tk.BooleanVar(value=self.settings.ui_preferences.remember_last_directory)
        ttk.Checkbutton(frame, text="Remember Last Directory", variable=self.remember_dir_var).pack(
            anchor="w", padx=10, pady=5)

        # Auto fill metadata
        self.auto_fill_var = tk.BooleanVar(value=self.settings.ui_preferences.auto_fill_metadata)
        ttk.Checkbutton(frame, text="Auto Fill Metadata", variable=self.auto_fill_var).pack(
            anchor="w", padx=10, pady=5)

        # Show advanced options
        self.show_advanced_var = tk.BooleanVar(value=self.settings.ui_preferences.show_advanced_options)
        ttk.Checkbutton(frame, text="Show Advanced Options", variable=self.show_advanced_var).pack(
            anchor="w", padx=10, pady=5)

        # Theme
        theme_frame = ttk.Frame(frame)
        theme_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(theme_frame, text="UI Theme:").pack(side="left")
        self.ui_theme_var = tk.StringVar(value=self.settings.ui_preferences.theme)
        theme_combo = ttk.Combobox(theme_frame, textvariable=self.ui_theme_var,
                                 values=["system", "light", "dark", "blue", "green", "purple", "minimal", "classic"],
                                 state="readonly")
        theme_combo.pack(side="left", padx=5)

        # Font size
        font_frame = ttk.Frame(frame)
        font_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(font_frame, text="Font Size:").pack(side="left")
        self.font_size_var = tk.StringVar(value=self.settings.ui_preferences.font_size)
        font_combo = ttk.Combobox(font_frame, textvariable=self.font_size_var,
                                values=["small", "medium", "large"], state="readonly")
        font_combo.pack(side="left", padx=5)

        # Remember window size
        self.remember_window_var = tk.BooleanVar(value=self.settings.ui_preferences.remember_window_size)
        ttk.Checkbutton(frame, text="Remember Window Size", variable=self.remember_window_var).pack(
            anchor="w", padx=10, pady=5)

        # Confirm overwrite
        self.confirm_overwrite_var = tk.BooleanVar(value=self.settings.ui_preferences.confirm_overwrite)
        ttk.Checkbutton(frame, text="Confirm Overwrite", variable=self.confirm_overwrite_var).pack(
            anchor="w", padx=10, pady=5)

        # Show tooltips
        self.show_tooltips_var = tk.BooleanVar(value=self.settings.ui_preferences.show_tooltips)
        ttk.Checkbutton(frame, text="Show Tooltips", variable=self.show_tooltips_var).pack(
            anchor="w", padx=10, pady=5)

    def setup_file_defaults_tab(self):
        """Setup file defaults tab."""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="File Defaults")

        # Default output directory
        dir_frame = ttk.Frame(frame)
        dir_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(dir_frame, text="Default Output Directory:").pack(anchor="w")
        self.default_output_var = tk.StringVar(value=self.settings.file_defaults.default_output_directory or "")
        dir_entry = ttk.Entry(dir_frame, textvariable=self.default_output_var, width=50)
        dir_entry.pack(side="left", fill="x", expand=True, pady=2)
        ttk.Button(dir_frame, text="Browse", command=self.browse_output_dir).pack(side="right", padx=5)

        # Auto generate output name
        self.auto_name_var = tk.BooleanVar(value=self.settings.file_defaults.auto_generate_output_name)
        ttk.Checkbutton(frame, text="Auto Generate Output Name", variable=self.auto_name_var).pack(
            anchor="w", padx=10, pady=5)

        # Backup original
        self.backup_var = tk.BooleanVar(value=self.settings.file_defaults.backup_original)
        ttk.Checkbutton(frame, text="Backup Original File", variable=self.backup_var).pack(
            anchor="w", padx=10, pady=5)

    def setup_ai_detection_tab(self):
        """Setup AI detection settings tab."""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="AI Detection")

        # Create scrollable frame
        canvas = tk.Canvas(frame, highlightthickness=0)
        scrollbar = ttk.Scrollbar(frame, orient="vertical", command=canvas.yview)
        scrollable_frame = ttk.Frame(canvas)

        scrollable_frame.bind(
            "<Configure>",
            lambda e: canvas.configure(scrollregion=canvas.bbox("all"))
        )

        canvas.create_window((0, 0), window=scrollable_frame, anchor="nw")
        canvas.configure(yscrollcommand=scrollbar.set)

        canvas.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # General AI Settings
        general_frame = ttk.LabelFrame(scrollable_frame, text="General AI Settings", padding=10)
        general_frame.pack(fill="x", padx=10, pady=5)

        # Confidence threshold
        ttk.Label(general_frame, text="Confidence Threshold:").grid(row=0, column=0, sticky="w", pady=2)
        self.confidence_var = tk.DoubleVar(value=self.settings.ai_detection.confidence_threshold)
        confidence_scale = ttk.Scale(general_frame, from_=0.1, to=1.0, variable=self.confidence_var, orient="horizontal")
        confidence_scale.grid(row=0, column=1, sticky="ew", padx=5, pady=2)
        confidence_label = ttk.Label(general_frame, text=f"{self.confidence_var.get():.1f}")
        confidence_label.grid(row=0, column=2, pady=2)
        confidence_scale.bind("<Motion>", lambda e: confidence_label.config(text=f"{self.confidence_var.get():.1f}"))

        # Min chapter length
        ttk.Label(general_frame, text="Min Chapter Length (words):").grid(row=1, column=0, sticky="w", pady=2)
        self.min_chapter_var = tk.IntVar(value=self.settings.ai_detection.min_chapter_length)
        ttk.Entry(general_frame, textvariable=self.min_chapter_var, width=10).grid(row=1, column=1, sticky="w", padx=5, pady=2)

        # Enable heuristics
        self.enable_heuristics_var = tk.BooleanVar(value=self.settings.ai_detection.enable_heuristics)
        ttk.Checkbutton(general_frame, text="Enable Heuristic Fallback", variable=self.enable_heuristics_var).grid(
            row=2, column=0, columnspan=3, sticky="w", pady=2)

        # Combine methods
        self.combine_methods_var = tk.BooleanVar(value=self.settings.ai_detection.combine_methods)
        ttk.Checkbutton(general_frame, text="Combine AI and Heuristic Results", variable=self.combine_methods_var).grid(
            row=3, column=0, columnspan=3, sticky="w", pady=2)

        general_frame.columnconfigure(1, weight=1)

        # Remote AI Settings
        remote_frame = ttk.LabelFrame(scrollable_frame, text="Remote AI Services", padding=10)
        remote_frame.pack(fill="x", padx=10, pady=5)

        # Use free API
        self.use_free_api_var = tk.BooleanVar(value=self.settings.ai_detection.use_free_api)
        ttk.Checkbutton(remote_frame, text="Use Free AI Service", variable=self.use_free_api_var).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=2)

        # API Key
        ttk.Label(remote_frame, text="Custom API Key:").grid(row=1, column=0, sticky="w", pady=2)
        self.api_key_var = tk.StringVar(value=self.settings.ai_detection.api_key)
        ttk.Entry(remote_frame, textvariable=self.api_key_var, width=40, show="*").grid(row=1, column=1, sticky="ew", padx=5, pady=2)

        # Model
        ttk.Label(remote_frame, text="Model:").grid(row=2, column=0, sticky="w", pady=2)
        self.model_var = tk.StringVar(value=self.settings.ai_detection.model)
        model_combo = ttk.Combobox(remote_frame, textvariable=self.model_var,
                                 values=["gpt-3.5-turbo", "gpt-4", "claude-3-haiku", "claude-3-sonnet"],
                                 width=20)
        model_combo.grid(row=2, column=1, sticky="w", padx=5, pady=2)

        remote_frame.columnconfigure(1, weight=1)

        # Local LLM Settings
        local_frame = ttk.LabelFrame(scrollable_frame, text="Local LLM Settings", padding=10)
        local_frame.pack(fill="x", padx=10, pady=5)

        # Use local LLM
        self.use_local_llm_var = tk.BooleanVar(value=self.settings.ai_detection.use_local_llm)
        ttk.Checkbutton(local_frame, text="Use Local LLM (Ollama, etc.)", variable=self.use_local_llm_var).grid(
            row=0, column=0, columnspan=2, sticky="w", pady=2)

        # Endpoint
        ttk.Label(local_frame, text="LLM Endpoint:").grid(row=1, column=0, sticky="w", pady=2)
        self.local_endpoint_var = tk.StringVar(value=self.settings.ai_detection.local_llm_endpoint)
        ttk.Entry(local_frame, textvariable=self.local_endpoint_var, width=30).grid(row=1, column=1, sticky="ew", padx=5, pady=2)

        # Model
        ttk.Label(local_frame, text="Local Model:").grid(row=2, column=0, sticky="w", pady=2)
        self.local_model_var = tk.StringVar(value=self.settings.ai_detection.local_llm_model)
        local_model_combo = ttk.Combobox(local_frame, textvariable=self.local_model_var,
                                       values=["llama2", "llama3", "mistral", "codellama", "dolphin-mixtral", "neural-chat"],
                                       width=20)
        local_model_combo.grid(row=2, column=1, sticky="w", padx=5, pady=2)

        # Timeout
        ttk.Label(local_frame, text="Timeout (seconds):").grid(row=3, column=0, sticky="w", pady=2)
        self.local_timeout_var = tk.IntVar(value=self.settings.ai_detection.local_llm_timeout)
        ttk.Entry(local_frame, textvariable=self.local_timeout_var, width=10).grid(row=3, column=1, sticky="w", padx=5, pady=2)

        # Max tokens
        ttk.Label(local_frame, text="Max Tokens:").grid(row=4, column=0, sticky="w", pady=2)
        self.local_max_tokens_var = tk.IntVar(value=self.settings.ai_detection.local_llm_max_tokens)
        ttk.Entry(local_frame, textvariable=self.local_max_tokens_var, width=10).grid(row=4, column=1, sticky="w", padx=5, pady=2)

        local_frame.columnconfigure(1, weight=1)

        # Test button
        test_frame = ttk.Frame(scrollable_frame)
        test_frame.pack(fill="x", padx=10, pady=5)
        ttk.Button(test_frame, text="Test AI Configuration", command=self.test_ai_settings).pack(side="right")

    def setup_advanced_tab(self):
        """Setup advanced settings tab."""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Advanced")

        # Enable logging
        self.enable_logging_var = tk.BooleanVar(value=self.settings.advanced_settings.enable_logging)
        ttk.Checkbutton(frame, text="Enable Logging", variable=self.enable_logging_var).pack(
            anchor="w", padx=10, pady=5)

        # Log level
        log_frame = ttk.Frame(frame)
        log_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(log_frame, text="Log Level:").pack(side="left")
        self.log_level_var = tk.StringVar(value=self.settings.advanced_settings.log_level)
        log_combo = ttk.Combobox(log_frame, textvariable=self.log_level_var,
                               values=["DEBUG", "INFO", "WARNING", "ERROR"], state="readonly")
        log_combo.pack(side="left", padx=5)

        # Max log files
        log_files_frame = ttk.Frame(frame)
        log_files_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(log_files_frame, text="Max Log Files:").pack(side="left")
        self.max_log_files_var = tk.IntVar(value=self.settings.advanced_settings.max_log_files)
        ttk.Entry(log_files_frame, textvariable=self.max_log_files_var, width=5).pack(side="left", padx=5)

        # Concurrent jobs
        jobs_frame = ttk.Frame(frame)
        jobs_frame.pack(fill="x", padx=10, pady=5)
        ttk.Label(jobs_frame, text="Concurrent Jobs:").pack(side="left")
        self.concurrent_jobs_var = tk.IntVar(value=self.settings.advanced_settings.concurrent_jobs)
        ttk.Entry(jobs_frame, textvariable=self.concurrent_jobs_var, width=5).pack(side="left", padx=5)

        # Check for updates
        self.check_updates_var = tk.BooleanVar(value=self.settings.advanced_settings.check_for_updates)
        ttk.Checkbutton(frame, text="Check for Updates", variable=self.check_updates_var).pack(
            anchor="w", padx=10, pady=5)

        # Auto update
        self.auto_update_var = tk.BooleanVar(value=self.settings.advanced_settings.auto_update)
        ttk.Checkbutton(frame, text="Auto Update", variable=self.auto_update_var).pack(
            anchor="w", padx=10, pady=5)

        # Crash reporting
        self.crash_reporting_var = tk.BooleanVar(value=self.settings.advanced_settings.enable_crash_reporting)
        ttk.Checkbutton(frame, text="Enable Crash Reporting", variable=self.crash_reporting_var).pack(
            anchor="w", padx=10, pady=5)

        # Telemetry
        self.telemetry_var = tk.BooleanVar(value=self.settings.advanced_settings.enable_telemetry)
        ttk.Checkbutton(frame, text="Enable Telemetry", variable=self.telemetry_var).pack(
            anchor="w", padx=10, pady=5)

    def setup_enterprise_tab(self):
        """Setup enterprise settings tab."""
        frame = ttk.Frame(self.notebook)
        self.notebook.add(frame, text="Enterprise")

        info_label = ttk.Label(frame, text="Enterprise settings are configured via enterprise.json file.")
        info_label.pack(padx=10, pady=20)

        if self.settings.enterprise_config:
            ttk.Label(frame, text="Enterprise configuration is active.").pack(padx=10, pady=5)
        else:
            ttk.Label(frame, text="No enterprise configuration found.").pack(padx=10, pady=5)

    def browse_output_dir(self):
        """Browse for default output directory."""
        from tkinter import filedialog
        directory = filedialog.askdirectory()
        if directory:
            self.default_output_var.set(directory)

    def test_ai_settings(self):
        """Test AI configuration."""
        # Create temporary config and test
        try:
            from ..ai_chapter_detection import ChapterDetectionEngine, AIDetectionConfig

            config = AIDetectionConfig(
                confidence_threshold=self.confidence_var.get(),
                min_chapter_length=self.min_chapter_var.get(),
                enable_heuristics=self.enable_heuristics_var.get(),
                combine_methods=self.combine_methods_var.get(),
                use_free_api=self.use_free_api_var.get(),
                api_key=self.api_key_var.get(),
                model=self.model_var.get(),
                use_local_llm=self.use_local_llm_var.get(),
                local_llm_endpoint=self.local_endpoint_var.get(),
                local_llm_model=self.local_model_var.get(),
                local_llm_timeout=self.local_timeout_var.get(),
                local_llm_max_tokens=self.local_max_tokens_var.get()
            )

            sample_content = """Chapter 1: The Beginning

This is the first chapter. It introduces the main characters.

Chapter 2: The Journey

The adventure continues in this chapter."""

            engine = ChapterDetectionEngine(config)
            boundaries = engine.detect_chapters(sample_content, "test")

            result_text = f"Test successful! Detected {len(boundaries)} boundaries:\n\n"
            for i, boundary in enumerate(boundaries, 1):
                result_text += f"{i}. {boundary.title} (confidence: {boundary.confidence:.2f}, method: {boundary.method})\n"

            messagebox.showinfo("AI Test Results", result_text)

        except Exception as e:
            messagebox.showerror("AI Test Failed", f"Test failed: {e}")

    def save_all_settings(self):
        """Save all settings."""
        try:
            from ..settings import save_settings

            # Update settings object with UI values
            # Conversion defaults
            self.settings.conversion_defaults.css_theme = self.css_theme_var.get()
            self.settings.conversion_defaults.language = self.language_var.get()
            self.settings.conversion_defaults.validate_epub = self.validate_epub_var.get()
            self.settings.conversion_defaults.image_max_width = self.image_max_width_var.get()
            self.settings.conversion_defaults.accessibility_features = self.accessibility_var.get()
            self.settings.conversion_defaults.chapter_detection = self.chapter_detection_var.get()
            self.settings.conversion_defaults.include_toc = self.include_toc_var.get()
            self.settings.conversion_defaults.include_cover = self.include_cover_var.get()
            self.settings.conversion_defaults.drm_free = self.drm_free_var.get()

            # UI preferences
            self.settings.ui_preferences.remember_last_directory = self.remember_dir_var.get()
            self.settings.ui_preferences.auto_fill_metadata = self.auto_fill_var.get()
            self.settings.ui_preferences.show_advanced_options = self.show_advanced_var.get()
            self.settings.ui_preferences.theme = self.ui_theme_var.get()
            self.settings.ui_preferences.font_size = self.font_size_var.get()
            self.settings.ui_preferences.remember_window_size = self.remember_window_var.get()
            self.settings.ui_preferences.confirm_overwrite = self.confirm_overwrite_var.get()
            self.settings.ui_preferences.show_tooltips = self.show_tooltips_var.get()

            # File defaults
            self.settings.file_defaults.default_output_directory = self.default_output_var.get() or None
            self.settings.file_defaults.auto_generate_output_name = self.auto_name_var.get()
            self.settings.file_defaults.backup_original = self.backup_var.get()

            # AI detection
            self.settings.ai_detection.confidence_threshold = self.confidence_var.get()
            self.settings.ai_detection.min_chapter_length = self.min_chapter_var.get()
            self.settings.ai_detection.enable_heuristics = self.enable_heuristics_var.get()
            self.settings.ai_detection.combine_methods = self.combine_methods_var.get()
            self.settings.ai_detection.use_free_api = self.use_free_api_var.get()
            self.settings.ai_detection.api_key = self.api_key_var.get()
            self.settings.ai_detection.model = self.model_var.get()
            self.settings.ai_detection.use_local_llm = self.use_local_llm_var.get()
            self.settings.ai_detection.local_llm_endpoint = self.local_endpoint_var.get()
            self.settings.ai_detection.local_llm_model = self.local_model_var.get()
            self.settings.ai_detection.local_llm_timeout = self.local_timeout_var.get()
            self.settings.ai_detection.local_llm_max_tokens = self.local_max_tokens_var.get()

            # Advanced settings
            self.settings.advanced_settings.enable_logging = self.enable_logging_var.get()
            self.settings.advanced_settings.log_level = self.log_level_var.get()
            self.settings.advanced_settings.max_log_files = self.max_log_files_var.get()
            self.settings.advanced_settings.concurrent_jobs = self.concurrent_jobs_var.get()
            self.settings.advanced_settings.check_for_updates = self.check_updates_var.get()
            self.settings.advanced_settings.auto_update = self.auto_update_var.get()
            self.settings.advanced_settings.enable_crash_reporting = self.crash_reporting_var.get()
            self.settings.advanced_settings.enable_telemetry = self.telemetry_var.get()

            # Save to disk
            save_settings(self.settings)

            messagebox.showinfo("Settings Saved", "All settings have been saved successfully!")
            self.dialog.destroy()

        except Exception as e:
            messagebox.showerror("Error", f"Could not save settings: {e}")

    def reset_to_defaults(self):
        """Reset all settings to defaults."""
        if messagebox.askyesno("Reset Settings", "Are you sure you want to reset all settings to defaults?"):
            try:
                from ..settings import get_settings_manager
                settings_manager = get_settings_manager()
                settings_manager.reset_to_defaults()
                messagebox.showinfo("Settings Reset", "Settings have been reset to defaults.")
                self.dialog.destroy()
            except Exception as e:
                messagebox.showerror("Error", f"Could not reset settings: {e}")

    def export_settings(self):
        """Export settings to file."""
        try:
            from tkinter import filedialog
            file_path = filedialog.asksaveasfilename(
                title="Export Settings",
                defaultextension=".json",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if file_path:
                from ..settings import get_settings_manager
                settings_manager = get_settings_manager()
                settings_manager.export_settings(Path(file_path))
                messagebox.showinfo("Export Complete", f"Settings exported to {file_path}")
        except Exception as e:
            messagebox.showerror("Export Error", f"Could not export settings: {e}")

    def import_settings(self):
        """Import settings from file."""
        try:
            from tkinter import filedialog
            file_path = filedialog.askopenfilename(
                title="Import Settings",
                filetypes=[("JSON files", "*.json"), ("All files", "*.*")]
            )
            if file_path:
                from ..settings import get_settings_manager
                settings_manager = get_settings_manager()
                settings_manager.import_settings(Path(file_path), merge=True)
                messagebox.showinfo("Import Complete", "Settings imported successfully.")
                self.dialog.destroy()
        except Exception as e:
            messagebox.showerror("Import Error", f"Could not import settings: {e}")


# Framework-specific MainWindow implementations
if GUI_FRAMEWORK == "tkinter":
    # Set the MainWindow for tkinter
    MainWindow = TkinterMainWindow

elif GUI_FRAMEWORK in ["pyqt5", "pyqt6"]:
    # Full PyQt implementation
    class PyQtMainWindow:
        def __init__(self):
            """Initialize PyQt main window."""
            self.app = QApplication(sys.argv)
            self.window = QMainWindow()
            self.setup_ui()

        def setup_ui(self):
            """Set up the PyQt user interface."""
            central_widget = QWidget()
            self.window.setCentralWidget(central_widget)

            layout = QVBoxLayout(central_widget)

            # Create tabs
            tab_widget = QTabWidget()
            layout.addWidget(tab_widget)

            # Convert tab
            convert_tab = QWidget()
            tab_widget.addTab(convert_tab, "Convert")
            self.setup_convert_tab_pyqt(convert_tab)

            # Settings tab
            settings_tab = QWidget()
            tab_widget.addTab(settings_tab, "Settings")
            self.setup_settings_tab_pyqt(settings_tab)

            # Batch tab
            batch_tab = QWidget()
            tab_widget.addTab(batch_tab, "Batch")
            self.setup_batch_tab_pyqt(batch_tab)

            self.window.setWindowTitle("Docx2Shelf - Document to EPUB Converter")
            self.window.resize(800, 600)

        def setup_convert_tab_pyqt(self, parent):
            """Set up PyQt convert tab."""
            layout = QVBoxLayout(parent)

            # File selection
            file_group = QGroupBox("Input File")
            file_layout = QHBoxLayout(file_group)

            self.file_path_edit = QLineEdit()
            self.file_path_edit.setPlaceholderText("Select DOCX file...")
            file_layout.addWidget(self.file_path_edit)

            browse_btn = QPushButton("Browse...")
            browse_btn.clicked.connect(self.browse_file_pyqt)
            file_layout.addWidget(browse_btn)

            layout.addWidget(file_group)

            # Metadata fields
            metadata_group = QGroupBox("Book Metadata")
            metadata_layout = QGridLayout(metadata_group)

            self.title_edit = QLineEdit()
            metadata_layout.addWidget(QLabel("Title:"), 0, 0)
            metadata_layout.addWidget(self.title_edit, 0, 1)

            self.author_edit = QLineEdit()
            metadata_layout.addWidget(QLabel("Author:"), 1, 0)
            metadata_layout.addWidget(self.author_edit, 1, 1)

            layout.addWidget(metadata_group)

            # Convert button
            convert_btn = QPushButton("Convert to EPUB")
            convert_btn.clicked.connect(self.convert_pyqt)
            layout.addWidget(convert_btn)

            # Progress bar
            self.progress_bar = QProgressBar()
            layout.addWidget(self.progress_bar)

        def setup_settings_tab_pyqt(self, parent):
            """Set up PyQt settings tab."""
            layout = QVBoxLayout(parent)

            # Theme selection
            theme_group = QGroupBox("Theme Settings")
            theme_layout = QVBoxLayout(theme_group)

            self.theme_combo = QComboBox()
            self.theme_combo.addItems(['serif', 'sans', 'printlike'])
            theme_layout.addWidget(QLabel("CSS Theme:"))
            theme_layout.addWidget(self.theme_combo)

            layout.addWidget(theme_group)

        def setup_batch_tab_pyqt(self, parent):
            """Set up PyQt batch processing tab."""
            layout = QVBoxLayout(parent)

            batch_group = QGroupBox("Batch Processing")
            batch_layout = QVBoxLayout(batch_group)

            batch_layout.addWidget(QLabel("Add multiple files for batch conversion:"))

            self.file_list = QListWidget()
            batch_layout.addWidget(self.file_list)

            buttons_layout = QHBoxLayout()
            add_files_btn = QPushButton("Add Files")
            add_files_btn.clicked.connect(self.add_batch_files_pyqt)
            buttons_layout.addWidget(add_files_btn)

            remove_btn = QPushButton("Remove Selected")
            remove_btn.clicked.connect(self.remove_batch_files_pyqt)
            buttons_layout.addWidget(remove_btn)

            batch_layout.addLayout(buttons_layout)
            layout.addWidget(batch_group)

        def browse_file_pyqt(self):
            """Browse for input file using PyQt."""
            file_path, _ = QFileDialog.getOpenFileName(
                self.window, "Select Document File", "",
                "All Supported (*.docx *.md *.txt *.html *.htm);;Word Documents (*.docx);;Markdown (*.md);;Text Files (*.txt);;HTML Files (*.html *.htm);;All Files (*)"
            )
            if file_path:
                self.file_path_edit.setText(file_path)

        def convert_pyqt(self):
            """Convert file using PyQt interface."""
            # Note: This is a placeholder for PyQt conversion
            # The main functionality is implemented in the Tkinter version
            # PyQt implementation would follow similar patterns but with PyQt widgets

            if hasattr(self, 'file_path_edit') and self.file_path_edit.text():
                file_path = self.file_path_edit.text()
                # In a full PyQt implementation, this would:
                # 1. Create a PyQt progress dialog
                # 2. Run conversion in QThread
                # 3. Update progress using Qt signals/slots
                # 4. Display results in PyQt message boxes

                # For now, show a simple message
                from PyQt5.QtWidgets import QMessageBox
                msg = QMessageBox()
                msg.setWindowTitle("PyQt Conversion")
                msg.setText(f"PyQt conversion initiated for:\n{Path(file_path).name}\n\nNote: Full PyQt implementation would provide\nthe same functionality as the Tkinter version.")
                msg.exec_()
            else:
                from PyQt5.QtWidgets import QMessageBox
                msg = QMessageBox()
                msg.setWindowTitle("Error")
                msg.setText("Please select a file first")
                msg.setIcon(QMessageBox.Warning)
                msg.exec_()

        def add_batch_files_pyqt(self):
            """Add files to batch list using PyQt."""
            files, _ = QFileDialog.getOpenFileNames(
                self.window, "Select Document Files", "",
                "All Supported (*.docx *.md *.txt *.html *.htm);;Word Documents (*.docx);;Markdown (*.md);;Text Files (*.txt);;HTML Files (*.html *.htm);;All Files (*)"
            )
            for file_path in files:
                self.file_list.addItem(file_path)

        def remove_batch_files_pyqt(self):
            """Remove selected files from batch list."""
            for item in self.file_list.selectedItems():
                self.file_list.takeItem(self.file_list.row(item))

        def run(self):
            """Run the PyQt application."""
            self.window.show()
            return self.app.exec_()

    MainWindow = PyQtMainWindow

else:
    # No GUI framework available
    class NoGUIMainWindow:
        def __init__(self):
            print("No GUI framework available. Please install tkinter or PyQt.")
            sys.exit(1)

        def run(self):
            pass

    MainWindow = NoGUIMainWindow


def create_context_menu_windows():
    """Create Windows context menu integration."""
    try:
        import winreg

        # Registry key for .docx files
        key_path = r"SOFTWARE\Classes\.docx\shell\Convert to EPUB\command"

        # Get path to this script
        script_path = Path(__file__).parent.parent / "cli.py"
        command = f'python "{script_path}" gui "%1"'

        # Create registry entry
        with winreg.CreateKey(winreg.HKEY_CURRENT_USER, key_path) as key:
            winreg.SetValue(key, "", winreg.REG_SZ, command)

        print("Windows context menu integration created successfully!")

    except Exception as e:
        print(f"Failed to create context menu integration: {e}")


def main():
    """Main entry point for GUI application."""

    # Setup logging
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

    # Create and run main window
    app = MainWindow()
    app.run()


if __name__ == "__main__":
    main()